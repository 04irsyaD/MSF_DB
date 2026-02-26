"""
================================================================================
SQL Table to Documentation Generator
================================================================================
Script untuk generate dokumentasi otomatis dari CREATE TABLE menggunakan AI lokal (Ollama)

Author  : 04irsyaD
Created : 2026-01-28
Project : PLSTX Tables Documentation

Usage:
    python sql_table_to_docs.py

Requirements:
    - Ollama installed dan running
    - Model AI sudah di-download (ollama pull llama3)
    - pip install -r requirements.txt
================================================================================
"""

import re
import os
from pathlib import Path
from datetime import datetime

try:
    import ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False
    print("⚠️  Package 'ollama' tidak ditemukan. Install dengan: pip install ollama")

try:
    from rich.console import Console
    from rich.progress import Progress, SpinnerColumn, TextColumn
    from rich.panel import Panel
    from rich.table import Table
    RICH_AVAILABLE = True
    console = Console()
except ImportError:
    RICH_AVAILABLE = False
    console = None

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import psycopg2
    PSYCOPG2_AVAILABLE = True
except ImportError:
    PSYCOPG2_AVAILABLE = False
    print("⚠️  Package 'psycopg2' tidak ditemukan. Install dengan: pip install psycopg2-binary")

import json


# ============================================
# CONFIGURATION
# ============================================
# Optimized for NVIDIA RTX 3050 (4GB VRAM)
CONFIG = {
    "sql_file": r"..\doc.sql",  # Relative path ke file SQL
    "output_dir": r".\output",
    
    # Model Ollama - Pakai yang sudah terinstall
    "model": "llama3:latest",         # 🔥 4.7GB - Sudah ada, bagus untuk dokumentasi
    # "model": "deepseek-r1:8b",      # 5.2GB - Reasoning model, lebih lambat tapi detail
    
    "output_formats": ["md", "docx"],  # Format output: md, docx, atau keduanya
    
    # Word Document Settings
    "word_font": {
        "name": "DM Sans",          # Font utama: DM Sans (Google Font)
        "size": 12,                 # Ukuran font body (pt)
        "heading_font": "DM Sans",  # Font untuk heading
        "code_font": "Consolas",    # Font untuk SQL code
        "code_size": 8,             # Ukuran font code (pt)
    },
    
    # GPU Settings untuk RTX 3050
    "gpu_layers": 35,  # Optimal untuk 4GB VRAM
    "num_ctx": 2048,   # Context window
    
    # ==========================================
    # KNOWLEDGE BASE DATABASE
    # ==========================================
    "use_db_knowledge": True,  # Set False untuk skip DB, langsung AI
    
    "db_config": {
        "host": "localhost",
        "port": 5414,
        "database": "knowledge_db",    # ← Ganti dengan nama database kamu
        "user": "postgres",            # ← Ganti dengan username
        "password": "1234",    # ← Ganti dengan password
    },
}


# ============================================
# KNOWLEDGE BASE FROM DATABASE
# ============================================
KNOWLEDGE_CACHE = {}  # Cache untuk mengurangi query DB
PREFIX_PATTERNS = {}  # Cache untuk prefix patterns (t_m_, L_, dll)

def load_knowledge_base_from_db():
    """
    Load semua knowledge base dari PostgreSQL ke cache
    Dipanggil sekali di awal untuk performa
    """
    global KNOWLEDGE_CACHE, PREFIX_PATTERNS
    
    if not CONFIG.get("use_db_knowledge", False):
        print("⚠️  Knowledge base DB dinonaktifkan (use_db_knowledge=False)")
        return False
    
    if not PSYCOPG2_AVAILABLE:
        print("⚠️  psycopg2 tidak tersedia, skip knowledge base DB")
        return False
    
    try:
        conn = psycopg2.connect(**CONFIG["db_config"])
        cur = conn.cursor()
        
        # Check if table exists
        cur.execute("""
            SELECT EXISTS (
                SELECT FROM information_schema.tables 
                WHERE table_name = 'doc_knowledge_base'
            )
        """)
        
        if not cur.fetchone()[0]:
            print("⚠️  Tabel doc_knowledge_base belum ada di database")
            print("   Jalankan create_knowledge_base.sql terlebih dahulu")
            cur.close()
            conn.close()
            return False
        
        # Load semua data ke cache
        cur.execute("""
            SELECT table_name, column_name, description 
            FROM doc_knowledge_base
            ORDER BY table_name, column_name
        """)
        
        KNOWLEDGE_CACHE = {
            'tables': {},       # {table_name: description}
            'columns': {},      # {table_name: {column_name: description}}
            'common': {}        # {column_name: description} untuk table_name='*'
        }
        
        row_count = 0
        for row in cur.fetchall():
            table_name, column_name, description = row
            row_count += 1
            
            if table_name == '*':
                # Common column
                KNOWLEDGE_CACHE['common'][column_name] = description
            elif column_name is None:
                # Table description
                KNOWLEDGE_CACHE['tables'][table_name] = description
            else:
                # Column description
                if table_name not in KNOWLEDGE_CACHE['columns']:
                    KNOWLEDGE_CACHE['columns'][table_name] = {}
                KNOWLEDGE_CACHE['columns'][table_name][column_name] = description
        
        # Load prefix patterns
        cur.execute("""
            SELECT EXISTS (
                SELECT FROM information_schema.tables 
                WHERE table_name = 'doc_prefix_patterns'
            )
        """)
        
        if cur.fetchone()[0]:
            cur.execute("""
                SELECT prefix, description 
                FROM doc_prefix_patterns
                ORDER BY LENGTH(prefix) DESC
            """)
            for row in cur.fetchall():
                PREFIX_PATTERNS[row[0]] = row[1]
        
        cur.close()
        conn.close()
        
        table_count = len(KNOWLEDGE_CACHE['tables'])
        common_count = len(KNOWLEDGE_CACHE['common'])
        specific_count = sum(len(cols) for cols in KNOWLEDGE_CACHE['columns'].values())
        prefix_count = len(PREFIX_PATTERNS)
        
        print(f"✅ Knowledge base loaded dari DB:")
        print(f"   📋 {table_count} deskripsi tabel")
        print(f"   📝 {specific_count} deskripsi kolom spesifik")
        print(f"   🔄 {common_count} common columns")
        print(f"   🏷️  {prefix_count} prefix patterns")
        
        return True
        
    except psycopg2.OperationalError as e:
        print(f"❌ Gagal connect ke database: {e}")
        print("   Cek konfigurasi db_config di CONFIG")
        return False
    except Exception as e:
        print(f"❌ Error loading knowledge base: {e}")
        return False


def get_description_from_db(table_name: str, column_name: str = None) -> str:
    """
    Ambil deskripsi dari cache (yang diload dari DB)
    
    Prioritas:
    1. Tabel/kolom spesifik
    2. Common columns (untuk kolom)
    3. Prefix patterns (untuk tabel)
    4. Return None → AI akan handle
    """
    if not KNOWLEDGE_CACHE:
        return None
    
    if column_name is None:
        # Ambil deskripsi tabel
        # 1. Cek spesifik dulu
        specific = KNOWLEDGE_CACHE.get('tables', {}).get(table_name)
        if specific:
            return specific
        
        # 2. Cek prefix patterns
        return get_table_prefix_description(table_name)
    
    # Ambil deskripsi kolom
    # 1. Cek spesifik dulu
    specific = KNOWLEDGE_CACHE.get('columns', {}).get(table_name, {}).get(column_name)
    if specific:
        return specific
    
    # 2. Fallback ke common
    return KNOWLEDGE_CACHE.get('common', {}).get(column_name)


def get_table_prefix_description(table_name: str) -> str:
    """
    Generate deskripsi tabel berdasarkan prefix dari DB
    
    Contoh:
    - t_m_division → "Tabel Master - division"
    - L_FAQ → "Log Data - FAQ"
    """
    if not PREFIX_PATTERNS:
        return None
    
    table_lower = table_name.lower()
    
    # Cari prefix yang match (sudah diurutkan dari terpanjang)
    for prefix, description in PREFIX_PATTERNS.items():
        if table_lower.startswith(prefix.lower()):
            # Ambil nama tanpa prefix
            name_part = table_name[len(prefix):]
            # Format: "Tabel Master - division" atau "Tabel Master" kalau nama kosong
            if name_part:
                return f"{description} - {name_part.replace('_', ' ').title()}"
            return description
    
    return None


# ============================================
# SQL PARSER FOR TABLES
# ============================================
def extract_tables_from_sql(sql_file: str) -> list[dict]:
    """
    Extract semua CREATE TABLE dari file SQL
    
    Args:
        sql_file: Path ke file SQL
        
    Returns:
        List of dict dengan keys: name, columns, constraints, definition
    """
    with open(sql_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    tables = []
    
    # Pattern yang lebih baik - capture sampai tutup kurung terakhir dan semicolon
    # Menggunakan greedy matching untuk table body
    pattern = r'[Cc][Rr][Ee][Aa][Tt][Ee]\s+[Tt][Aa][Bb][Ll][Ee]\s+(?:public\.)?"?(\w+)"?\s*\(([\s\S]*?)\)(?:\s*;)?'
    
    # Cari semua CREATE TABLE dengan cara yang lebih robust
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Check for CREATE TABLE
        create_match = re.match(r'^\s*[Cc][Rr][Ee][Aa][Tt][Ee]\s+[Tt][Aa][Bb][Ll][Ee]\s+(?:public\.)?"?(\w+)"?\s*\(', line, re.IGNORECASE)
        if create_match:
            table_name = create_match.group(1)
            
            # Collect full table definition
            full_def = [line]
            depth = line.count('(') - line.count(')')
            i += 1
            
            while i < len(lines) and depth > 0:
                full_def.append(lines[i])
                depth += lines[i].count('(') - lines[i].count(')')
                i += 1
            
            # Join and parse
            table_sql = '\n'.join(full_def)
            
            # Extract body between first ( and last )
            body_match = re.search(r'\(([\s\S]*)\)', table_sql)
            if body_match:
                table_body = body_match.group(1)
                
                # Parse columns dan constraints
                columns = []
                constraints = []
                
                parts = split_table_body(table_body)
                
                for part in parts:
                    part = part.strip()
                    if not part:
                        continue
                    
                    if re.match(r'^\s*(CONSTRAINT|PRIMARY\s+KEY|FOREIGN\s+KEY|UNIQUE|CHECK|constraint)', part, re.IGNORECASE):
                        constraints.append(parse_constraint(part))
                    else:
                        col = parse_column(part)
                        if col:
                            columns.append(col)
                
                tables.append({
                    'name': table_name,
                    'columns': columns,
                    'constraints': constraints,
                    'definition': table_sql.strip()
                })
        else:
            i += 1
    
    return tables


def split_table_body(body: str) -> list[str]:
    """Split table body by comma, respecting parentheses"""
    parts = []
    current = ""
    depth = 0
    
    for char in body:
        if char == '(':
            depth += 1
            current += char
        elif char == ')':
            depth -= 1
            current += char
        elif char == ',' and depth == 0:
            parts.append(current.strip())
            current = ""
        else:
            current += char
    
    if current.strip():
        parts.append(current.strip())
    
    return parts


def parse_column(col_str: str) -> dict:
    """Parse column definition"""
    col_str = col_str.strip()
    
    # Skip if it looks like a constraint
    if re.match(r'^\s*(CONSTRAINT|PRIMARY|FOREIGN|UNIQUE|CHECK)', col_str, re.IGNORECASE):
        return None
    
    # Pattern: column_name data_type [constraints]
    # Handle quoted column names
    match = re.match(r'^"?(\w+)"?\s+(\w+(?:\([^)]+\))?)\s*(.*)?$', col_str, re.IGNORECASE)
    
    if match:
        name = match.group(1)
        data_type = match.group(2)
        rest = match.group(3) or ""
        
        return {
            'name': name,
            'data_type': data_type,
            'rest': rest  # Contains NULL/NOT NULL, DEFAULT, etc.
        }
    
    return None


def parse_constraint(const_str: str) -> dict:
    """Parse constraint definition"""
    const_str = const_str.strip()
    
    # PRIMARY KEY
    pk_match = re.search(r'CONSTRAINT\s+"?(\w+)"?\s+PRIMARY\s+KEY\s*\(([^)]+)\)', const_str, re.IGNORECASE)
    if pk_match:
        return {
            'name': pk_match.group(1),
            'type': 'PRIMARY KEY',
            'columns': pk_match.group(2)
        }
    
    # FOREIGN KEY
    fk_match = re.search(r'(?:CONSTRAINT\s+"?(\w+)"?\s+)?FOREIGN\s+KEY\s*\(([^)]+)\)\s*REFERENCES\s+"?(\w+)"?\s*\(([^)]+)\)', const_str, re.IGNORECASE)
    if fk_match:
        return {
            'name': fk_match.group(1) or 'unnamed_fk',
            'type': 'FOREIGN KEY',
            'columns': fk_match.group(2),
            'ref_table': fk_match.group(3),
            'ref_columns': fk_match.group(4)
        }
    
    # Simple constraint match
    simple_match = re.match(r'CONSTRAINT\s+"?(\w+)"?\s+(\w+)', const_str, re.IGNORECASE)
    if simple_match:
        return {
            'name': simple_match.group(1),
            'type': simple_match.group(2).upper(),
            'detail': const_str
        }
    
    return {
        'name': 'unknown',
        'type': 'UNKNOWN',
        'detail': const_str
    }


# ============================================
# AI DOCUMENTATION GENERATOR (DB + AI)
# ============================================
def generate_table_documentation_with_ai(table: dict, model: str = "llama3:latest") -> dict:
    """
    Generate dokumentasi dengan prioritas:
    1. Knowledge Base dari DB
    2. AI (Ollama) untuk yang tidak ada di DB
    3. Fallback basic
    
    Returns:
        Dict dengan keys: table_description, columns
    """
    table_name = table['name']
    
    # ==========================================
    # STEP 1: Cek Knowledge Base (DB) dulu
    # ==========================================
    db_table_desc = get_description_from_db(table_name)
    db_columns = {}
    need_ai_columns = []  # Kolom yang perlu AI generate
    
    for col in table['columns']:
        db_col_desc = get_description_from_db(table_name, col['name'])
        if db_col_desc:
            db_columns[col['name']] = db_col_desc
        else:
            need_ai_columns.append(col)
    
    # Hitung statistik
    total_cols = len(table['columns'])
    db_cols_count = len(db_columns)
    ai_cols_count = len(need_ai_columns)
    
    # Jika SEMUA sudah ada di DB, skip AI
    if db_table_desc and ai_cols_count == 0:
        if RICH_AVAILABLE:
            console.print(f"     [green]📊 100% dari DB[/green]")
        else:
            print(f"     📊 100% dari DB")
        return {
            'table_description': db_table_desc,
            'columns': db_columns,
            'source': 'db'
        }
    
    # ==========================================
    # STEP 2: Generate dengan AI (untuk yang tidak ada)
    # ==========================================
    if not OLLAMA_AVAILABLE:
        # Gabung DB + fallback
        result = generate_table_documentation_basic(table)
        if db_table_desc:
            result['table_description'] = db_table_desc
        for col_name, desc in db_columns.items():
            result['columns'][col_name] = desc
        return result
    
    # Build prompt HANYA untuk kolom yang perlu AI
    columns_for_ai = "\n".join([
        f"- {col['name']} ({col['data_type']})" 
        for col in need_ai_columns
    ])
    
    # Konteks dari DB jika ada
    context = db_table_desc if db_table_desc else f"Tabel {table_name}"
    
    prompt = f"""Kamu adalah database documentation expert. Analisis PostgreSQL table berikut dan buat dokumentasi dalam Bahasa Indonesia.

NAMA TABLE: {table_name}
KONTEKS: {context}

KOLOM YANG PERLU DESKRIPSI:
{columns_for_ai}

Buatkan dokumentasi dengan format JSON seperti ini:
{{
    "table_description": "Deskripsi singkat 1 kalimat saja",
    "columns": {{
        "nama_kolom1": "Deskripsi singkat max 5 kata",
        "nama_kolom2": "Deskripsi singkat max 5 kata"
    }}
}}

PENTING:
- Jawab HANYA dalam format JSON yang valid
- Deskripsi dalam Bahasa Indonesia
- DESKRIPSI KOLOM HARUS SINGKAT (maksimal 5-8 kata)
"""

    try:
        response = ollama.chat(
            model=model, 
            messages=[{'role': 'user', 'content': prompt}],
            options={
                'num_gpu': CONFIG.get('gpu_layers', 35),
                'num_ctx': CONFIG.get('num_ctx', 2048),
                'temperature': 0.3,
            }
        )
        
        response_text = response['message']['content']
        
        # Parse JSON dari response
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            try:
                ai_result = json.loads(json_match.group())
                
                # ==========================================
                # STEP 3: Gabungkan DB + AI
                # ==========================================
                final_columns = db_columns.copy()
                final_columns.update(ai_result.get('columns', {}))
                
                # Table desc: prioritas DB
                final_desc = db_table_desc or ai_result.get('table_description', f"Tabel {table_name}")
                
                # Print statistik
                db_pct = int(db_cols_count / total_cols * 100) if total_cols > 0 else 0
                if RICH_AVAILABLE:
                    console.print(f"     [cyan]📊 DB: {db_pct}% ({db_cols_count}) | 🤖 AI: {100-db_pct}% ({ai_cols_count})[/cyan]")
                else:
                    print(f"     📊 DB: {db_pct}% ({db_cols_count}) | 🤖 AI: {100-db_pct}% ({ai_cols_count})")
                
                return {
                    'table_description': final_desc,
                    'columns': final_columns,
                    'source': 'db+ai' if db_columns else 'ai'
                }
            except json.JSONDecodeError:
                pass
        
        # Fallback
        result = generate_table_documentation_basic(table, response_text)
        if db_table_desc:
            result['table_description'] = db_table_desc
        for col_name, desc in db_columns.items():
            result['columns'][col_name] = desc
        return result
        
    except Exception as e:
        print(f"⚠️  Error calling Ollama: {e}")
        result = generate_table_documentation_basic(table)
        if db_table_desc:
            result['table_description'] = db_table_desc
        for col_name, desc in db_columns.items():
            result['columns'][col_name] = desc
        return result


def generate_table_documentation_basic(table: dict, ai_response: str = None) -> dict:
    """Generate dokumentasi basic tanpa AI (fallback)"""
    
    # Common column descriptions - SINGKAT
    common_cols = {
        'id': 'Primary key unik',
        'created_at': 'Waktu pembuatan record',
        'updated_at': 'Waktu update terakhir',
        'deleted_at': 'Waktu soft delete',
        'is_active': 'Status aktif record',
        'created_by_id': 'ID user pembuat',
        'updated_by_id': 'ID user yang update',
        'order_data': 'Urutan tampilan data',
        'email': 'Alamat email',
        'username': 'Nama pengguna',
        'password': 'Password terenkripsi',
        'phone_number': 'Nomor telepon',
        'note': 'Catatan tambahan',
    }
    
    columns_desc = {}
    for col in table['columns']:
        col_name = col['name'].lower()
        if col_name in common_cols:
            columns_desc[col['name']] = common_cols[col_name]
        elif '_id' in col_name:
            columns_desc[col['name']] = f"FK ke tabel terkait"
        elif '_at' in col_name:
            columns_desc[col['name']] = f"Timestamp {col['name'].replace('_', ' ')}"
        elif '_name' in col_name or '_nm' in col_name:
            columns_desc[col['name']] = f"Nama {col['name'].replace('_name', '').replace('_nm', '')}"
        elif '_code' in col_name or '_cd' in col_name:
            columns_desc[col['name']] = f"Kode {col['name'].replace('_code', '').replace('_cd', '')}"
        elif '_path' in col_name:
            columns_desc[col['name']] = f"Path file {col['name'].replace('_path', '')}"
        else:
            columns_desc[col['name']] = f"Field {col['name'].replace('_', ' ')}"
    
    # Try to extract description from AI response if available
    desc = f"Tabel {table['name']} untuk menyimpan data terkait sistem."
    if ai_response:
        # Try to find a description sentence
        lines = ai_response.split('\n')
        for line in lines:
            if 'deskripsi' in line.lower() or 'tabel' in line.lower():
                desc = line.strip()
                break
    
    return {
        'table_description': desc,
        'columns': columns_desc
    }


# ============================================
# OUTPUT GENERATORS
# ============================================
def save_as_markdown(tables_docs: list[dict], output_path: str):
    """Save dokumentasi ke file Markdown"""
    
    content = f"""# 📚 Dokumentasi Database Schema - PLSTX System

> **Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  
> **Source:** sm.sql  
> **Total Tables:** {len(tables_docs)}

---

## 📋 Daftar Tables

| No | Nama Table | Deskripsi |
|----|------------|-----------|
"""
    
    for i, table in enumerate(tables_docs, 1):
        desc = table.get('ai_doc', {}).get('table_description', '-')
        # Truncate for table view
        if len(desc) > 80:
            desc = desc[:77] + "..."
        content += f"| {i} | [`{table['name']}`](#{table['name'].lower().replace('_', '-')}) | {desc} |\n"
    
    content += "\n---\n\n"
    
    for i, table in enumerate(tables_docs, 1):
        ai_doc = table.get('ai_doc', {})
        columns_desc = ai_doc.get('columns', {})
        
        content += f"""## {i}. `{table['name']}`

### Deskripsi
{ai_doc.get('table_description', 'Dokumentasi belum tersedia.')}

### Kolom

| No | Nama Field | Tipe Data | Deskripsi |
|----|------------|-----------|-----------|
"""
        
        for j, col in enumerate(table['columns'], 1):
            col_desc = columns_desc.get(col['name'], f"Field {col['name']}")
            content += f"| {j} | {col['name']} | {col['data_type']} | {col_desc} |\n"
        
        # Constraints
        if table['constraints']:
            content += "\n### Constraints\n\n"
            content += "| Nama | Tipe | Detail |\n"
            content += "|------|------|--------|\n"
            
            for const in table['constraints']:
                if const['type'] == 'PRIMARY KEY':
                    detail = f"({const.get('columns', '')})"
                elif const['type'] == 'FOREIGN KEY':
                    detail = f"{const.get('columns', '')} → {const.get('ref_table', '')}({const.get('ref_columns', '')})"
                else:
                    detail = const.get('detail', '-')[:50]
                
                content += f"| {const['name']} | {const['type']} | {detail} |\n"
        
        # Foreign Key Relations
        fk_constraints = [c for c in table['constraints'] if c['type'] == 'FOREIGN KEY']
        if fk_constraints:
            content += "\n### Relasi (Foreign Keys)\n\n"
            content += "```\n"
            content += f"{table['name']}\n"
            for fk in fk_constraints:
                content += f"    └──> {fk.get('ref_table', '')} ({fk.get('columns', '')})\n"
            content += "```\n"
        
        # SQL Definition
        content += f"""
<details>
<summary>📝 <strong>SQL Definition</strong></summary>

```sql
{table['definition']}
```

</details>

---

"""
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return output_path


def save_as_docx(tables_docs: list[dict], output_path: str):
    """Save dokumentasi ke file Word (.docx)"""
    
    if not DOCX_AVAILABLE:
        print("⚠️  python-docx tidak tersedia. Skip export DOCX.")
        return None
    
    doc = Document()
    
    # Get font settings from config
    font_cfg = CONFIG.get("word_font", {})
    main_font = font_cfg.get("name", "Arial")
    main_size = font_cfg.get("size", 11)
    heading_font = font_cfg.get("heading_font", "Arial")
    code_font = font_cfg.get("code_font", "Consolas")
    code_size = font_cfg.get("code_size", 8)
    
    # Set default font for document
    style = doc.styles['Normal']
    style.font.name = main_font
    style.font.size = Pt(main_size)
    
    # Set heading fonts
    for i in range(10):
        try:
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = heading_font
        except KeyError:
            pass
    
    # Title
    title = doc.add_heading('Dokumentasi Database Schema - PLSTX System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Tables: {len(tables_docs)}")
    doc.add_paragraph("─" * 50)
    
    # Table of Contents
    doc.add_heading('Daftar Tables', level=1)
    toc_table = doc.add_table(rows=1, cols=3)
    toc_table.style = 'Table Grid'
    hdr_cells = toc_table.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Nama Table'
    hdr_cells[2].text = 'Deskripsi'
    
    for i, table in enumerate(tables_docs, 1):
        row_cells = toc_table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = table['name']
        desc = table.get('ai_doc', {}).get('table_description', '-')
        row_cells[2].text = desc[:60] + "..." if len(desc) > 60 else desc
    
    doc.add_page_break()
    
    # Detail each table
    for i, table in enumerate(tables_docs, 1):
        ai_doc = table.get('ai_doc', {})
        columns_desc = ai_doc.get('columns', {})
        
        doc.add_heading(f"{i}. {table['name']}", level=1)
        
        # Description
        doc.add_heading('Deskripsi', level=2)
        doc.add_paragraph(ai_doc.get('table_description', 'Dokumentasi belum tersedia.'))
        
        # Columns table
        doc.add_heading('Kolom', level=2)
        col_table = doc.add_table(rows=1, cols=4)
        col_table.style = 'Table Grid'
        
        # Set column widths (px to cm conversion: 1px ≈ 0.0264cm at 96dpi)
        # No: 60px ≈ 1.6cm, Nama Field: 200px ≈ 5.3cm, Tipe Data: 170px ≈ 4.5cm, Deskripsi: 210px ≈ 5.5cm
        col_widths = [Cm(1.5), Cm(5.0), Cm(4.5), Cm(5.5)]
        
        # Apply column widths
        for row in col_table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = col_widths[idx]
        
        hdr = col_table.rows[0].cells
        hdr[0].text = 'No'
        hdr[1].text = 'Nama Field'
        hdr[2].text = 'Tipe Data'
        hdr[3].text = 'Deskripsi'
        
        # Make header bold
        for cell in hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        for j, col in enumerate(table['columns'], 1):
            row = col_table.add_row().cells
            row[0].text = str(j)
            row[1].text = col['name']
            row[2].text = col['data_type']
            row[3].text = columns_desc.get(col['name'], f"Field {col['name']}")
            
            # Apply column widths to new rows
            for idx, cell in enumerate(row):
                cell.width = col_widths[idx]
        
        # SQL Definition
        doc.add_heading('SQL Definition', level=2)
        sql_para = doc.add_paragraph()
        sql_run = sql_para.add_run(table['definition'])
        sql_run.font.name = code_font
        sql_run.font.size = Pt(code_size)
        
        doc.add_paragraph("─" * 50)
    
    doc.save(output_path)
    return output_path


# ============================================
# MAIN
# ============================================
def main():
    """Main function"""
    
    print_header()
    check_gpu_status()
    
    # ==========================================
    # Load Knowledge Base dari Database
    # ==========================================
    print()
    print("=" * 50)
    print("📚 KNOWLEDGE BASE")
    print("=" * 50)
    
    if CONFIG.get("use_db_knowledge", False):
        kb_loaded = load_knowledge_base_from_db()
        if not kb_loaded:
            print("⚠️  Lanjut tanpa knowledge base, AI akan generate semua")
    else:
        print("ℹ️  Knowledge base dinonaktifkan (use_db_knowledge=False)")
    
    print()
    
    # Resolve paths
    script_dir = Path(__file__).parent
    sql_file = (script_dir / CONFIG['sql_file']).resolve()
    output_dir = (script_dir / CONFIG['output_dir']).resolve()
    
    # Create output directory
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Check SQL file exists
    if not sql_file.exists():
        print(f"❌ File tidak ditemukan: {sql_file}")
        print(f"   Pastikan file sm.sql ada di: {sql_file.parent}")
        return
    
    print(f"📂 SQL File: {sql_file}")
    print(f"📁 Output Dir: {output_dir}")
    print()
    
    # Extract tables
    print("📖 Extracting tables dari SQL file...")
    tables = extract_tables_from_sql(str(sql_file))
    print(f"   ✅ Ditemukan {len(tables)} tables")
    print()
    
    if not tables:
        print("❌ Tidak ada tables yang ditemukan!")
        return
    
    # Generate documentation for each table
    tables_docs = []
    
    print("=" * 50)
    print("🤖 GENERATING DOKUMENTASI (DB + AI)")
    print("=" * 50)
    print("   Prioritas: DB Knowledge Base → AI → Fallback")
    print()
    
    for i, table in enumerate(tables, 1):
        print(f"   [{i}/{len(tables)}] Processing: {table['name']}")
        
        ai_doc = generate_table_documentation_with_ai(table, CONFIG['model'])
        
        tables_docs.append({
            'name': table['name'],
            'columns': table['columns'],
            'constraints': table['constraints'],
            'definition': table['definition'],
            'ai_doc': ai_doc
        })
    
    print()
    print("💾 Saving documentation...")
    
    # Save outputs
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    if 'md' in CONFIG['output_formats']:
        md_path = output_dir / f"PLSTX_Tables_Documentation_{timestamp}.md"
        save_as_markdown(tables_docs, str(md_path))
        print(f"   ✅ Markdown: {md_path}")
    
    if 'docx' in CONFIG['output_formats']:
        docx_path = output_dir / f"PLSTX_Tables_Documentation_{timestamp}.docx"
        result = save_as_docx(tables_docs, str(docx_path))
        if result:
            print(f"   ✅ Word: {docx_path}")
    
    print()
    print("🎉 Selesai! Dokumentasi berhasil di-generate.")


def print_header():
    """Print header"""
    header = """
╔══════════════════════════════════════════════════════╗
║     SQL Table to Documentation Generator             ║
║     Powered by Ollama AI + GPU 🚀                    ║
╚══════════════════════════════════════════════════════╝
"""
    print(header)


def check_gpu_status():
    """Check dan print GPU status"""
    try:
        import subprocess
        result = subprocess.run(
            ['nvidia-smi', '--query-gpu=name,memory.total,memory.free', '--format=csv,noheader'],
            capture_output=True, text=True, timeout=5
        )
        if result.returncode == 0:
            gpu_info = result.stdout.strip()
            print(f"🎮 GPU Detected: {gpu_info}")
            print(f"   Model: {CONFIG['model']}")
            print()
            return True
    except:
        pass
    
    print("⚠️  GPU tidak terdeteksi, menggunakan CPU")
    print(f"   Model: {CONFIG['model']}")
    print()
    return False


if __name__ == "__main__":
    main()
