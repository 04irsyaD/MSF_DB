"""
================================================================================
SQL to Documentation Generator
================================================================================
Script untuk generate dokumentasi otomatis dari SQL views menggunakan AI lokal (Ollama)

Author  : 04irsyaD
Created : 2026-01-28
Project : ERM Views Documentation

Usage:
    python sql_to_docs.py

Requirements:
    - Ollama installed dan running
    - Model AI sudah di-download (ollama pull llama3.2)
    - pip install -r requirements.txt
================================================================================
"""

import re
import os
from pathlib import Path
from datetime import datetime

try:
    import ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False
    print("⚠️  Package 'ollama' tidak ditemukan. Install dengan: pip install ollama")

try:
    from rich.console import Console
    from rich.progress import Progress, SpinnerColumn, TextColumn
    from rich.panel import Panel
    from rich.table import Table
    RICH_AVAILABLE = True
    console = Console()
except ImportError:
    RICH_AVAILABLE = False
    console = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ============================================
# CONFIGURATION
# ============================================
# Optimized for NVIDIA RTX 3050 (4GB VRAM)
CONFIG = {
    "sql_file": r"..\final_view.sql",  # Relative path ke file SQL
    "output_dir": r".\output",
    
    # Model Ollama - Pakai yang sudah terinstall
    "model": "llama3:latest",         # 🔥 4.7GB - Sudah ada, bagus untuk dokumentasi
    # "model": "deepseek-r1:8b",      # 5.2GB - Reasoning model, lebih lambat tapi detail
    # "model": "qwen3:30b",           # 18GB - Terlalu besar untuk RTX 3050
    
    "output_formats": ["md", "docx"],  # Format output: md, docx, atau keduanya
    
    # GPU Settings untuk RTX 3050
    "gpu_layers": 35,  # Optimal untuk 4GB VRAM (tidak full, sisakan untuk system)
    "num_ctx": 2048,   # Context window (lebih kecil = lebih hemat VRAM)
}


# ============================================
# SQL PARSER
# ============================================
def extract_views_from_sql(sql_file: str) -> list[dict]:
    """
    Extract semua CREATE VIEW dari file SQL
    
    Args:
        sql_file: Path ke file SQL
        
    Returns:
        List of dict dengan keys: name, definition, comment
    """
    with open(sql_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    views = []
    
    # Pattern untuk menangkap comment sebelum view dan definisi view
    # Mendukung: CREATE VIEW, CREATE OR REPLACE VIEW
    pattern = r'(?:--\s*(.+?)\n)?(?:CREATE\s+(?:OR\s+REPLACE\s+)?VIEW\s+(?:public\.)?(\w+)\s+AS\s+)(.*?)(?=(?:CREATE\s+(?:OR\s+REPLACE\s+)?VIEW|ALTER\s+VIEW|;?\s*$))'
    
    # Simplified pattern - lebih reliable
    lines = content.split('\n')
    current_view = None
    current_comment = ""
    current_definition = []
    in_view = False
    
    for i, line in enumerate(lines):
        # Check for comment before view
        if line.strip().startswith('--') and not in_view:
            current_comment = line.strip('- ').strip()
            continue
            
        # Check for CREATE VIEW
        create_match = re.match(r'CREATE\s+(?:OR\s+REPLACE\s+)?VIEW\s+(?:public\.)?(\w+)', line, re.IGNORECASE)
        if create_match:
            # Save previous view if exists
            if current_view and current_definition:
                views.append({
                    'name': current_view,
                    'comment': current_comment if current_comment else "",
                    'definition': '\n'.join(current_definition)
                })
            
            current_view = create_match.group(1)
            current_definition = [line]
            in_view = True
            continue
        
        # If we're inside a view definition
        if in_view:
            # Check if this line starts a new view or ALTER
            if re.match(r'CREATE\s+(?:OR\s+REPLACE\s+)?VIEW', line, re.IGNORECASE):
                # Save current and start new
                if current_view:
                    views.append({
                        'name': current_view,
                        'comment': current_comment,
                        'definition': '\n'.join(current_definition)
                    })
                current_view = None
                current_definition = []
                in_view = False
                current_comment = ""
            elif re.match(r'ALTER\s+VIEW', line, re.IGNORECASE):
                # End current view, don't start new
                if current_view:
                    views.append({
                        'name': current_view,
                        'comment': current_comment,
                        'definition': '\n'.join(current_definition)
                    })
                current_view = None
                current_definition = []
                in_view = False
                current_comment = ""
            else:
                current_definition.append(line)
    
    # Don't forget the last view
    if current_view and current_definition:
        views.append({
            'name': current_view,
            'comment': current_comment,
            'definition': '\n'.join(current_definition)
        })
    
    return views


# ============================================
# AI DOCUMENTATION GENERATOR
# ============================================
def generate_documentation_with_ai(view: dict, model: str = "llama3.2") -> str:
    """
    Generate dokumentasi menggunakan Ollama AI
    
    Args:
        view: Dict dengan keys name, definition, comment
        model: Nama model Ollama
        
    Returns:
        String dokumentasi dalam format Markdown
    """
    if not OLLAMA_AVAILABLE:
        return generate_documentation_basic(view)
    
    # Limit definition length untuk context window
    definition = view['definition'][:3000] if len(view['definition']) > 3000 else view['definition']
    
    prompt = f"""Kamu adalah database documentation expert. Analisis PostgreSQL view berikut dan buat dokumentasi dalam Bahasa Indonesia.

NAMA VIEW: {view['name']}
COMMENT: {view.get('comment', 'Tidak ada')}

SQL DEFINITION:
```sql
{definition}
```

Buatkan dokumentasi dengan format berikut:

### Deskripsi
[Jelaskan tujuan dan fungsi view ini dalam 2-3 kalimat]

### Kolom Output
| Kolom | Deskripsi |
|-------|-----------|
[List semua kolom yang dihasilkan view]

### Tabel Sumber
[List tabel-tabel yang digunakan dengan penjelasan singkat]

### Filter & Kondisi
[Jelaskan WHERE clause dan kondisi yang digunakan]

### Catatan Penting
[Hal-hal penting yang perlu diperhatikan]

Jawab dalam Bahasa Indonesia, format Markdown yang rapi."""

    try:
        response = ollama.chat(
            model=model, 
            messages=[{'role': 'user', 'content': prompt}],
            options={
                'num_gpu': CONFIG.get('gpu_layers', 35),  # RTX 3050 optimal
                'num_ctx': CONFIG.get('num_ctx', 2048),   # Context window
                'temperature': 0.3,  # Lebih fokus/konsisten
            }
        )
        return response['message']['content']
    except Exception as e:
        print(f"⚠️  Error calling Ollama: {e}")
        return generate_documentation_basic(view)


def generate_documentation_basic(view: dict) -> str:
    """
    Generate dokumentasi basic tanpa AI (fallback)
    """
    definition = view['definition']
    
    # Extract tables from JOIN clauses
    tables = re.findall(r'FROM\s+(\w+)|JOIN\s+(\w+)', definition, re.IGNORECASE)
    tables = list(set([t[0] or t[1] for t in tables if t[0] or t[1]]))
    
    # Extract columns (simplified)
    select_match = re.search(r'SELECT\s+(.*?)\s+FROM', definition, re.IGNORECASE | re.DOTALL)
    columns = []
    if select_match:
        cols_str = select_match.group(1)
        # Simple column extraction
        col_patterns = re.findall(r'(?:AS\s+["\']?(\w+)["\']?|(\w+)\s*,)', cols_str, re.IGNORECASE)
        columns = [c[0] or c[1] for c in col_patterns if c[0] or c[1]]
    
    doc = f"""### Deskripsi
{view.get('comment', 'Dokumentasi belum tersedia.')}

### Kolom Output
| Kolom | Deskripsi |
|-------|-----------|
"""
    for col in columns[:10]:  # Limit to 10 columns
        doc += f"| {col} | - |\n"
    
    doc += f"""
### Tabel Sumber
"""
    for table in tables:
        doc += f"- `{table}`\n"
    
    doc += """
### Filter & Kondisi
- Lihat definisi SQL untuk detail kondisi

### Catatan Penting
- Dokumentasi ini di-generate otomatis
"""
    return doc


# ============================================
# OUTPUT GENERATORS
# ============================================
def save_as_markdown(views_docs: list[dict], output_path: str):
    """Save dokumentasi ke file Markdown"""
    
    content = f"""# 📚 Dokumentasi Views - ERM System

> **Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  
> **Source:** final_view.sql  
> **Total Views:** {len(views_docs)}

---

## 📋 Daftar Views

| No | Nama View | Deskripsi Singkat |
|----|-----------|-------------------|
"""
    
    for i, view in enumerate(views_docs, 1):
        # Extract deskripsi dari dokumentasi AI (baris pertama setelah "### Deskripsi")
        desc = extract_short_description(view.get('documentation', ''))
        content += f"| {i} | [`{view['name']}`](#{view['name'].lower()}) | {desc} |\n"
    
    content += "\n---\n\n"
    
    for i, view in enumerate(views_docs, 1):
        content += f"""## {i}. `{view['name']}`

{view['documentation']}

<details>
<summary>📝 <strong>SQL Definition</strong> (klik untuk expand)</summary>

```sql
{view.get('definition', 'SQL tidak tersedia')}
```

</details>

---

"""
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return output_path


def extract_short_description(documentation: str) -> str:
    """Extract deskripsi singkat dari dokumentasi AI"""
    lines = documentation.split('\n')
    capture = False
    desc_lines = []
    
    for line in lines:
        if '### Deskripsi' in line or '**Deskripsi**' in line:
            capture = True
            continue
        if capture:
            # Stop at next header
            if line.startswith('#') or line.startswith('|') or '### ' in line:
                break
            if line.strip():
                desc_lines.append(line.strip())
                # Ambil maksimal 2 kalimat
                if len(desc_lines) >= 2:
                    break
    
    desc = ' '.join(desc_lines)
    
    # Ambil sampai titik pertama atau kedua (1-2 kalimat penuh)
    sentences = desc.split('.')
    if len(sentences) >= 2:
        desc = sentences[0] + '.'
    
    return desc if desc else "Lihat detail di bawah"


def save_as_docx(views_docs: list[dict], output_path: str):
    """Save dokumentasi ke file Word (.docx)"""
    
    if not DOCX_AVAILABLE:
        print("⚠️  python-docx tidak tersedia. Skip export DOCX.")
        return None
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Dokumentasi Views - ERM System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Views: {len(views_docs)}")
    doc.add_paragraph("─" * 50)
    
    # Table of Contents
    doc.add_heading('Daftar Views', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Nama View'
    hdr_cells[2].text = 'Deskripsi'
    
    for i, view in enumerate(views_docs, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = view['name']
        # Extract deskripsi dari dokumentasi AI - tidak dipotong
        desc = extract_short_description(view.get('documentation', ''))
        row_cells[2].text = desc
    
    doc.add_page_break()
    
    # Detail each view
    for i, view in enumerate(views_docs, 1):
        doc.add_heading(f"{i}. {view['name']}", level=1)
        
        # Add documentation content
        for line in view['documentation'].split('\n'):
            if line.startswith('###'):
                doc.add_heading(line.replace('#', '').strip(), level=2)
            elif line.startswith('|'):
                # Skip markdown tables in docx (complex to convert)
                continue
            elif line.strip():
                doc.add_paragraph(line)
        
        # Add SQL Definition section
        doc.add_heading('SQL Definition', level=2)
        sql_para = doc.add_paragraph()
        sql_run = sql_para.add_run(view.get('definition', 'SQL tidak tersedia'))
        sql_run.font.name = 'Consolas'
        sql_run.font.size = Pt(9)
        
        doc.add_paragraph("─" * 50)
    
    doc.save(output_path)
    return output_path


# ============================================
# MAIN
# ============================================
def main():
    """Main function"""
    
    print_header()
    check_gpu_status()
    
    # Resolve paths
    script_dir = Path(__file__).parent
    sql_file = (script_dir / CONFIG['sql_file']).resolve()
    output_dir = (script_dir / CONFIG['output_dir']).resolve()
    
    # Create output directory
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Check SQL file exists
    if not sql_file.exists():
        print(f"❌ File tidak ditemukan: {sql_file}")
        print(f"   Pastikan file final_view.sql ada di: {sql_file.parent}")
        return
    
    print(f"📂 SQL File: {sql_file}")
    print(f"📁 Output Dir: {output_dir}")
    print()
    
    # Extract views
    print("📖 Extracting views dari SQL file...")
    views = extract_views_from_sql(str(sql_file))
    print(f"   ✅ Ditemukan {len(views)} views")
    print()
    
    if not views:
        print("❌ Tidak ada views yang ditemukan!")
        return
    
    # Generate documentation for each view
    views_docs = []
    
    print("🤖 Generating dokumentasi dengan AI...")
    print("   (Proses ini membutuhkan waktu beberapa menit)")
    print()
    
    for i, view in enumerate(views, 1):
        print(f"   [{i}/{len(views)}] Processing: {view['name']}")
        
        doc = generate_documentation_with_ai(view, CONFIG['model'])
        
        views_docs.append({
            'name': view['name'],
            'comment': view.get('comment', ''),
            'definition': view['definition'],
            'documentation': doc
        })
    
    print()
    print("💾 Saving documentation...")
    
    # Save outputs
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    if 'md' in CONFIG['output_formats']:
        md_path = output_dir / f"ERM_Views_Documentation_{timestamp}.md"
        save_as_markdown(views_docs, str(md_path))
        print(f"   ✅ Markdown: {md_path}")
    
    if 'docx' in CONFIG['output_formats']:
        docx_path = output_dir / f"ERM_Views_Documentation_{timestamp}.docx"
        result = save_as_docx(views_docs, str(docx_path))
        if result:
            print(f"   ✅ Word: {docx_path}")
    
    print()
    print("🎉 Selesai! Dokumentasi berhasil di-generate.")


def print_header():
    """Print header"""
    header = """
╔══════════════════════════════════════════════════════╗
║     SQL to Documentation Generator                   ║
║     Powered by Ollama AI + GPU 🚀                    ║
╚══════════════════════════════════════════════════════╝
"""
    print(header)


def check_gpu_status():
    """Check dan print GPU status"""
    try:
        import subprocess
        result = subprocess.run(
            ['nvidia-smi', '--query-gpu=name,memory.total,memory.free', '--format=csv,noheader'],
            capture_output=True, text=True, timeout=5
        )
        if result.returncode == 0:
            gpu_info = result.stdout.strip()
            print(f"🎮 GPU Detected: {gpu_info}")
            print(f"   Model: {CONFIG['model']}")
            print()
            return True
    except:
        pass
    
    print("⚠️  GPU tidak terdeteksi, menggunakan CPU")
    print(f"   Model: {CONFIG['model']}")
    print()
    return False


if __name__ == "__main__":
    main()
