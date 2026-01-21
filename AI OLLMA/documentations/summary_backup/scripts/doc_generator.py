# doc_generator.py - Template preserving document generator
from docxtpl import DocxTemplate
from datetime import datetime
import os
import zipfile
import shutil
from xml.etree import ElementTree as ET

class TemplatePreservingGenerator:
    """Generator yang mempertahankan format visual template asli"""
    
    def __init__(self, template_path):
        self.template_path = template_path
        self.template_backup = None
        
    def analyze_template_structure(self):
        """Analisis struktur template untuk preservasi format"""
        try:
            doc = DocxTemplate(self.template_path)
            
            # Ekstrak informasi template
            template_info = {
                'has_headers': bool(doc.docx.sections[0].header.paragraphs if doc.docx.sections else False),
                'has_footers': bool(doc.docx.sections[0].footer.paragraphs if doc.docx.sections else False),
                'styles_count': len(doc.docx.styles),
                'sections_count': len(doc.docx.sections),
                'original_paragraphs': len(doc.docx.paragraphs),
                'original_tables': len(doc.docx.tables)
            }
            
            return template_info
            
        except Exception as e:
            print(f"⚠️ Template analysis failed: {e}")
            return {}
    
    def backup_template(self):
        """Backup template asli untuk preservasi"""
        backup_path = self.template_path + '.backup'
        shutil.copy2(self.template_path, backup_path)
        self.template_backup = backup_path
        return backup_path
    
    def restore_template_formatting(self, output_path):
        """Restore formatting template asli jika diperlukan"""
        if not self.template_backup:
            return False
            
        try:
            # Implementasi khusus untuk mempertahankan format
            # Untuk saat ini, kita gunakan pendekatan sederhana
            return True
        except Exception as e:
            print(f"⚠️ Format restoration failed: {e}")
            return False
    
    def generate_with_format_preservation(self, output_path, db_name, tables_data, template_context=None):
        """Generate dokumen dengan preservasi format maksimal"""
        print(f"📄 Generating document with format preservation...")
        
        try:
            # Backup template
            self.backup_template()
            
            # Analisis struktur
            template_info = self.analyze_template_structure()
            print(f"   Template info: {template_info.get('styles_count', 0)} styles, {template_info.get('original_paragraphs', 0)} paragraphs")
            
            # Siapkan konteks data yang kaya
            enhanced_context = self._prepare_enhanced_context(db_name, tables_data, template_context)
            
            # Generate menggunakan DocxTemplate
            doc = DocxTemplate(self.template_path)
            doc.render(enhanced_context)
            
            # Pastikan folder output ada
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Simpan dokumen
            doc.save(output_path)
            
            # Verifikasi hasil
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                print(f"   ✅ Document generated: {file_size} bytes")
                
                # Validasi content
                validation_result = self._validate_generated_content(output_path, tables_data)
                if validation_result['success']:
                    print(f"   ✅ Content validation passed")
                else:
                    print(f"   ⚠️ Content validation issues: {validation_result['issues']}")
                
                return output_path
            else:
                raise Exception("Output file was not created")
                
        except Exception as e:
            print(f"❌ Generation failed: {e}")
            return None
        finally:
            # Cleanup backup if needed
            if self.template_backup and os.path.exists(self.template_backup):
                os.remove(self.template_backup)
    
    def _prepare_enhanced_context(self, db_name, tables_data, template_context):
        """Siapkan konteks yang diperkaya untuk template"""
        # Base context
        context = {
            'db_name': db_name,
            'generated_date': datetime.now().strftime("%d %B %Y"),
            'generated_time': datetime.now().strftime("%H:%M:%S"),
            'tables': []
        }
        
        # Enhanced table data
        total_columns = 0
        categories = {}
        
        for table_info in tables_data:
            # Hitung statistik
            columns = table_info.get('columns', [])
            if isinstance(columns, list):
                total_columns += len(columns)
            
            # Kategorisasi
            category = table_info.get('category', 'general')
            categories[category] = categories.get(category, 0) + 1
            
            # Enhanced table info
            enhanced_table = {
                'table_name': table_info.get('table_name', ''),
                'description': table_info.get('description', ''),
                'columns': columns,
                'column_count': len(columns) if isinstance(columns, list) else 0,
                'category': category,
                'has_primary_key': any(col.get('is_primary_key', False) for col in columns) if isinstance(columns, list) else False,
                'has_foreign_keys': any(col.get('is_foreign_key', False) for col in columns) if isinstance(columns, list) else False,
                'primary_keys': [col['name'] for col in columns if col.get('is_primary_key')] if isinstance(columns, list) else [],
                'foreign_keys': [col['name'] for col in columns if col.get('is_foreign_key')] if isinstance(columns, list) else []
            }
            
            context['tables'].append(enhanced_table)
        
        # Summary statistics
        context['summary'] = {
            'total_tables': len(tables_data),
            'total_columns': total_columns,
            'categories': categories,
            'avg_columns_per_table': round(total_columns / len(tables_data), 1) if tables_data else 0
        }
        
        # Template context integration
        if template_context:
            context['template_style'] = template_context.get('style', 'standard')
            context['template_format'] = template_context.get('format', 'professional')
        
        return context
    
    def _validate_generated_content(self, output_path, expected_tables):
        """Validasi konten yang dihasilkan"""
        try:
            # Basic file validation
            if not os.path.exists(output_path):
                return {'success': False, 'issues': ['File not found']}
            
            file_size = os.path.getsize(output_path)
            if file_size < 1000:  # Too small, likely empty or error
                return {'success': False, 'issues': ['File too small']}
            
            # Try to open and check basic structure
            try:
                from docx import Document
                doc = Document(output_path)
                
                # Count paragraphs and tables
                paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
                table_count = len(doc.tables)
                
                if paragraph_count < 5:  # Too few paragraphs
                    return {'success': False, 'issues': ['Too few paragraphs']}
                
                return {
                    'success': True, 
                    'stats': {
                        'paragraphs': paragraph_count,
                        'tables': table_count,
                        'file_size': file_size
                    }
                }
                
            except ImportError:
                # If python-docx not available, do basic validation
                return {'success': True, 'stats': {'file_size': file_size}}
            
        except Exception as e:
            return {'success': False, 'issues': [f'Validation error: {e}']}

def generate_document(template_path, output_path, db_name, tables_data, template_context=None):
    """Main function with backward compatibility"""
    generator = TemplatePreservingGenerator(template_path)
    return generator.generate_with_format_preservation(output_path, db_name, tables_data, template_context)

# Legacy function for compatibility
def generate_document_simple(template_path, output_path, db_name, tables_data):
    """Simple generation for backward compatibility"""
    try:
        doc = DocxTemplate(template_path)
        context = {
            'db_name': db_name,
            'generated_date': datetime.now().strftime("%d %B %Y"),
            'tables': tables_data
        }
        doc.render(context)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"Simple generation failed: {e}")
        return None

if __name__ == '__main__':
    tpl = os.path.join(os.path.dirname(__file__), '..', 'template', 'template_dokumentasi.docx')
    out = os.path.join(os.path.dirname(__file__), '..', 'output', 'Dokumentasi_Auto.docx')
    sample_tables = [{
        'table_name':'customer',
        'description':'Tabel menyimpan data pelanggan.',
        'columns':[{'name':'id','type':'integer','nullable':'NO'},{'name':'name','type':'varchar','nullable':'NO'}]
    }]
    print(generate_document(tpl, out, 'sample_db', sample_tables))
