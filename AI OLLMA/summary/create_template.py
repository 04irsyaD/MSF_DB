# create_proper_template.py - Buat template DOCX yang benar
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_template():
    doc = Document()
    
    # Title
    title = doc.add_heading('DOKUMENTASI DATABASE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Database info
    doc.add_paragraph()
    p1 = doc.add_paragraph()
    p1.add_run('Database: ').bold = True
    p1.add_run('{{db_name}}')
    
    p2 = doc.add_paragraph()
    p2.add_run('Tanggal Generate: ').bold = True  
    p2.add_run('{{generated_date}}')
    
    doc.add_paragraph()
    
    # Tables section
    doc.add_heading('DAFTAR TABEL', level=1)
    
    # Loop untuk setiap tabel - ini yang penting untuk Jinja2
    doc.add_paragraph('{% for table in tables %}')
    
    # Table name sebagai heading
    doc.add_heading('{{table.table_name}}', level=2)
    
    # Description
    p_desc = doc.add_paragraph()
    p_desc.add_run('Deskripsi: ').bold = True
    doc.add_paragraph('{{table.description}}')
    
    # Columns table
    doc.add_paragraph()
    p_cols = doc.add_paragraph()
    p_cols.add_run('Kolom-kolom:').bold = True
    
    # Create table for columns
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nama Kolom'
    hdr_cells[1].text = 'Tipe Data'  
    hdr_cells[2].text = 'Nullable'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Jinja2 loop untuk kolom
    doc.add_paragraph('{% for col in table.columns %}')
    
    # Add row (ini akan dirender oleh Jinja2)
    doc.add_paragraph('{{col.name}} | {{col.type}} | {{col.nullable}}')
    
    doc.add_paragraph('{% endfor %}')
    
    doc.add_paragraph('---')
    doc.add_paragraph('{% endfor %}')
    
    # Save template
    template_path = os.path.join('template', 'template_dokumentasi_new.docx')
    doc.save(template_path)
    print(f'✅ Template baru dibuat: {template_path}')

if __name__ == '__main__':
    create_template()