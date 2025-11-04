# template_filler_docx.py - Sistem yang menggunakan template asli dan isi data ke placeholder
import os
import shutil
from datetime import datetime

class TemplateFillDocxSystem:
    """Sistem yang menggunakan template DOCX asli sebagai base dan mengisi data"""
    
    def __init__(self):
        self.db_config = {
            'host': 'localhost',
            'port': '5414', 
            'dbname': 'deverm',
            'user': 'postgres',
            'password': '1234'
        }
        
        self.template_path = os.path.join('..', 'template', 'template_dokumentasi.docx')
        self.output_dir = os.path.join('..', 'output')
        
    def check_and_copy_template(self):
        """Check template dan copy sebagai base untuk editing"""
        
        if not os.path.exists(self.template_path):
            print(f"   ❌ Template tidak ditemukan: {self.template_path}")
            return None
        
        # Copy template ke output sebagai working file
        output_path = os.path.join(self.output_dir, "Template_Filled_Documentation.docx")
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        try:
            shutil.copy2(self.template_path, output_path)
            print(f"   ✅ Template copied to: {os.path.basename(output_path)}")
            return output_path
        except Exception as e:
            print(f"   ❌ Error copying template: {e}")
            return None
    
    def get_database_summary(self):
        """Get ringkasan database untuk mengisi template"""
        print("🗄️ Getting database summary for template...")
        
        try:
            import psycopg2
            
            conn = psycopg2.connect(
                host=self.db_config['host'],
                port=self.db_config['port'],
                dbname=self.db_config['dbname'],
                user=self.db_config['user'],
                password=self.db_config['password']
            )
            
            cur = conn.cursor()
            
            # Get comprehensive summary
            cur.execute("""
                SELECT 
                    COUNT(DISTINCT t.table_name) as total_tables,
                    COUNT(c.column_name) as total_columns,
                    COUNT(CASE WHEN pk.column_name IS NOT NULL THEN 1 END) as total_pks,
                    COUNT(CASE WHEN fk.column_name IS NOT NULL THEN 1 END) as total_fks
                FROM information_schema.tables t
                LEFT JOIN information_schema.columns c ON t.table_name = c.table_name AND c.table_schema = 'public'
                LEFT JOIN (
                    SELECT kcu.column_name, kcu.table_name
                    FROM information_schema.table_constraints tc
                    JOIN information_schema.key_column_usage kcu ON tc.constraint_name = kcu.constraint_name
                    WHERE tc.constraint_type = 'PRIMARY KEY'
                ) pk ON c.column_name = pk.column_name AND c.table_name = pk.table_name
                LEFT JOIN (
                    SELECT kcu.column_name, kcu.table_name
                    FROM information_schema.table_constraints AS tc
                    JOIN information_schema.key_column_usage AS kcu ON tc.constraint_name = kcu.constraint_name
                    WHERE tc.constraint_type = 'FOREIGN KEY'
                ) fk ON c.column_name = fk.column_name AND c.table_name = fk.table_name
                WHERE t.table_schema = 'public';
            """)
            
            summary = cur.fetchone()
            
            # Get top tables untuk contoh
            cur.execute("""
                SELECT 
                    t.table_name,
                    COUNT(c.column_name) as column_count
                FROM information_schema.tables t
                LEFT JOIN information_schema.columns c ON t.table_name = c.table_name AND c.table_schema = 'public'
                WHERE t.table_schema = 'public'
                GROUP BY t.table_name
                ORDER BY column_count DESC
                LIMIT 10;
            """)
            
            top_tables = cur.fetchall()
            
            cur.close()
            conn.close()
            
            db_summary = {
                'total_tables': summary[0] if summary[0] else 0,
                'total_columns': summary[1] if summary[1] else 0,
                'total_pks': summary[2] if summary[2] else 0,
                'total_fks': summary[3] if summary[3] else 0,
                'top_tables': top_tables,
                'database_name': self.db_config['dbname'],
                'generated_date': datetime.now().strftime("%d %B %Y"),
                'generated_time': datetime.now().strftime("%H:%M WIB")
            }
            
            print(f"   ✅ Database summary: {db_summary['total_tables']} tables, {db_summary['total_columns']} columns")
            return db_summary
            
        except Exception as e:
            print(f"   ❌ Database error: {e}")
            return None
    
    def fill_template_with_data(self, template_file, db_summary):
        """Isi template dengan data database"""
        
        try:
            from docx import Document
            
            print("📝 Filling template with database data...")
            
            # Buka template yang sudah di-copy
            doc = Document(template_file)
            
            # Dictionary untuk replacement
            replacements = {
                '{DB_NAME}': db_summary['database_name'].upper(),
                '{TOTAL_TABLES}': str(db_summary['total_tables']),
                '{TOTAL_COLUMNS}': str(db_summary['total_columns']),
                '{TOTAL_PKS}': str(db_summary['total_pks']),
                '{TOTAL_FKS}': str(db_summary['total_fks']),
                '{GENERATED_DATE}': db_summary['generated_date'],
                '{GENERATED_TIME}': db_summary['generated_time'],
                '{AVG_COLUMNS}': str(round(db_summary['total_columns'] / db_summary['total_tables'], 1) if db_summary['total_tables'] > 0 else 0)
            }
            
            # Replace dalam paragraphs
            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        # Replace text sambil mempertahankan formatting
                        for run in paragraph.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, value)
            
            # Replace dalam tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in replacements.items():
                            if key in cell.text:
                                cell.text = cell.text.replace(key, value)
            
            # Tambah data tabel jika ada placeholder khusus
            self.add_tables_to_template(doc, db_summary['top_tables'])
            
            # Save hasil
            doc.save(template_file)
            
            print(f"   ✅ Template filled with data")
            return True
            
        except Exception as e:
            print(f"   ❌ Error filling template: {e}")
            return False
    
    def add_tables_to_template(self, doc, top_tables):
        """Tambah tabel database ke template jika perlu"""
        
        try:
            # Cari apakah ada placeholder untuk tabel
            table_placeholder_found = False
            
            for paragraph in doc.paragraphs:
                if '{TABLE_LIST}' in paragraph.text or 'DAFTAR_TABEL' in paragraph.text.upper():
                    table_placeholder_found = True
                    
                    # Clear existing text
                    paragraph.clear()
                    
                    # Add table title
                    run = paragraph.add_run("Daftar 10 Tabel Terbesar:")
                    run.bold = True
                    
                    break
            
            if table_placeholder_found or len(doc.tables) > 0:
                # Tambah tabel baru setelah paragraf terakhir
                if top_tables:
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    
                    # Header
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'No.'
                    hdr_cells[1].text = 'Nama Tabel'
                    hdr_cells[2].text = 'Jumlah Kolom'
                    
                    # Data
                    for i, (table_name, column_count) in enumerate(top_tables, 1):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(i)
                        row_cells[1].text = str(table_name)
                        row_cells[2].text = str(column_count)
            
        except Exception as e:
            print(f"   ⚠️ Warning adding tables: {e}")
    
    def create_template_filled_documentation(self):
        """Buat dokumentasi dengan mengisi template asli"""
        
        print("📋 TEMPLATE FILLER DOCX SYSTEM")
        print("=" * 60)
        print("🎯 Uses original template and fills data into placeholders")
        print()
        
        start_time = datetime.now()
        
        try:
            # Copy template sebagai base
            working_file = self.check_and_copy_template()
            if not working_file:
                print("❌ Cannot proceed without template file")
                return False
            
            # Get database summary
            db_summary = self.get_database_summary()
            if not db_summary:
                print("❌ Cannot get database information")
                return False
            
            # Fill template with data
            success = self.fill_template_with_data(working_file, db_summary)
            
            if success:
                duration = datetime.now() - start_time
                file_size = os.path.getsize(working_file)
                
                print(f"\n🎉 TEMPLATE FILLED SUCCESSFULLY!")
                print("=" * 60)
                print(f"📁 File: {os.path.basename(working_file)}")
                print(f"📊 Size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
                print(f"📋 Data filled: {db_summary['total_tables']} tables info")
                print(f"⏱️ Duration: {duration}")
                print(f"🎨 Format: Original template preserved")
                print(f"✅ Method: Template + Data replacement")
                
                print(f"\n🎯 TEMPLATE FEATURES PRESERVED:")
                print(f"   ✅ Original headers dan styling")
                print(f"   ✅ Template fonts dan colors")
                print(f"   ✅ Layout dan spacing asli")
                print(f"   ✅ Company branding (jika ada)")
                print(f"   ✅ Custom formatting dari template")
                
                print(f"\n📊 DATA YANG DIISI:")
                print(f"   • Database: {db_summary['database_name']}")
                print(f"   • Total Tables: {db_summary['total_tables']}")
                print(f"   • Total Columns: {db_summary['total_columns']}")
                print(f"   • Generated: {db_summary['generated_date']} {db_summary['generated_time']}")
                
                return working_file
            else:
                print("\n❌ Failed to fill template with data!")
                return False
                
        except Exception as e:
            print(f"\n💥 ERROR: {e}")
            return False

def main():
    """Main function"""
    system = TemplateFillDocxSystem()
    
    print("🚀 Template Filler DOCX System")
    print("📋 Preserves original template format and fills with database data")
    print("🎯 This maintains all your template's headers, fonts, and styling!")
    print()
    
    result = system.create_template_filled_documentation()
    
    if result:
        print(f"\n✨ SUCCESS! Template filled with database data:")
        print(f"📁 {result}")
        print(f"\n🎨 This should look EXACTLY like your template!")
        print(f"💡 All headers, fonts, colors preserved from original!")
        print(f"\n📋 Template placeholders yang bisa digunakan:")
        print(f"   • {{DB_NAME}} - Nama database")
        print(f"   • {{TOTAL_TABLES}} - Jumlah tabel") 
        print(f"   • {{TOTAL_COLUMNS}} - Jumlah kolom")
        print(f"   • {{GENERATED_DATE}} - Tanggal generate")
        print(f"   • {{GENERATED_TIME}} - Waktu generate")
        print(f"   • {{TABLE_LIST}} - Placeholder untuk daftar tabel")
    else:
        print(f"\n💥 FAILED! Check template file and try again")

if __name__ == '__main__':
    main()