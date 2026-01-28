# template_analyzer.py - Analisis template dengan Ollama
import subprocess
import os
from docx import Document
# template_analyzer.py - Analisis template dengan Ollama
import subprocess
import os
from docx import Document
import zipfile
import xml.etree.ElementTree as ET

def extract_template_structure(template_path):
    """Extract text dan structure dari template DOCX"""
    try:
        # Baca dokumen
        doc = Document(template_path)
        
        # Extract semua teks dari template
        template_text = []
        
        # Ambil paragraf
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                template_text.append(paragraph.text.strip())
        
        # Ambil teks dari tabel jika ada
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        template_text.append(cell.text.strip())
        
        return "\n".join(template_text)
    
    except Exception as e:
        print(f"⚠️ Error membaca template: {e}")
        return ""

def analyze_template_with_ollama(template_path, model='llama3'):
    """Analisis template dengan Ollama untuk memahami struktur dan placeholder"""
    
    print("🔍 Menganalisis struktur template dengan AI...")
    
    # Extract konten template
    template_content = extract_template_structure(template_path)
    
    if not template_content:
        print("⚠️ Template kosong atau tidak bisa dibaca")
        return None
    
    # Prompt untuk Ollama
    # NOTE: build the prompt without f-string to avoid formatting errors due to JSON-like braces
    prompt = (
        "Analisis template dokumentasi berikut dan berikan panduan untuk mengisinya:\n\n"
        "KONTEN TEMPLATE:\n"
        + template_content
        + "\n\n"
        "Tolong identifikasi:\n"
        "1. Apakah ada placeholder seperti {{variabel}} atau merge fields?\n"
        "2. Bagaimana struktur template (heading, paragraf, tabel)?\n"
        "3. Di mana data tabel database harus diisi?\n"
        "4. Format apa yang diharapkan untuk deskripsi tabel dan kolom?\n\n"
        "Berikan jawaban dalam format JSON seperti ini:\n"
        "{\n"
        "    \"has_placeholders\": true/false,\n"
        "    \"placeholder_format\": \"jinja2/mailmerge/manual\",\n"
        "    \"sections\": [\"section1\", \"section2\"],\n"
        "    \"table_format\": \"deskripsi format tabel\",\n"
        "    \"recommendations\": \"saran untuk mengisi template\"\n"
        "}\n"
    )
    
    try:
        # Jalankan Ollama
        # Pass the prompt via stdin to avoid Windows command-line length limits
        result = subprocess.run(
            ["ollama", "run", model],
            input=prompt,
            capture_output=True,
            text=True,
            check=True,
            timeout=60,
            encoding='utf-8',
            errors='replace'
        )

        response = result.stdout.strip()
        print("✅ Analisis template selesai")
        return response
        
    except subprocess.TimeoutExpired:
        print("⚠️ Timeout saat menganalisis template")
        return None
    except Exception as e:
        print(f"⚠️ Error analisis template: {e}")
        return None

def get_filling_strategy_from_ollama(template_analysis, sample_table_data, model='llama3'):
    """Dapatkan strategi pengisian template berdasarkan analisis"""
    
    print("🤖 Menentukan strategi pengisian template...")
    
    prompt = f"""
Berdasarkan analisis template ini:
{template_analysis}

Dan contoh data tabel:
{sample_table_data}

Berikan strategi terbaik untuk mengisi template:
1. Apakah perlu memodifikasi template untuk menambah placeholder?
2. Bagaimana format data yang harus disiapkan?
3. Langkah-langkah pengisian yang tepat?

Jawab dalam format yang jelas dan actionable.
"""
    
    try:
        result = subprocess.run(
            ["ollama", "run", model, prompt],
            capture_output=True,
            text=True,
            check=True,
            timeout=60,
            encoding='utf-8',
            errors='replace'
        )
        
        return result.stdout.strip()
        
    except Exception as e:
        print(f"⚠️ Error mendapatkan strategi: {e}")
        return None

def analyze_and_prepare_template(template_path):
    """Fungsi utama untuk menganalisis template dan menyiapkan strategi pengisian"""
    
    if not os.path.exists(template_path):
        print(f"❌ Template tidak ditemukan: {template_path}")
        return None
    
    print(f"📋 Menganalisis template: {os.path.basename(template_path)}")
    
    # Langkah 1: Analisis struktur template
    analysis = analyze_template_with_ollama(template_path)
    
    if not analysis:
        print("❌ Gagal menganalisis template")
        return None
    
    print(f"📄 Hasil analisis template:")
    print(analysis[:500] + "..." if len(analysis) > 500 else analysis)
    
    # Langkah 2: Contoh data untuk strategi
    sample_data = {
        "table_name": "users",
        "description": "Tabel untuk menyimpan data pengguna",
        "columns": [
            {"name": "id", "type": "integer", "nullable": "NO"},
            {"name": "username", "type": "varchar(50)", "nullable": "NO"}
        ]
    }
    
    # Langkah 3: Dapatkan strategi pengisian
    strategy = get_filling_strategy_from_ollama(analysis, str(sample_data))
    
    if strategy:
        print(f"\n💡 Strategi pengisian:")
        print(strategy[:500] + "..." if len(strategy) > 500 else strategy)
    
    return {
        "analysis": analysis,
        "strategy": strategy,
        "template_path": template_path
    }

if __name__ == '__main__':
    # Test analisis template
    template_path = os.path.join('..', 'template', 'template_dokumentasi.docx')
    result = analyze_and_prepare_template(template_path)
    
    if result:
        print("\n✅ Analisis template berhasil!")
        print("   Gunakan hasil ini untuk menyesuaikan pengisian data.")
    else:
        print("\n❌ Analisis template gagal.")