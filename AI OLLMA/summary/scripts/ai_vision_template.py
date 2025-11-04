# ai_vision_template.py - AI Vision untuk analisis template Word
import os
import subprocess
import json
from docx import Document
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
import zipfile
import xml.etree.ElementTree as ET

class TemplateVisualAnalyzer:
    """Analyzer untuk mengekstrak struktur visual template Word"""
    
    def __init__(self, template_path):
        self.template_path = template_path
        self.visual_data = {}
        
    def extract_visual_elements(self):
        """Ekstrak elemen visual dari template DOCX"""
        print("🔍 Mengekstrak elemen visual dari template...")
        
        try:
            # Buka dokumen
            doc = Document(self.template_path)
            
            # Ekstrak informasi visual
            self.visual_data = {
                'fonts': self._extract_fonts(doc),
                'colors': self._extract_colors(doc),
                'styles': self._extract_styles(doc),
                'structure': self._extract_structure(doc),
                'headers_footers': self._extract_headers_footers(doc),
                'tables': self._extract_table_formats(doc)
            }
            
            return True
            
        except Exception as e:
            print(f"❌ Error mengekstrak visual: {e}")
            return False
    
    def _extract_fonts(self, doc):
        """Ekstrak informasi font"""
        fonts = set()
        
        # Font dari paragraf
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name:
                    fonts.add(run.font.name)
        
        # Font dari styles
        for style in doc.styles:
            if hasattr(style, 'font') and style.font.name:
                fonts.add(style.font.name)
                
        return list(fonts)
    
    def _extract_colors(self, doc):
        """Ekstrak informasi warna"""
        colors = set()
        
        # Warna dari paragraf
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.color.rgb:
                    colors.add(str(run.font.color.rgb))
        
        return list(colors)
    
    def _extract_styles(self, doc):
        """Ekstrak informasi styles"""
        styles_info = {}
        
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                styles_info[style.name] = {
                    'font_name': style.font.name if style.font.name else 'Default',
                    'font_size': str(style.font.size) if style.font.size else 'Default',
                    'bold': style.font.bold if style.font.bold is not None else False,
                    'italic': style.font.italic if style.font.italic is not None else False
                }
                
        return styles_info
    
    def _extract_structure(self, doc):
        """Ekstrak struktur dokumen"""
        structure = {
            'headings': [],
            'paragraphs': [],
            'tables': len(doc.tables),
            'sections': len(doc.sections)
        }
        
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                structure['headings'].append({
                    'level': paragraph.style.name,
                    'text': paragraph.text[:50] + '...' if len(paragraph.text) > 50 else paragraph.text
                })
            elif paragraph.text.strip():
                structure['paragraphs'].append({
                    'style': paragraph.style.name,
                    'text': paragraph.text[:30] + '...' if len(paragraph.text) > 30 else paragraph.text
                })
                
        return structure
    
    def _extract_headers_footers(self, doc):
        """Ekstrak informasi header dan footer"""
        headers_footers = {
            'has_header': False,
            'has_footer': False,
            'header_content': [],
            'footer_content': []
        }
        
        for section in doc.sections:
            if section.header.paragraphs:
                headers_footers['has_header'] = True
                for p in section.header.paragraphs:
                    if p.text.strip():
                        headers_footers['header_content'].append(p.text)
                        
            if section.footer.paragraphs:
                headers_footers['has_footer'] = True
                for p in section.footer.paragraphs:
                    if p.text.strip():
                        headers_footers['footer_content'].append(p.text)
                        
        return headers_footers
    
    def _extract_table_formats(self, doc):
        """Ekstrak format tabel"""
        table_formats = []
        
        for i, table in enumerate(doc.tables):
            table_info = {
                'index': i,
                'rows': len(table.rows),
                'cols': len(table.columns) if table.rows else 0,
                'style': table.style.name if table.style else 'Default'
            }
            table_formats.append(table_info)
            
        return table_formats
    
    def analyze_with_ai(self):
        """Gunakan AI untuk menganalisis struktur visual"""
        print("🤖 Menganalisis struktur visual dengan AI...")
        
        # Buat deskripsi visual untuk AI
        visual_description = self._create_visual_description()
        
        prompt = f"""
Analisis template dokumentasi Word berikut berdasarkan struktur visualnya:

INFORMASI VISUAL:
{visual_description}

Berikan analisis dalam format JSON:
{{
    "document_type": "jenis dokumen (laporan/dokumentasi/manual)",
    "layout_style": "gaya layout (formal/modern/simple)",
    "content_areas": ["area1", "area2"],
    "placeholder_locations": "di mana data harus diisi",
    "preservation_priority": "elemen yang harus dipertahankan",
    "filling_strategy": "strategi mengisi data",
    "expected_content": "jenis konten yang diharapkan"
}}

Fokus pada bagaimana mempertahankan format asli sambil mengisi data database.
"""

        try:
            result = subprocess.run(
                ["ollama", "run", "llama3", prompt],
                capture_output=True, text=True, timeout=60,
                encoding='utf-8', errors='replace'
            )
            
            if result.stdout.strip():
                self.ai_analysis = result.stdout.strip()
                print("✅ Analisis AI selesai")
                return True
            else:
                print("⚠️ AI tidak memberikan response")
                return False
                
        except Exception as e:
            print(f"⚠️ Error AI analysis: {e}")
            return False
    
    def _create_visual_description(self):
        """Buat deskripsi visual untuk AI"""
        desc = f"""
FONTS: {', '.join(self.visual_data.get('fonts', []))}
COLORS: {len(self.visual_data.get('colors', []))} warna berbeda
STRUKTUR:
- Headings: {len(self.visual_data.get('structure', {}).get('headings', []))} 
- Paragraphs: {len(self.visual_data.get('structure', {}).get('paragraphs', []))}
- Tables: {self.visual_data.get('structure', {}).get('tables', 0)}
- Sections: {self.visual_data.get('structure', {}).get('sections', 1)}

HEADERS/FOOTERS: {'Ada header' if self.visual_data.get('headers_footers', {}).get('has_header') else 'Tidak ada header'}, {'Ada footer' if self.visual_data.get('headers_footers', {}).get('has_footer') else 'Tidak ada footer'}

STYLES TERDETEKSI:
{json.dumps(self.visual_data.get('styles', {}), indent=2)}

STRUKTUR HEADINGS:
{json.dumps(self.visual_data.get('structure', {}).get('headings', []), indent=2)}
"""
        return desc
    
    def get_filling_instructions(self):
        """Dapatkan instruksi untuk mengisi template"""
        if not hasattr(self, 'ai_analysis'):
            return None
            
        return {
            'visual_data': self.visual_data,
            'ai_analysis': self.ai_analysis,
            'instructions': self._generate_filling_instructions()
        }
    
    def _generate_filling_instructions(self):
        """Generate instruksi spesifik untuk pengisian"""
        instructions = {
            'preserve_fonts': self.visual_data.get('fonts', []),
            'preserve_colors': self.visual_data.get('colors', []),
            'preserve_styles': list(self.visual_data.get('styles', {}).keys()),
            'preserve_structure': True,
            'content_areas': self._identify_content_areas(),
            'data_insertion_points': self._identify_insertion_points()
        }
        
        return instructions
    
    def _identify_content_areas(self):
        """Identifikasi area untuk konten"""
        areas = []
        
        # Dari struktur paragraf
        for p in self.visual_data.get('structure', {}).get('paragraphs', []):
            if any(keyword in p['text'].lower() for keyword in ['tabel', 'database', 'kolom', 'data']):
                areas.append({
                    'type': 'database_content',
                    'style': p['style'],
                    'sample_text': p['text']
                })
                
        return areas
    
    def _identify_insertion_points(self):
        """Identifikasi titik insersi data"""
        points = []
        
        # Cari placeholder potensial
        for p in self.visual_data.get('structure', {}).get('paragraphs', []):
            if any(marker in p['text'] for marker in ['{{', '[]', '<>', 'XXXXX']):
                points.append({
                    'location': 'paragraph',
                    'style': p['style'],
                    'placeholder': p['text']
                })
                
        return points

def analyze_template_visual(template_path):
    """Fungsi utama untuk analisis visual template"""
    print(f"📋 Memulai analisis visual template: {os.path.basename(template_path)}")
    
    analyzer = TemplateVisualAnalyzer(template_path)
    
    # Ekstrak elemen visual
    if not analyzer.extract_visual_elements():
        return None
    
    # Analisis dengan AI
    if not analyzer.analyze_with_ai():
        print("⚠️ Melanjutkan tanpa analisis AI")
    
    # Dapatkan instruksi pengisian
    instructions = analyzer.get_filling_instructions()
    
    print("✅ Analisis visual selesai")
    return instructions

if __name__ == '__main__':
    template_path = os.path.join('..', 'template', 'template_dokumentasi.docx')
    
    if os.path.exists(template_path):
        result = analyze_template_visual(template_path)
        
        if result:
            print("\n📊 HASIL ANALISIS VISUAL:")
            print("="*50)
            
            # Tampilkan ringkasan
            visual = result['visual_data']
            print(f"Fonts: {', '.join(visual.get('fonts', []))}")
            print(f"Styles: {len(visual.get('styles', {}))}")
            print(f"Structure: {visual.get('structure', {}).get('tables', 0)} tabel, {len(visual.get('structure', {}).get('headings', []))} headings")
            
            if 'ai_analysis' in result:
                print(f"\nAI Analysis Preview:")
                print(result['ai_analysis'][:200] + "...")
                
        else:
            print("❌ Gagal menganalisis template")
    else:
        print(f"❌ Template tidak ditemukan: {template_path}")