# template_inspector.py - Untuk melihat isi template secara detail
import os
import zipfile
import xml.etree.ElementTree as ET

def inspect_template_content():
    """Inspect template content untuk melihat placeholder dan struktur"""
    
    template_path = os.path.join('..', 'template', 'template_dokumentasi.docx')
    
    if not os.path.exists(template_path):
        print(f"❌ Template tidak ditemukan: {template_path}")
        return
    
    print("🔍 TEMPLATE CONTENT INSPECTOR")
    print("=" * 60)
    
    try:
        from docx import Document
        
        # Baca template dengan python-docx
        doc = Document(template_path)
        
        print("📄 DOCUMENT PARAGRAPHS:")
        print("-" * 40)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"{i+1:2d}. '{para.text}'")
                print(f"     Style: {para.style.name}")
                
                # Cek runs dalam paragraf
                for j, run in enumerate(para.runs):
                    if run.text.strip():
                        print(f"     Run {j+1}: '{run.text}' (Font: {run.font.name}, Bold: {run.font.bold})")
                print()
        
        print("🎨 HEADERS:")
        print("-" * 40)
        for section_idx, section in enumerate(doc.sections):
            print(f"Section {section_idx + 1}:")
            if section.header:
                for i, para in enumerate(section.header.paragraphs):
                    if para.text.strip():
                        print(f"  Header {i+1}: '{para.text}'")
                        print(f"    Style: {para.style.name}")
                        
                        for j, run in enumerate(para.runs):
                            if run.text.strip():
                                print(f"    Run {j+1}: '{run.text}' (Font: {run.font.name})")
            else:
                print("  No header found")
            print()
        
        print("📋 FOOTERS:")
        print("-" * 40)
        for section_idx, section in enumerate(doc.sections):
            print(f"Section {section_idx + 1}:")
            if section.footer:
                for i, para in enumerate(section.footer.paragraphs):
                    if para.text.strip():
                        print(f"  Footer {i+1}: '{para.text}'")
                        
                        for j, run in enumerate(para.runs):
                            if run.text.strip():
                                print(f"    Run {j+1}: '{run.text}'")
            else:
                print("  No footer found")
            print()
        
    except Exception as e:
        print(f"❌ Error with python-docx: {e}")
    
    # Inspect XML level
    print("🔍 XML LEVEL INSPECTION:")
    print("-" * 40)
    
    try:
        with zipfile.ZipFile(template_path, 'r') as zip_file:
            
            # Check document.xml
            if 'word/document.xml' in zip_file.namelist():
                doc_xml = zip_file.read('word/document.xml').decode('utf-8')
                
                print("📄 Document.xml content preview:")
                # Cari text content
                import re
                text_matches = re.findall(r'<w:t[^>]*>([^<]*)</w:t>', doc_xml)
                for i, text in enumerate(text_matches[:15], 1):  # First 15 text elements
                    if text.strip():
                        print(f"  {i:2d}. '{text}'")
                
                # Cari placeholder patterns
                placeholders = re.findall(r'[\{\[\(]([^}\]\)]+)[\}\]\)]', doc_xml)
                if placeholders:
                    print(f"\n📝 Potential placeholders found:")
                    for placeholder in set(placeholders):
                        print(f"  • {{{placeholder}}}")
                else:
                    print(f"\n📝 No clear placeholders found in document")
            
            # Check headers
            header_files = [f for f in zip_file.namelist() if 'header' in f and f.endswith('.xml')]
            for header_file in header_files:
                print(f"\n🎨 {header_file}:")
                try:
                    header_xml = zip_file.read(header_file).decode('utf-8')
                    text_matches = re.findall(r'<w:t[^>]*>([^<]*)</w:t>', header_xml)
                    for i, text in enumerate(text_matches, 1):
                        if text.strip():
                            print(f"  {i}. Header text: '{text}'")
                    
                    # Cari placeholder di header
                    header_placeholders = re.findall(r'[\{\[\(]([^}\]\)]+)[\}\]\)]', header_xml)
                    if header_placeholders:
                        print(f"  Placeholders in header:")
                        for placeholder in set(header_placeholders):
                            print(f"    • {{{placeholder}}}")
                            
                except Exception as e:
                    print(f"  Error reading {header_file}: {e}")
    
    except Exception as e:
        print(f"❌ Error inspecting XML: {e}")

if __name__ == '__main__':
    inspect_template_content()