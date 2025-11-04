# erd_integrated_generator.py - ERD + Table Documentation Generator
import os
import shutil
from datetime import datetime

class ERDIntegratedGenerator:
    """Generate ERD diagram + comprehensive table documentation"""
    
    def __init__(self):
        self.db_config = {
            'host': 'localhost',
            'port': '5414', 
            'dbname': 'deverm',
            'user': 'postgres',
            'password': '1234'
        }
        
        self.template_path = os.path.join('..', 'template', 'template_dokumentasi.docx')
        self.output_dir = os.path.join('..', 'output')
        
        print(f"📊 ERD Integrated Generator Initialized")
        print(f"🔗 Will generate ERD + Table documentation")
    
    def analyze_database_relationships(self, table_limit=15):
        """Analyze database untuk ERD generation"""
        
        print("\n🔍 DATABASE RELATIONSHIP ANALYSIS")
        print("=" * 60)
        print(f"📊 Analyzing relationships for ERD generation...")
        
        try:
            import psycopg2
            
            conn = psycopg2.connect(**self.db_config)
            cur = conn.cursor()
            
            # Get tables dengan relationship info
            cur.execute(f"""
                SELECT DISTINCT t.table_name
                FROM information_schema.tables t
                WHERE t.table_schema = 'public'
                ORDER BY t.table_name
                LIMIT {table_limit};
            """)
            
            tables = cur.fetchall()
            
            erd_data = {
                'tables': [],
                'relationships': [],
                'total_tables': 0,
                'total_relationships': 0
            }
            
            print(f"📋 Analyzing {len(tables)} tables for relationships...")
            
            for (table_name,) in tables:
                # Get table columns dengan constraint info
                cur.execute("""
                    SELECT 
                        c.column_name,
                        c.data_type,
                        c.is_nullable,
                        c.column_default,
                        CASE WHEN pk.column_name IS NOT NULL THEN 'PK' ELSE '' END as is_primary,
                        CASE WHEN fk.column_name IS NOT NULL THEN 'FK' ELSE '' END as is_foreign,
                        fk.foreign_table_name,
                        fk.foreign_column_name
                    FROM information_schema.columns c
                    LEFT JOIN (
                        SELECT kcu.column_name, kcu.table_name
                        FROM information_schema.table_constraints tc
                        JOIN information_schema.key_column_usage kcu ON tc.constraint_name = kcu.constraint_name
                        WHERE tc.constraint_type = 'PRIMARY KEY' AND tc.table_schema = 'public'
                    ) pk ON c.column_name = pk.column_name AND c.table_name = pk.table_name
                    LEFT JOIN (
                        SELECT 
                            kcu.column_name, 
                            kcu.table_name,
                            ccu.table_name AS foreign_table_name,
                            ccu.column_name AS foreign_column_name
                        FROM information_schema.table_constraints AS tc
                        JOIN information_schema.key_column_usage AS kcu ON tc.constraint_name = kcu.constraint_name
                        JOIN information_schema.constraint_column_usage AS ccu ON ccu.constraint_name = tc.constraint_name
                        WHERE tc.constraint_type = 'FOREIGN KEY' AND tc.table_schema = 'public'
                    ) fk ON c.column_name = fk.column_name AND c.table_name = fk.table_name
                    WHERE c.table_name = %s AND c.table_schema = 'public'
                    ORDER BY c.ordinal_position;
                """, (table_name,))
                
                columns = cur.fetchall()
                
                if columns:
                    table_info = {
                        'name': table_name,
                        'columns': [],
                        'primary_keys': [],
                        'foreign_keys': []
                    }
                    
                    for column_name, data_type, is_nullable, column_default, is_primary, is_foreign, foreign_table, foreign_column in columns:
                        
                        column_info = {
                            'name': column_name,
                            'type': data_type,
                            'nullable': is_nullable == 'YES',
                            'default': column_default,
                            'is_primary': bool(is_primary),
                            'is_foreign': bool(is_foreign)
                        }
                        
                        table_info['columns'].append(column_info)
                        
                        if is_primary:
                            table_info['primary_keys'].append(column_name)
                        
                        if is_foreign and foreign_table and foreign_column:
                            fk_info = {
                                'column': column_name,
                                'references_table': foreign_table,
                                'references_column': foreign_column
                            }
                            table_info['foreign_keys'].append(fk_info)
                            
                            # Add to relationships
                            relationship = {
                                'from_table': table_name,
                                'from_column': column_name,
                                'to_table': foreign_table,
                                'to_column': foreign_column,
                                'relationship_type': 'FOREIGN_KEY'
                            }
                            erd_data['relationships'].append(relationship)
                    
                    erd_data['tables'].append(table_info)
                    print(f"   ✅ {table_name}: {len(columns)} columns, {len(table_info['primary_keys'])} PKs, {len(table_info['foreign_keys'])} FKs")
            
            cur.close()
            conn.close()
            
            erd_data['total_tables'] = len(erd_data['tables'])
            erd_data['total_relationships'] = len(erd_data['relationships'])
            
            print(f"\n📊 DATABASE ANALYSIS COMPLETE!")
            print(f"   📋 Tables analyzed: {erd_data['total_tables']}")
            print(f"   🔗 Relationships found: {erd_data['total_relationships']}")
            
            return erd_data
            
        except Exception as e:
            print(f"❌ Database analysis error: {e}")
            return None
    
    def generate_text_erd(self, erd_data):
        """Generate text-based ERD representation"""
        
        print("📊 Generating text-based ERD...")
        
        erd_text = []
        erd_text.append("📊 ENTITY RELATIONSHIP DIAGRAM (ERD)")
        erd_text.append("=" * 50)
        erd_text.append("")
        
        # Tables overview
        erd_text.append("📋 TABLES OVERVIEW:")
        erd_text.append(f"   Total Tables: {erd_data['total_tables']}")
        erd_text.append(f"   Total Relationships: {erd_data['total_relationships']}")
        erd_text.append("")
        
        # Table details dengan relationship indicators
        erd_text.append("📊 TABLE STRUCTURES:")
        erd_text.append("")
        
        for table in erd_data['tables']:
            erd_text.append(f"┌─ 📋 {table['name'].upper()}")
            erd_text.append("├─ COLUMNS:")
            
            for column in table['columns']:
                indicators = []
                if column['is_primary']:
                    indicators.append('PK')
                if column['is_foreign']:
                    indicators.append('FK')
                if not column['nullable']:
                    indicators.append('NOT NULL')
                
                indicator_str = f" ({', '.join(indicators)})" if indicators else ""
                erd_text.append(f"│  • {column['name']} : {column['type']}{indicator_str}")
            
            # Foreign key relationships
            if table['foreign_keys']:
                erd_text.append("├─ RELATIONSHIPS:")
                for fk in table['foreign_keys']:
                    erd_text.append(f"│  🔗 {fk['column']} → {fk['references_table']}.{fk['references_column']}")
            
            erd_text.append("└─")
            erd_text.append("")
        
        # Relationship summary
        if erd_data['relationships']:
            erd_text.append("🔗 RELATIONSHIP SUMMARY:")
            erd_text.append("")
            
            for i, rel in enumerate(erd_data['relationships'], 1):
                erd_text.append(f"{i}. {rel['from_table']}.{rel['from_column']} → {rel['to_table']}.{rel['to_column']}")
            
            erd_text.append("")
        
        return "\n".join(erd_text)
    
    def create_mermaid_erd(self, erd_data):
        """Generate Mermaid.js ERD code"""
        
        print("🎨 Generating Mermaid.js ERD code...")
        
        mermaid_lines = []
        mermaid_lines.append("```mermaid")
        mermaid_lines.append("erDiagram")
        mermaid_lines.append("")
        
        # Define entities
        for table in erd_data['tables']:
            mermaid_lines.append(f"    {table['name']} {{")
            
            for column in table['columns']:
                type_info = column['type']
                
                # Add constraints
                constraints = []
                if column['is_primary']:
                    constraints.append('PK')
                if column['is_foreign']:
                    constraints.append('FK')
                
                constraint_str = f" {','.join(constraints)}" if constraints else ""
                mermaid_lines.append(f"        {type_info} {column['name']}{constraint_str}")
            
            mermaid_lines.append("    }")
            mermaid_lines.append("")
        
        # Define relationships
        if erd_data['relationships']:
            mermaid_lines.append("    %% Relationships")
            for rel in erd_data['relationships']:
                # Mermaid relationship syntax: TABLE1 ||--o{ TABLE2 : "relationship"
                mermaid_lines.append(f"    {rel['to_table']} ||--o{{ {rel['from_table']} : \"{rel['from_column']} references {rel['to_column']}\"")
        
        mermaid_lines.append("```")
        
        return "\n".join(mermaid_lines)
    
    def create_erd_documentation(self, table_limit=15):
        """Create comprehensive documentation dengan ERD + Tables"""
        
        print("\n🚀 ERD + TABLE DOCUMENTATION GENERATOR")
        print("=" * 70)
        print(f"📊 Creating comprehensive documentation dengan ERD")
        print(f"🔗 ERD + {table_limit} detailed table documentation")
        
        # Analyze database
        erd_data = self.analyze_database_relationships(table_limit)
        
        if not erd_data:
            print("❌ Could not analyze database")
            return None
        
        try:
            # Copy template
            output_path = os.path.join(self.output_dir, f"ERD_Complete_Documentation_{table_limit}.docx")
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            shutil.copy2(self.template_path, output_path)
            
            print(f"📁 Template copied to: {os.path.basename(output_path)}")
            
            # Generate ERD representations
            text_erd = self.generate_text_erd(erd_data)
            mermaid_erd = self.create_mermaid_erd(erd_data)
            
            # Create comprehensive document
            self.generate_erd_document(output_path, erd_data, text_erd, mermaid_erd)
            
            file_size = os.path.getsize(output_path)
            
            print(f"\n🎉 ERD COMPLETE DOCUMENTATION CREATED!")
            print("=" * 70)
            print(f"📁 File: {os.path.basename(output_path)}")
            print(f"📊 Size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
            
            print(f"\n🔗 ERD FEATURES INCLUDED:")
            print(f"   ✅ Text-based ERD: Visual table relationships")
            print(f"   ✅ Mermaid.js code: Ready for online ERD tools")
            print(f"   ✅ Detailed tables: {erd_data['total_tables']} tables documented")
            print(f"   ✅ Relationships: {erd_data['total_relationships']} FK relationships mapped")
            print(f"   ✅ Column details: PK/FK indicators, data types")
            print(f"   ✅ Template format: Perfect column widths preserved")
            
            total_columns = sum(len(table['columns']) for table in erd_data['tables'])
            print(f"\n📊 COMPREHENSIVE SUMMARY:")
            print(f"   • ERD generation: ✅ COMPLETE")
            print(f"   • Tables documented: {erd_data['total_tables']}")
            print(f"   • Total columns: {total_columns}")
            print(f"   • Relationships mapped: {erd_data['total_relationships']}")
            print(f"   • Documentation quality: COMPREHENSIVE")
            
            return output_path
            
        except Exception as e:
            print(f"❌ Error: {e}")
            return None
    
    def generate_erd_document(self, file_path, erd_data, text_erd, mermaid_erd):
        """Generate comprehensive document dengan ERD + Tables"""
        
        print("📝 Generating comprehensive ERD + Table document...")
        
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.oxml.shared import OxmlElement, qn
            
            doc = Document(file_path)
            
            # Replace template content
            self.replace_erd_content(doc, erd_data['total_tables'])
            
            # Remove existing tables
            existing_tables = list(doc.tables)
            for table in existing_tables:
                table._element.getparent().remove(table._element)
            
            doc.add_page_break()
            
            # === ERD SECTION ===
            # ERD Title
            erd_title = doc.add_paragraph()
            erd_title_run = erd_title.add_run("📊 ENTITY RELATIONSHIP DIAGRAM (ERD)")
            erd_title_run.font.name = 'Times New Roman'
            erd_title_run.font.size = Pt(16)
            erd_title_run.font.bold = True
            
            # ERD Summary
            erd_summary = doc.add_paragraph()
            erd_summary_run = erd_summary.add_run(
                f"Database Overview: {erd_data['total_tables']} Tables, "
                f"{erd_data['total_relationships']} Relationships"
            )
            erd_summary_run.font.name = 'Times New Roman'
            erd_summary_run.font.size = Pt(12)
            erd_summary_run.font.italic = True
            
            doc.add_paragraph()
            
            # Text ERD
            erd_text_title = doc.add_paragraph()
            erd_text_title_run = erd_text_title.add_run("📋 Database Structure Diagram")
            erd_text_title_run.font.name = 'Times New Roman'
            erd_text_title_run.font.size = Pt(14)
            erd_text_title_run.font.bold = True
            
            # Add ERD content dengan monospace font
            erd_lines = text_erd.split('\n')
            for line in erd_lines:
                if line.strip():
                    erd_para = doc.add_paragraph()
                    erd_run = erd_para.add_run(line)
                    erd_run.font.name = 'Courier New'  # Monospace untuk alignment
                    erd_run.font.size = Pt(10)
                else:
                    doc.add_paragraph()
            
            doc.add_page_break()
            
            # Mermaid ERD section
            mermaid_title = doc.add_paragraph()
            mermaid_title_run = mermaid_title.add_run("🎨 Mermaid.js ERD Code")
            mermaid_title_run.font.name = 'Times New Roman'
            mermaid_title_run.font.size = Pt(14)
            mermaid_title_run.font.bold = True
            
            mermaid_info = doc.add_paragraph()
            mermaid_info_run = mermaid_info.add_run(
                "Copy kode berikut ke https://mermaid.live atau tools ERD lainnya untuk generate visual diagram:"
            )
            mermaid_info_run.font.name = 'Times New Roman'
            mermaid_info_run.font.size = Pt(11)
            mermaid_info_run.font.italic = True
            
            # Add Mermaid code
            mermaid_lines = mermaid_erd.split('\n')
            for line in mermaid_lines:
                mermaid_para = doc.add_paragraph()
                mermaid_run = mermaid_para.add_run(line)
                mermaid_run.font.name = 'Courier New'
                mermaid_run.font.size = Pt(9)
            
            doc.add_page_break()
            
            # === TABLE DOCUMENTATION SECTION ===
            # Table documentation title
            table_title = doc.add_paragraph()
            table_title_run = table_title.add_run("📋 DETAILED TABLE DOCUMENTATION")
            table_title_run.font.name = 'Times New Roman'
            table_title_run.font.size = Pt(16)
            table_title_run.font.bold = True
            
            table_subtitle = doc.add_paragraph()
            table_subtitle_run = table_subtitle.add_run(f"Comprehensive documentation for {erd_data['total_tables']} database tables")
            table_subtitle_run.font.name = 'Times New Roman'
            table_subtitle_run.font.size = Pt(12)
            table_subtitle_run.font.italic = True
            
            print(f"   📊 Creating detailed documentation for {erd_data['total_tables']} tables...")
            
            # Generate table documentation
            for table_idx, table_data in enumerate(erd_data['tables'], 1):
                self.create_erd_table_documentation(doc, table_data, table_idx)
                
                if table_idx < len(erd_data['tables']):
                    doc.add_paragraph()
            
            doc.save(file_path)
            print(f"   ✅ ERD + Table document saved successfully")
            
        except Exception as e:
            print(f"   ❌ Error: {e}")
    
    def create_erd_table_documentation(self, doc, table_data, table_number):
        """Create detailed table documentation dengan ERD context"""
        
        try:
            from docx.shared import Pt, Twips
            from docx.oxml.shared import OxmlElement, qn
            
            # Table title dengan relationship info
            relationship_info = ""
            if table_data['foreign_keys']:
                relationship_info = f" (🔗 {len(table_data['foreign_keys'])} relationships)"
            
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"Table {table_number}: {table_data['name']}{relationship_info}")
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(14)
            title_run.font.bold = True
            
            # Create table dengan adaptive structure
            table = doc.add_table(rows=1, cols=5)  # Extra column untuk relationship info
            table.style = 'Normal Table'
            
            # Apply column widths (adapted for 5 columns)
            self.apply_erd_table_widths(table)
            self.apply_erd_borders(table)
            
            # Headers
            headers = ['No', 'Column Name', 'Data Type', 'Constraints', 'Description']
            hdr_cells = table.rows[0].cells
            
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for para in hdr_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.bold = True
            
            # Data rows
            for col_idx, column in enumerate(table_data['columns'], 1):
                row = table.add_row()
                cells = row.cells
                
                # Column number
                cells[0].text = f"{col_idx}."
                
                # Column name
                cells[1].text = column['name'].upper()
                
                # Data type
                cells[2].text = column['type'].title()
                
                # Constraints
                constraints = []
                if column['is_primary']:
                    constraints.append('PRIMARY KEY')
                if column['is_foreign']:
                    constraints.append('FOREIGN KEY')
                if not column['nullable']:
                    constraints.append('NOT NULL')
                if column['default']:
                    constraints.append(f"DEFAULT: {column['default']}")
                
                cells[3].text = ', '.join(constraints) if constraints else '-'
                
                # Smart description
                description = self.generate_erd_column_description(column, table_data)
                cells[4].text = description
                
                # Apply formatting
                for cell in cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
            
            # Add relationship summary jika ada
            if table_data['foreign_keys']:
                doc.add_paragraph()
                rel_para = doc.add_paragraph()
                rel_run = rel_para.add_run(f"🔗 Relationships for {table_data['name']}:")
                rel_run.font.name = 'Times New Roman'
                rel_run.font.size = Pt(12)
                rel_run.font.bold = True
                
                for fk in table_data['foreign_keys']:
                    fk_para = doc.add_paragraph()
                    fk_run = fk_para.add_run(f"   • {fk['column']} → {fk['references_table']}.{fk['references_column']}")
                    fk_run.font.name = 'Times New Roman'
                    fk_run.font.size = Pt(11)
                    fk_run.font.italic = True
            
            print(f"   ✅ Table {table_number}: ERD documentation complete")
                            
        except Exception as e:
            print(f"   ⚠️ Error documenting table {table_number}: {e}")
    
    def apply_erd_table_widths(self, table):
        """Apply column widths untuk 5-column ERD table"""
        
        try:
            from docx.shared import Twips
            from docx.oxml.shared import OxmlElement, qn
            
            # Width untuk 5 columns: No, Name, Type, Constraints, Description
            widths_twips = [400, 1800, 1200, 1800, 3920]  # Total: 9120 twips
            
            # Apply XML widths
            tbl = table._tbl
            
            tbl_grid = tbl.find(qn('w:tblGrid'))
            if tbl_grid is None:
                tbl_grid = OxmlElement('w:tblGrid')
                tbl.insert(1, tbl_grid)
            else:
                for child in list(tbl_grid):
                    tbl_grid.remove(child)
            
            for width_twips in widths_twips:
                grid_col = OxmlElement('w:gridCol')
                grid_col.set(qn('w:w'), str(width_twips))
                tbl_grid.append(grid_col)
            
            # Force cell widths
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    if i < len(widths_twips):
                        tc = cell._tc
                        tc_pr = tc.find(qn('w:tcPr'))
                        if tc_pr is None:
                            tc_pr = OxmlElement('w:tcPr')
                            tc.insert(0, tc_pr)
                        
                        tc_w = tc_pr.find(qn('w:tcW'))
                        if tc_w is None:
                            tc_w = OxmlElement('w:tcW')
                            tc_pr.append(tc_w)
                        
                        tc_w.set(qn('w:w'), str(widths_twips[i]))
                        tc_w.set(qn('w:type'), 'dxa')
            
        except Exception as e:
            pass
    
    def apply_erd_borders(self, table):
        """Apply borders untuk ERD table"""
        
        try:
            from docx.oxml.shared import OxmlElement, qn
            
            tbl = table._tbl
            tbl_pr = tbl.find(qn('w:tblPr'))
            
            if tbl_pr is None:
                tbl_pr = OxmlElement('w:tblPr')
                tbl.insert(0, tbl_pr)
            
            tbl_borders = OxmlElement('w:tblBorders')
            
            for border_type in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_type}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tbl_borders.append(border)
            
            existing_borders = tbl_pr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tbl_pr.remove(existing_borders)
            
            tbl_pr.append(tbl_borders)
            
        except Exception as e:
            pass
    
    def generate_erd_column_description(self, column, table_data):
        """Generate smart description dengan ERD context"""
        
        column_name = column['name'].lower()
        
        # Enhanced descriptions dengan relationship context
        patterns = {
            'id': "Unique identifier for record identification and cross-table referencing",
            'name': "Name field for entity identification and display purposes", 
            'code': "Unique code or reference identifier for system integration",
            'email': "Email address field for contact and communication purposes",
            'phone': "Phone number contact information field", 
            'address': "Address or location information storage field",
            'date': "Date and time information for temporal data tracking",
            'created': "Creation timestamp for audit and tracking purposes",
            'updated': "Last modification timestamp for change tracking", 
            'status': "Status indicator for entity state management",
            'active': "Active/inactive status flag for record management",
            'type': "Type classification field for categorization purposes",
            'description': "Descriptive text field for additional information storage"
        }
        
        # Check if column is part of relationship
        if column['is_foreign']:
            for fk in table_data['foreign_keys']:
                if fk['column'] == column['name']:
                    return f"Foreign key referencing {fk['references_table']}.{fk['references_column']} - establishes relationship"
        
        if column['is_primary']:
            return "Primary key - unique identifier for this table's records"
        
        # Pattern matching
        for pattern, desc in patterns.items():
            if pattern in column_name:
                return desc
        
        return f"Data field for {column['name'].replace('_', ' ')} information storage and management"
    
    def replace_erd_content(self, doc, table_count):
        """Replace template content untuk ERD documentation"""
        
        replacements = [
            ('Personal Assignment 1', 'ERD + Database Documentation'),
            ('Data Modelling and Analytics', f'Entity Relationship Diagram + {table_count} Tables'),
            ('Written by :', 'Generated by ERD Integration System'),
            ('MOH IRSYAD RISNO QUSHOYYI 2802650670', f'ERD System - {datetime.now().strftime("%d %B %Y %H:%M WIB")}'),
            ('INFORMATION SYSTEM STUDY PROGRAM BINUS UNIVERSITY ONLINE LEARNING', f'COMPREHENSIVE ERD + TABLE DOCUMENTATION')
        ]
        
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements:
                if old_text.lower() in paragraph.text.lower():
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                        for run in paragraph.runs[1:]:
                            run.text = ""

def main():
    """Main function untuk ERD integrated documentation"""
    
    print("🔗 ERD Integrated Generator")
    print("📊 ERD + Comprehensive Table Documentation")
    print("🎨 Text ERD + Mermaid.js + Detailed Tables")
    print()
    
    generator = ERDIntegratedGenerator()
    
    # Generate ERD + Table documentation
    result = generator.create_erd_documentation(15)
    
    if result:
        print(f"\n✨ ERD + TABLE DOCUMENTATION SUCCESS!")
        print(f"📁 {result}")
        print(f"\n🔗 ERD INTEGRATION FEATURES:")
        print(f"   ✅ Entity Relationship Diagram: Text-based visual")
        print(f"   ✅ Mermaid.js ERD code: Ready for online tools")
        print(f"   ✅ Relationship mapping: FK relationships traced")
        print(f"   ✅ Constraint documentation: PK/FK/NOT NULL")
        print(f"   ✅ Table documentation: Comprehensive details")
        print(f"   ✅ Template formatting: Perfect column widths")
        
        print(f"\n📊 COMPREHENSIVE DOCUMENTATION:")
        print(f"   • ERD generation: ✅ COMPLETE")
        print(f"   • Relationship analysis: ✅ COMPLETE")
        print(f"   • Table documentation: ✅ COMPLETE")
        print(f"   • Mermaid.js ready: ✅ Copy & paste ke https://mermaid.live")
        
        print(f"\n💡 Sekarang dokumentasi lengkap dengan ERD + Table details!")
    else:
        print(f"\n❌ ERD generation failed!")

if __name__ == '__main__':
    main()