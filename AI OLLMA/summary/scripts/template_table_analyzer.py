# template_table_analyzer.py - Analisis detail struktur tabel template
import os
import zipfile
import xml.etree.ElementTree as ET

def analyze_template_table_structure():
    """Analisis struktur tabel template secara detail"""
    
    template_path = os.path.join('..', 'template', 'template_dokumentasi.docx')
    
    if not os.path.exists(template_path):
        print(f"❌ Template tidak ditemukan: {template_path}")
        return
    
    print("🔍 TEMPLATE TABLE STRUCTURE ANALYZER")
    print("=" * 60)
    
    try:
        from docx import Document
        
        # Baca template dengan python-docx
        doc = Document(template_path)
        
        print("📊 ANALYZING TABLES IN TEMPLATE:")
        print("-" * 40)
        
        table_structures = []
        
        for table_idx, table in enumerate(doc.tables):
            print(f"\n📋 TABLE {table_idx + 1}:")
            print(f"   Dimensions: {len(table.rows)} rows × {len(table.columns)} columns")
            print(f"   Style: {table.style.name if table.style else 'No style'}")
            
            table_structure = {
                'index': table_idx,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'style': table.style.name if table.style else None,
                'content': []
            }
            
            # Analisis setiap cell
            for row_idx, row in enumerate(table.rows):
                row_content = []
                print(f"\n   ROW {row_idx + 1}:")
                
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    cell_info = {
                        'text': cell_text,
                        'col_index': col_idx,
                        'formatting': []
                    }
                    
                    print(f"     Col {col_idx + 1}: '{cell_text}'")
                    
                    # Analisis formatting dalam cell
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text.strip():
                                format_info = {
                                    'text': run.text,
                                    'font_name': run.font.name,
                                    'font_size': run.font.size.pt if run.font.size else None,
                                    'bold': run.font.bold,
                                    'italic': run.font.italic
                                }
                                cell_info['formatting'].append(format_info)
                                
                                if format_info['font_name'] or format_info['font_size']:
                                    print(f"         Format: {format_info['font_name']}, {format_info['font_size']}pt, Bold:{format_info['bold']}")
                    
                    row_content.append(cell_info)
                
                table_structure['content'].append(row_content)
            
            table_structures.append(table_structure)
            
            # Analisis border/garis tabel
            print(f"\n   🎨 TABLE STYLING:")
            print(f"     Style applied: {table.style.name if table.style else 'None'}")
            
        # Analisis XML level untuk border details
        print(f"\n🔍 XML LEVEL TABLE ANALYSIS:")
        print("-" * 40)
        
        with zipfile.ZipFile(template_path, 'r') as zip_file:
            if 'word/document.xml' in zip_file.namelist():
                doc_xml = zip_file.read('word/document.xml').decode('utf-8')
                
                # Cari table elements
                import re
                
                # Cari table borders
                border_matches = re.findall(r'<w:tblBorders[^>]*>(.*?)</w:tblBorders>', doc_xml, re.DOTALL)
                if border_matches:
                    print(f"   📐 Table borders found: {len(border_matches)} definitions")
                    for i, border in enumerate(border_matches):
                        if 'w:top' in border:
                            print(f"     Border {i+1}: Has top border")
                        if 'w:bottom' in border:
                            print(f"     Border {i+1}: Has bottom border")
                        if 'w:left' in border:
                            print(f"     Border {i+1}: Has left border")
                        if 'w:right' in border:
                            print(f"     Border {i+1}: Has right border")
                        if 'w:insideH' in border:
                            print(f"     Border {i+1}: Has inside horizontal borders")
                        if 'w:insideV' in border:
                            print(f"     Border {i+1}: Has inside vertical borders")
                
                # Cari table styles
                style_matches = re.findall(r'<w:tblStyle w:val="([^"]*)"', doc_xml)
                if style_matches:
                    print(f"   🎨 Table styles used: {', '.join(set(style_matches))}")
                
                # Cari table content structure
                cell_matches = re.findall(r'<w:tc[^>]*>(.*?)</w:tc>', doc_xml, re.DOTALL)
                print(f"   📊 Total cells found in XML: {len(cell_matches)}")
        
        return table_structures
        
    except Exception as e:
        print(f"❌ Error analyzing template: {e}")
        return []

def print_template_structure_summary(table_structures):
    """Print ringkasan struktur untuk implementasi"""
    
    print(f"\n📋 TEMPLATE STRUCTURE SUMMARY FOR IMPLEMENTATION:")
    print("=" * 60)
    
    for i, table_struct in enumerate(table_structures):
        print(f"\nTable {i+1} Implementation Guide:")
        print(f"  Dimensions: {table_struct['rows']} rows × {table_struct['cols']} columns")
        print(f"  Style to apply: '{table_struct['style']}'")
        
        if table_struct['content']:
            print(f"  Content Structure:")
            for row_idx, row_content in enumerate(table_struct['content']):
                print(f"    Row {row_idx + 1}:")
                for col_idx, cell_info in enumerate(row_content):
                    if cell_info['text']:
                        print(f"      Col {col_idx + 1}: '{cell_info['text'][:30]}...'")
                        if cell_info['formatting']:
                            format_info = cell_info['formatting'][0]
                            print(f"        → Font: {format_info['font_name']}, Size: {format_info['font_size']}, Bold: {format_info['bold']}")

def main():
    """Main function"""
    print("🚀 Template Table Structure Analyzer")
    print("🔍 Analyzing template table structure for exact replication")
    print()
    
    table_structures = analyze_template_table_structure()
    
    if table_structures:
        print_template_structure_summary(table_structures)
        
        print(f"\n✨ ANALYSIS COMPLETE!")
        print(f"📊 Found {len(table_structures)} table(s) in template")
        print(f"🎯 Structure data ready for implementation")
        
        return table_structures
    else:
        print(f"\n❌ No tables found or analysis failed")
        return []

if __name__ == '__main__':
    main()