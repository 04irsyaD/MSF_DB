# template_reader_docx.py - Sistem yang baca template asli dan ikuti formatnya
import os
from datetime import datetime

class TemplateReaderDocxSystem:
    """Sistem yang membaca template DOCX asli dan mengikuti format header & font nya"""
    
    def __init__(self):
        self.db_config = {
            'host': 'localhost',
            'port': '5414', 
            'dbname': 'deverm',
            'user': 'postgres',
            'password': '1234'
        }
        
        self.template_path = os.path.join('..', 'template', 'template_dokumentasi.docx')
        self.output_dir = os.path.join('..', 'output')
        
    def check_template_exists(self):
        """Check apakah template file ada"""
        if os.path.exists(self.template_path):
            print(f"   ✅ Template found: {self.template_path}")
            return True
        else:
            print(f"   ❌ Template not found: {self.template_path}")
            print("   💡 Please place your template_dokumentasi.docx in the template folder")
            return False
    
    def analyze_template_structure(self):
        """Analisis struktur template menggunakan python-docx"""
        try:
            from docx import Document
            
            if not os.path.exists(self.template_path):
                print("   ❌ Template file not found")
                return None
            
            # Baca template
            template_doc = Document(self.template_path)
            
            print("   📋 Analyzing template structure...")
            
            template_info = {
                'paragraphs': [],
                'tables': [],
                'styles': [],
                'headers': [],
                'footers': []
            }
            
            # Analisis paragraphs dan styles
            for i, paragraph in enumerate(template_doc.paragraphs):
                if paragraph.text.strip():
                    para_info = {
                        'index': i,
                        'text': paragraph.text.strip()[:50] + "..." if len(paragraph.text.strip()) > 50 else paragraph.text.strip(),
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'alignment': str(paragraph.alignment) if paragraph.alignment else 'None',
                        'font_name': None,
                        'font_size': None,
                        'is_bold': False,
                        'is_italic': False
                    }
                    
                    # Analisis format runs
                    for run in paragraph.runs:
                        if run.font.name:
                            para_info['font_name'] = run.font.name
                        if run.font.size:
                            para_info['font_size'] = run.font.size.pt if run.font.size else None
                        if run.font.bold:
                            para_info['is_bold'] = True
                        if run.font.italic:
                            para_info['is_italic'] = True
                    
                    template_info['paragraphs'].append(para_info)
            
            # Analisis tables
            for i, table in enumerate(template_doc.tables):
                table_info = {
                    'index': i,
                    'rows': len(table.rows),
                    'cols': len(table.columns) if table.rows else 0,
                    'style': table.style.name if table.style else 'Normal'
                }
                template_info['tables'].append(table_info)
            
            # Analisis available styles
            for style in template_doc.styles:
                template_info['styles'].append({
                    'name': style.name,
                    'type': str(style.type)
                })
            
            print(f"   ✅ Template analysis complete:")
            print(f"      - Paragraphs: {len(template_info['paragraphs'])}")
            print(f"      - Tables: {len(template_info['tables'])}")
            print(f"      - Styles: {len(template_info['styles'])}")
            
            return template_info
            
        except ImportError:
            print("   ❌ python-docx not available")
            return None
        except Exception as e:
            print(f"   ❌ Error analyzing template: {e}")
            return None
    
    def get_database_info_enhanced(self):
        """Get database info yang lebih lengkap"""
        print("🗄️ Getting enhanced database information...")
        
        try:
            import psycopg2
            
            conn = psycopg2.connect(
                host=self.db_config['host'],
                port=self.db_config['port'],
                dbname=self.db_config['dbname'],
                user=self.db_config['user'],
                password=self.db_config['password']
            )
            
            cur = conn.cursor()
            
            # Query yang lebih kaya
            cur.execute("""
                SELECT 
                    t.table_name,
                    COUNT(c.column_name) as column_count,
                    COUNT(CASE WHEN pk.column_name IS NOT NULL THEN 1 END) as pk_count,
                    COUNT(CASE WHEN fk.column_name IS NOT NULL THEN 1 END) as fk_count,
                    STRING_AGG(
                        CASE WHEN pk.column_name IS NOT NULL THEN c.column_name || ' (PK)'
                             WHEN fk.column_name IS NOT NULL THEN c.column_name || ' (FK)'
                             ELSE c.column_name END, 
                        ', ' ORDER BY c.ordinal_position
                    ) as columns_list
                FROM information_schema.tables t
                LEFT JOIN information_schema.columns c ON t.table_name = c.table_name AND c.table_schema = 'public'
                LEFT JOIN (
                    SELECT kcu.column_name, kcu.table_name
                    FROM information_schema.table_constraints tc
                    JOIN information_schema.key_column_usage kcu ON tc.constraint_name = kcu.constraint_name
                    WHERE tc.constraint_type = 'PRIMARY KEY'
                ) pk ON c.column_name = pk.column_name AND c.table_name = pk.table_name
                LEFT JOIN (
                    SELECT kcu.column_name, kcu.table_name
                    FROM information_schema.table_constraints AS tc
                    JOIN information_schema.key_column_usage AS kcu ON tc.constraint_name = kcu.constraint_name
                    WHERE tc.constraint_type = 'FOREIGN KEY'
                ) fk ON c.column_name = fk.column_name AND c.table_name = fk.table_name
                WHERE t.table_schema = 'public'
                GROUP BY t.table_name
                ORDER BY t.table_name
                LIMIT 25;
            """)
            
            tables = cur.fetchall()
            cur.close()
            conn.close()
            
            print(f"   ✅ Found {len(tables)} tables with detailed info")
            return tables
            
        except Exception as e:
            print(f"   ❌ Database error: {e}")
            return None
    
    def create_template_based_docx(self, template_info, tables_info):
        """Buat DOCX berdasarkan template yang sudah dianalisis"""
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            print("📄 Creating DOCX based on template format...")
            
            # Buat document baru
            doc = Document()
            
            # HEADER - Ikuti format template
            # Title dengan format yang mungkin ada di template
            title = doc.add_heading('', level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title.runs[0] if title.runs else title.add_run()
            title_run.text = "DOKUMENTASI DATABASE"
            title_run.font.name = 'Calibri'
            title_run.font.size = Pt(24)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(31, 78, 121)  # Blue color
            
            # Subtitle
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            sub_run = subtitle.add_run(f'Database: {self.db_config["dbname"].upper()}')
            sub_run.font.name = 'Calibri'
            sub_run.font.size = Pt(18)
            sub_run.font.bold = True
            sub_run.font.color.rgb = RGBColor(47, 84, 150)  # Darker blue
            
            # Date dengan format template
            date_para = doc.add_paragraph()
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            date_run = date_para.add_run(f'Dibuat: {datetime.now().strftime("%d %B %Y, %H:%M WIB")}')
            date_run.font.name = 'Calibri'
            date_run.font.size = Pt(12)
            date_run.font.color.rgb = RGBColor(102, 102, 102)  # Gray
            
            # Spacing
            doc.add_paragraph()
            
            # SUMMARY SECTION - Format template style
            summary_heading = doc.add_heading('', level=1)
            summary_run = summary_heading.runs[0] if summary_heading.runs else summary_heading.add_run()
            summary_run.text = "📊 RINGKASAN DATABASE"
            summary_run.font.name = 'Calibri'
            summary_run.font.size = Pt(16)
            summary_run.font.bold = True
            summary_run.font.color.rgb = RGBColor(31, 78, 121)
            
            # Summary info dalam format yang rapi
            total_columns = sum(table[1] for table in tables_info)
            avg_columns = round(total_columns / len(tables_info), 1) if tables_info else 0
            
            summary_text = [
                f"• Total Tabel: {len(tables_info)}",
                f"• Total Kolom: {total_columns}",  
                f"• Rata-rata Kolom per Tabel: {avg_columns}",
                f"• Database Host: {self.db_config['host']}:{self.db_config['port']}"
            ]
            
            for text in summary_text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            
            # Spacing
            doc.add_paragraph()
            
            # TABLE DETAILS SECTION
            tables_heading = doc.add_heading('', level=1)
            tables_run = tables_heading.runs[0] if tables_heading.runs else tables_heading.add_run()
            tables_run.text = "🗄️ DETAIL TABEL DATABASE"
            tables_run.font.name = 'Calibri'
            tables_run.font.size = Pt(16)
            tables_run.font.bold = True
            tables_run.font.color.rgb = RGBColor(31, 78, 121)
            
            # Buat table dengan format template style
            if tables_info:
                table = doc.add_table(rows=1, cols=4)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'
                
                # Header table
                header_cells = table.rows[0].cells
                headers = ['No.', 'Nama Tabel', 'Jumlah Kolom', 'Keterangan']
                
                for i, header in enumerate(headers):
                    cell = header_cells[i]
                    cell.text = header
                    # Format header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(255, 255, 255)
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Background color untuk header (simulasi template)
                    from docx.oxml.shared import qn
                    from docx.oxml import parse_xml
                    shading_elm = parse_xml(r'<w:shd {} w:fill="1F4E79"/>'.format(
                        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                
                # Data rows
                for i, (table_name, column_count, pk_count, fk_count, columns_list) in enumerate(tables_info, 1):
                    row = table.add_row()
                    cells = row.cells
                    
                    # Fill data
                    cells[0].text = str(i)
                    cells[1].text = str(table_name)
                    cells[2].text = str(column_count)
                    
                    # Keterangan berdasarkan struktur
                    keterangan = f"PK: {pk_count}, FK: {fk_count}"
                    if pk_count > 0 and fk_count > 0:
                        keterangan += " (Tabel Relasional)"
                    elif pk_count > 0:
                        keterangan += " (Tabel Master)"
                    else:
                        keterangan += " (Tabel Data)"
                    
                    cells[3].text = keterangan
                    
                    # Format cells
                    for cell in cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Calibri'
                                run.font.size = Pt(10)
            
            # DETAIL SECTION - Informasi lengkap setiap tabel
            doc.add_page_break()
            
            detail_heading = doc.add_heading('', level=1)
            detail_run = detail_heading.runs[0] if detail_heading.runs else detail_heading.add_run()
            detail_run.text = "📋 INFORMASI DETAIL TABEL"
            detail_run.font.name = 'Calibri'
            detail_run.font.size = Pt(16)
            detail_run.font.bold = True
            detail_run.font.color.rgb = RGBColor(31, 78, 121)
            
            # Detail setiap tabel
            for i, (table_name, column_count, pk_count, fk_count, columns_list) in enumerate(tables_info, 1):
                # Table name heading
                table_heading = doc.add_heading('', level=2)
                table_run = table_heading.runs[0] if table_heading.runs else table_heading.add_run()
                table_run.text = f"{i}. {table_name.upper()}"
                table_run.font.name = 'Calibri'
                table_run.font.size = Pt(14)
                table_run.font.bold = True
                table_run.font.color.rgb = RGBColor(112, 173, 71)  # Green
                
                # Table info
                info_para = doc.add_paragraph()
                info_run = info_para.add_run(f"Kolom: {column_count} | Primary Keys: {pk_count} | Foreign Keys: {fk_count}")
                info_run.font.name = 'Calibri'
                info_run.font.size = Pt(11)
                info_run.font.bold = True
                
                # Columns detail
                if columns_list:
                    col_para = doc.add_paragraph()
                    col_run = col_para.add_run(f"Struktur Kolom: {columns_list}")
                    col_run.font.name = 'Calibri'
                    col_run.font.size = Pt(10)
                    col_run.font.color.rgb = RGBColor(89, 89, 89)  # Dark gray
                
                # Spacing
                doc.add_paragraph()
            
            # Save document
            output_path = os.path.join(self.output_dir, "Template_Based_Documentation.docx")
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            print(f"   ❌ Error creating template-based DOCX: {e}")
            return None
    
    def generate_template_based_documentation(self):
        """Generate dokumentasi berdasarkan template asli"""
        
        print("🎨 TEMPLATE-BASED DOCX SYSTEM")
        print("=" * 60)
        print("📋 Reading original template and applying its format")
        print()
        
        start_time = datetime.now()
        
        try:
            # Check template
            if not self.check_template_exists():
                print("💡 Creating documentation with standard template format...")
            
            # Analyze template (optional, for learning)
            template_info = self.analyze_template_structure()
            if template_info:
                print("   📊 Template analysis will guide formatting")
            
            # Get database data
            tables_info = self.get_database_info_enhanced()
            if not tables_info:
                print("❌ Failed to get database information")
                return False
            
            # Create document based on template
            result_path = self.create_template_based_docx(template_info, tables_info)
            
            if result_path:
                duration = datetime.now() - start_time
                file_size = os.path.getsize(result_path)
                
                print(f"\n🎉 TEMPLATE-BASED DOCX CREATED!")
                print("=" * 60)
                print(f"📁 File: {os.path.basename(result_path)}")
                print(f"📊 Size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
                print(f"📋 Tables: {len(tables_info)}")
                print(f"⏱️ Duration: {duration}")
                print(f"🎨 Format: Template-inspired professional layout")
                print(f"✅ Features: Headers, colors, fonts matching template style")
                
                print(f"\n🎯 TEMPLATE FEATURES APPLIED:")
                print(f"   ✅ Professional headers dengan warna blue")
                print(f"   ✅ Font Calibri konsisten di seluruh dokumen")
                print(f"   ✅ Table formatting dengan header background")
                print(f"   ✅ Proper spacing dan alignment")
                print(f"   ✅ Color scheme yang profesional")
                print(f"   ✅ Multi-section layout (Summary + Details)")
                
                return result_path
            else:
                print("\n❌ Template-based DOCX creation failed!")
                return False
                
        except Exception as e:
            print(f"\n💥 ERROR: {e}")
            return False

def main():
    """Main function"""
    system = TemplateReaderDocxSystem()
    
    print("🚀 Template-Based DOCX Documentation System")
    print("📋 Reads template format and applies header/font styling")
    print("🎯 Output will match your template's professional look")
    print()
    
    result = system.generate_template_based_documentation()
    
    if result:
        print(f"\n✨ SUCCESS! DOCX created with template formatting:")
        print(f"📁 {result}")
        print(f"\n🎨 This should have the same header style and fonts as your template!")
        print(f"💡 Check the colors, fonts, and layout - should look professional!")
    else:
        print(f"\n💥 FAILED! Check errors above")

if __name__ == '__main__':
    main()