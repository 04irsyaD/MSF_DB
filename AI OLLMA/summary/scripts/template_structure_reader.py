# template_structure_reader.py - AI khusus membaca daftar isi dan struktur template
import os
from datetime import datetime

class TemplateStructureReader:
    """AI khusus untuk membaca daftar isi, headers, dan struktur template"""
    
    def __init__(self):
        self.template_path = os.path.join('template', 'template_dokumentasi.docx')
        
        print(f"📖 Template Structure Reader AI")
        print(f"🔍 Specialized in reading table of contents and document structure")
        print(f"📋 Will extract all structural elements from template")
    
    def read_template_structure(self):
        """Baca struktur template: daftar isi, headers, sections"""
        
        print("\n" + "="*80)
        print("📖 TEMPLATE STRUCTURE & TABLE OF CONTENTS ANALYSIS")
        print("="*80)
        print(f"📁 Template: {self.template_path}")
        print(f"⏰ Analysis Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print("="*80)
        
        if not os.path.exists(self.template_path):
            print(f"❌ Template not found: {self.template_path}")
            return None
        
        try:
            from docx import Document
            
            # Load template
            doc = Document(self.template_path)
            
            print(f"📊 Document loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
            print()
            
            # Analyze document structure
            structure_data = {
                'document_outline': self.extract_document_outline(doc),
                'table_of_contents': self.extract_table_of_contents(doc),
                'section_headers': self.extract_section_headers(doc),
                'document_flow': self.analyze_document_flow(doc),
                'content_hierarchy': self.build_content_hierarchy(doc)
            }
            
            return structure_data
            
        except Exception as e:
            print(f"❌ Structure reading error: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def extract_document_outline(self, doc):
        """Extract dokumen outline dan struktur utama"""
        
        print("📋 DOCUMENT OUTLINE EXTRACTION")
        print("-" * 60)
        
        outline_elements = {
            'cover_page': [],
            'table_of_contents': [],
            'main_sections': [],
            'appendices': [],
            'references': []
        }
        
        current_section = 'unknown'
        
        # Scan paragraphs untuk outline
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            text_lower = text.lower()
            
            # Detect different sections
            if self.is_cover_element(text):
                outline_elements['cover_page'].append({
                    'index': para_idx,
                    'text': text,
                    'type': self.categorize_cover_element(text)
                })
                current_section = 'cover'
            
            elif self.is_toc_element(text):
                outline_elements['table_of_contents'].append({
                    'index': para_idx,
                    'text': text,
                    'level': self.get_toc_level(text)
                })
                current_section = 'toc'
            
            elif self.is_main_section_header(text):
                outline_elements['main_sections'].append({
                    'index': para_idx,
                    'text': text,
                    'level': self.get_header_level(text, para),
                    'section': current_section
                })
                current_section = 'main'
            
            elif self.is_appendix_element(text):
                outline_elements['appendices'].append({
                    'index': para_idx,
                    'text': text,
                    'type': 'appendix'
                })
                current_section = 'appendix'
        
        # Display results
        print("📖 COVER PAGE ELEMENTS:")
        for i, element in enumerate(outline_elements['cover_page'][:10]):
            print(f"   {i+1}. [{element['type']}] {element['text'][:60]}...")
        
        print(f"\n📋 TABLE OF CONTENTS FOUND:")
        if outline_elements['table_of_contents']:
            for i, toc_item in enumerate(outline_elements['table_of_contents'][:15]):
                level_indent = "  " * (toc_item.get('level', 0))
                print(f"   {level_indent}• {toc_item['text']}")
        else:
            print("   ❌ No table of contents detected")
        
        print(f"\n📚 MAIN SECTIONS FOUND:")
        for i, section in enumerate(outline_elements['main_sections'][:20]):
            level = section.get('level', 0)
            indent = "  " * level
            print(f"   {indent}📂 Level {level}: {section['text'][:80]}...")
        
        return outline_elements
    
    def extract_table_of_contents(self, doc):
        """Extract detailed table of contents"""
        
        print(f"\n📋 DETAILED TABLE OF CONTENTS ANALYSIS")
        print("-" * 60)
        
        toc_data = {
            'detected_toc': [],
            'chapter_structure': {},
            'numbering_pattern': [],
            'page_references': []
        }
        
        # Look for TOC patterns
        toc_indicators = [
            'daftar isi', 'table of contents', 'contents', 'indeks',
            'bab', 'chapter', 'bagian', 'section'
        ]
        
        numbering_patterns = [
            r'^\d+\.', r'^[IVX]+\.', r'^[A-Z]\.', r'^\d+\.\d+',
            r'^\d+\.\d+\.\d+', r'^[a-z]\)', r'^\([a-z]\)'
        ]
        
        in_toc_section = False
        current_chapter = None
        
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            text_lower = text.lower()
            
            # Check if we're entering TOC section
            if any(indicator in text_lower for indicator in toc_indicators):
                in_toc_section = True
                toc_data['detected_toc'].append({
                    'start_index': para_idx,
                    'header': text,
                    'type': 'toc_start'
                })
                print(f"📋 TOC Section Started: \"{text}\"")
                continue
            
            # If in TOC, analyze structure
            if in_toc_section:
                # Check for numbering patterns
                import re
                for pattern in numbering_patterns:
                    if re.match(pattern, text):
                        level = self.determine_toc_level(text)
                        
                        toc_item = {
                            'index': para_idx,
                            'text': text,
                            'level': level,
                            'pattern': pattern,
                            'number': re.match(pattern, text).group() if re.match(pattern, text) else None
                        }
                        
                        toc_data['detected_toc'].append(toc_item)
                        toc_data['numbering_pattern'].append(pattern)
                        
                        # Build chapter structure
                        if level == 1:
                            current_chapter = text
                            toc_data['chapter_structure'][current_chapter] = []
                        elif level > 1 and current_chapter:
                            if current_chapter not in toc_data['chapter_structure']:
                                toc_data['chapter_structure'][current_chapter] = []
                            toc_data['chapter_structure'][current_chapter].append(text)
                        
                        break
                
                # Check for page numbers
                if '...' in text or '.' * 3 in text:
                    page_match = re.search(r'(\d+)$', text)
                    if page_match:
                        toc_data['page_references'].append({
                            'text': text,
                            'page': page_match.group(1)
                        })
                
                # Check if TOC section ended
                if len(text) > 100 or 'pendahuluan' in text_lower or 'introduction' in text_lower:
                    in_toc_section = False
        
        # Display detailed TOC
        print(f"📚 DETAILED TABLE OF CONTENTS:")
        if toc_data['detected_toc']:
            current_level = 0
            for item in toc_data['detected_toc'][:25]:
                if 'level' in item:
                    level = item['level']
                    indent = "  " * level
                    number = item.get('number', '')
                    text_clean = item['text'].replace(number, '').strip()
                    print(f"   {indent}{number} {text_clean}")
                else:
                    print(f"   📋 {item.get('header', item.get('text', 'Unknown'))}")
        
        print(f"\n📊 TOC STATISTICS:")
        print(f"   Total TOC Items: {len([item for item in toc_data['detected_toc'] if 'level' in item])}")
        print(f"   Numbering Patterns: {len(set(toc_data['numbering_pattern']))}")
        print(f"   Page References: {len(toc_data['page_references'])}")
        print(f"   Chapters Detected: {len(toc_data['chapter_structure'])}")
        
        if toc_data['chapter_structure']:
            print(f"\n📚 CHAPTER STRUCTURE:")
            for chapter, subsections in list(toc_data['chapter_structure'].items())[:5]:
                print(f"   📂 {chapter}")
                for subsection in subsections[:3]:
                    print(f"      📄 {subsection}")
                if len(subsections) > 3:
                    print(f"      ... and {len(subsections) - 3} more subsections")
        
        return toc_data
    
    def extract_section_headers(self, doc):
        """Extract all section headers and their hierarchy"""
        
        print(f"\n📚 SECTION HEADERS ANALYSIS")
        print("-" * 60)
        
        headers_data = {
            'all_headers': [],
            'header_hierarchy': {},
            'formatting_patterns': {},
            'header_types': {}
        }
        
        header_indicators = [
            'bab', 'chapter', 'bagian', 'section', 'sub',
            'pendahuluan', 'introduction', 'kesimpulan', 'conclusion',
            'metodologi', 'methodology', 'hasil', 'results',
            'pembahasan', 'discussion', 'daftar', 'list'
        ]
        
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this looks like a header
            is_header = False
            header_type = 'unknown'
            level = 0
            
            # Method 1: Check paragraph style
            if para.style and 'heading' in para.style.name.lower():
                is_header = True
                header_type = 'style_based'
                level = self.extract_heading_level_from_style(para.style.name)
            
            # Method 2: Check formatting (bold, size, etc.)
            elif self.is_formatted_header(para):
                is_header = True
                header_type = 'format_based'
                level = self.determine_format_level(para)
            
            # Method 3: Check content patterns
            elif self.contains_header_indicators(text, header_indicators):
                is_header = True
                header_type = 'content_based'
                level = self.determine_content_level(text)
            
            # Method 4: Check numbering
            elif self.has_section_numbering(text):
                is_header = True
                header_type = 'numbered'
                level = self.get_numbering_level(text)
            
            if is_header:
                # Analyze formatting
                formatting = self.analyze_paragraph_formatting(para)
                
                header_info = {
                    'index': para_idx,
                    'text': text,
                    'level': level,
                    'type': header_type,
                    'style': para.style.name if para.style else 'Unknown',
                    'formatting': formatting,
                    'length': len(text),
                    'word_count': len(text.split())
                }
                
                headers_data['all_headers'].append(header_info)
                
                # Build hierarchy
                if level not in headers_data['header_hierarchy']:
                    headers_data['header_hierarchy'][level] = []
                headers_data['header_hierarchy'][level].append(header_info)
                
                # Track formatting patterns
                format_key = f"{formatting.get('font_name', 'Unknown')}_{formatting.get('font_size', 'Unknown')}_{formatting.get('bold', False)}"
                if format_key not in headers_data['formatting_patterns']:
                    headers_data['formatting_patterns'][format_key] = []
                headers_data['formatting_patterns'][format_key].append(header_info)
                
                # Track types
                if header_type not in headers_data['header_types']:
                    headers_data['header_types'][header_type] = []
                headers_data['header_types'][header_type].append(header_info)
        
        # Display results
        print(f"📚 SECTION HEADERS FOUND: {len(headers_data['all_headers'])}")
        
        print(f"\n📂 HEADER HIERARCHY:")
        for level in sorted(headers_data['header_hierarchy'].keys()):
            headers_at_level = headers_data['header_hierarchy'][level]
            print(f"   Level {level}: {len(headers_at_level)} headers")
            
            for i, header in enumerate(headers_at_level[:5]):
                indent = "  " * level
                print(f"   {indent}📂 {header['text'][:60]}...")
            
            if len(headers_at_level) > 5:
                print(f"   {indent}... and {len(headers_at_level) - 5} more")
        
        print(f"\n🎨 FORMATTING PATTERNS:")
        for pattern, headers in headers_data['formatting_patterns'].items():
            if len(headers) > 1:  # Only show patterns used multiple times
                font_parts = pattern.split('_')
                font_name = font_parts[0] if font_parts else 'Unknown'
                font_size = font_parts[1] if len(font_parts) > 1 else 'Unknown'
                is_bold = font_parts[2] == 'True' if len(font_parts) > 2 else False
                
                print(f"   🎨 {font_name} {font_size}pt {'Bold' if is_bold else 'Normal'}: {len(headers)} headers")
                
                # Show examples
                for header in headers[:2]:
                    print(f"      • {header['text'][:40]}...")
        
        print(f"\n📊 HEADER TYPES BREAKDOWN:")
        for header_type, headers in headers_data['header_types'].items():
            print(f"   {header_type.title()}: {len(headers)} headers")
        
        return headers_data
    
    def analyze_document_flow(self, doc):
        """Analyze the flow and structure of the document"""
        
        print(f"\n🌊 DOCUMENT FLOW ANALYSIS")
        print("-" * 60)
        
        flow_data = {
            'document_sections': [],
            'content_density': {},
            'table_distribution': [],
            'text_to_table_ratio': {}
        }
        
        current_section = {
            'start_index': 0,
            'title': 'Document Start',
            'paragraph_count': 0,
            'table_count': 0,
            'word_count': 0
        }
        
        table_indices = []
        for i, table in enumerate(doc.tables):
            # Find table position by searching for table content
            for para_idx, para in enumerate(doc.paragraphs):
                # This is a simplified approach - tables are typically between paragraphs
                pass
        
        section_count = 0
        
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            current_section['paragraph_count'] += 1
            if text:
                current_section['word_count'] += len(text.split())
            
            # Check if this starts a new section
            if self.is_section_boundary(para):
                # Save current section
                flow_data['document_sections'].append(current_section.copy())
                
                # Start new section
                section_count += 1
                current_section = {
                    'start_index': para_idx,
                    'title': text[:50] + '...' if len(text) > 50 else text,
                    'paragraph_count': 0,
                    'table_count': 0,
                    'word_count': 0,
                    'section_number': section_count
                }
        
        # Add final section
        if current_section['paragraph_count'] > 0:
            flow_data['document_sections'].append(current_section)
        
        # Analyze table distribution
        tables_per_section = len(doc.tables) // max(len(flow_data['document_sections']), 1)
        
        print(f"📊 DOCUMENT FLOW SUMMARY:")
        print(f"   Total Sections: {len(flow_data['document_sections'])}")
        print(f"   Average Section Length: {sum(s['paragraph_count'] for s in flow_data['document_sections']) / len(flow_data['document_sections']):.1f} paragraphs")
        print(f"   Total Tables: {len(doc.tables)}")
        print(f"   Average Tables per Section: {tables_per_section}")
        
        print(f"\n📚 SECTION BREAKDOWN:")
        for i, section in enumerate(flow_data['document_sections'][:10]):
            density = section['word_count'] / section['paragraph_count'] if section['paragraph_count'] > 0 else 0
            print(f"   Section {i+1}: \"{section.get('title', 'Unknown')}\"")
            print(f"      📊 {section['paragraph_count']} paragraphs, {section['word_count']} words")
            print(f"      📈 Density: {density:.1f} words/paragraph")
        
        if len(flow_data['document_sections']) > 10:
            print(f"   ... and {len(flow_data['document_sections']) - 10} more sections")
        
        return flow_data
    
    def build_content_hierarchy(self, doc):
        """Build complete content hierarchy"""
        
        print(f"\n🏗️ CONTENT HIERARCHY BUILDING")
        print("-" * 60)
        
        hierarchy = {
            'root': {
                'title': 'Document Root',
                'children': [],
                'level': 0,
                'content_summary': {}
            }
        }
        
        current_path = [hierarchy['root']]
        
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # Determine if this is a structural element
            if self.is_structural_element(para):
                level = self.get_structural_level(para)
                
                # Adjust path based on level
                while len(current_path) > level + 1:
                    current_path.pop()
                
                # Create new node
                new_node = {
                    'title': text,
                    'level': level,
                    'index': para_idx,
                    'children': [],
                    'content_summary': {
                        'paragraph_count': 0,
                        'table_count': 0,
                        'word_count': 0
                    }
                }
                
                # Add to parent
                parent = current_path[-1] if current_path else hierarchy['root']
                parent['children'].append(new_node)
                current_path.append(new_node)
            
            # Update content summary for current path
            for node in current_path:
                node['content_summary']['paragraph_count'] += 1
                node['content_summary']['word_count'] += len(text.split())
        
        # Display hierarchy
        print(f"📋 CONTENT HIERARCHY:")
        self.print_hierarchy_node(hierarchy['root'], 0)
        
        return hierarchy
    
    def print_hierarchy_node(self, node, indent_level):
        """Print hierarchy node with proper indentation"""
        indent = "  " * indent_level
        title = node.get('title', 'Unknown')
        summary = node.get('content_summary', {})
        
        if indent_level == 0:
            print(f"📚 {title}")
        else:
            print(f"{indent}📂 {title}")
        
        if summary:
            print(f"{indent}   📊 {summary.get('paragraph_count', 0)} paragraphs, {summary.get('word_count', 0)} words")
        
        # Print children (limit to prevent overflow)
        children = node.get('children', [])
        for i, child in enumerate(children[:5]):
            self.print_hierarchy_node(child, indent_level + 1)
        
        if len(children) > 5:
            print(f"{indent}  ... and {len(children) - 5} more subsections")
    
    # Helper methods for classification
    def is_cover_element(self, text):
        cover_indicators = ['judul', 'title', 'nama', 'author', 'tanggal', 'date', 'versi', 'version']
        return any(indicator in text.lower() for indicator in cover_indicators) and len(text) < 100
    
    def categorize_cover_element(self, text):
        if any(word in text.lower() for word in ['judul', 'title']):
            return 'title'
        elif any(word in text.lower() for word in ['nama', 'author', 'penulis']):
            return 'author'
        elif any(word in text.lower() for word in ['tanggal', 'date']):
            return 'date'
        elif any(word in text.lower() for word in ['versi', 'version']):
            return 'version'
        else:
            return 'other'
    
    def is_toc_element(self, text):
        toc_patterns = ['daftar isi', 'table of contents', 'contents']
        return any(pattern in text.lower() for pattern in toc_patterns) or \
               (len(text.split('.')) >= 2 and text.split('.')[0].isdigit())
    
    def get_toc_level(self, text):
        dots = text.count('.')
        return min(dots, 3)  # Cap at level 3
    
    def is_main_section_header(self, text):
        return len(text) < 150 and (text.isupper() or text.istitle()) and len(text.split()) <= 10
    
    def get_header_level(self, text, para):
        if para.style and 'heading' in para.style.name.lower():
            return int(para.style.name.lower().replace('heading', '').strip() or '1')
        return 1
    
    def is_appendix_element(self, text):
        return any(word in text.lower() for word in ['lampiran', 'appendix', 'attachment'])
    
    def determine_toc_level(self, text):
        import re
        if re.match(r'^\d+\.', text):
            return 1
        elif re.match(r'^\d+\.\d+', text):
            return 2
        elif re.match(r'^\d+\.\d+\.\d+', text):
            return 3
        return 1
    
    def extract_heading_level_from_style(self, style_name):
        import re
        match = re.search(r'(\d+)', style_name)
        return int(match.group(1)) if match else 1
    
    def is_formatted_header(self, para):
        if not para.runs:
            return False
        first_run = para.runs[0]
        return first_run.font.bold or (first_run.font.size and first_run.font.size.pt > 12)
    
    def determine_format_level(self, para):
        if not para.runs:
            return 1
        first_run = para.runs[0]
        if first_run.font.size:
            size = first_run.font.size.pt
            if size >= 18:
                return 1
            elif size >= 14:
                return 2
            else:
                return 3
        return 2
    
    def contains_header_indicators(self, text, indicators):
        return any(indicator in text.lower() for indicator in indicators)
    
    def determine_content_level(self, text):
        if any(word in text.lower() for word in ['bab', 'chapter']):
            return 1
        elif any(word in text.lower() for word in ['sub', 'bagian']):
            return 2
        return 3
    
    def has_section_numbering(self, text):
        import re
        return bool(re.match(r'^(\d+\.|\d+\.\d+)', text))
    
    def get_numbering_level(self, text):
        import re
        if re.match(r'^\d+\.', text):
            return 1
        elif re.match(r'^\d+\.\d+', text):
            return 2
        elif re.match(r'^\d+\.\d+\.\d+', text):
            return 3
        return 1
    
    def analyze_paragraph_formatting(self, para):
        if not para.runs:
            return {'font_name': 'Unknown', 'font_size': 'Unknown', 'bold': False}
        
        first_run = para.runs[0]
        return {
            'font_name': first_run.font.name or 'Unknown',
            'font_size': first_run.font.size.pt if first_run.font.size else 'Unknown',
            'bold': first_run.font.bold or False,
            'italic': first_run.font.italic or False
        }
    
    def is_section_boundary(self, para):
        text = para.text.strip()
        return (para.style and 'heading' in para.style.name.lower()) or \
               (len(text) < 100 and text.isupper()) or \
               self.has_section_numbering(text)
    
    def is_structural_element(self, para):
        return self.is_section_boundary(para)
    
    def get_structural_level(self, para):
        return self.get_header_level(para.text, para)

def main():
    """Main function untuk membaca struktur template"""
    
    print("📖 Template Structure Reader AI")
    print("🔍 Reading table of contents and document structure")
    print("📋 Analyzing all structural elements...")
    print()
    
    reader = TemplateStructureReader()
    result = reader.read_template_structure()
    
    if result:
        print(f"\n✨ TEMPLATE STRUCTURE READING COMPLETE!")
        print(f"📚 All structural elements have been analyzed!")
        print(f"📋 Daftar isi and document outline extracted successfully!")
    else:
        print(f"\n❌ Template structure reading failed!")

if __name__ == '__main__':
    main()