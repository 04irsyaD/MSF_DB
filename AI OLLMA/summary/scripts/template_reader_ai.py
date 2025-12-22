# template_reader_ai.py - AI khusus membaca dan menganalisis template 
import os
from datetime import datetime

class TemplateReaderAI:
    """AI khusus untuk membaca dan menganalisis template secara detail"""
    
    def __init__(self):
        self.template_path = os.path.join('template', 'template_dokumentasi.docx')
        
        print(f"🤖 Template Reader AI v2.0 Initialized")
        print(f"📚 Specialized in comprehensive template analysis")
        print(f"🔍 Will extract ALL template details and log findings")
    
    def read_and_analyze_template(self):
        """Baca dan analisis template dengan detail comprehensive"""
        
        print("\n" + "="*80)
        print("🔍 TEMPLATE READER AI - COMPREHENSIVE ANALYSIS")
        print("="*80)
        print(f"📁 Template File: {self.template_path}")
        print(f"⏰ Analysis Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print("="*80)
        
        if not os.path.exists(self.template_path):
            print(f"❌ Template not found: {self.template_path}")
            return None
        
        try:
            from docx import Document
            
            # Load template
            doc = Document(self.template_path)
            
            # Get file size
            file_size = os.path.getsize(self.template_path)
            print(f"📊 Template Size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
            
            # Comprehensive analysis
            results = {
                'basic_info': self.analyze_basic_structure(doc),
                'content_details': self.analyze_content_in_detail(doc),
                'formatting_patterns': self.analyze_formatting_patterns(doc),
                'table_structures': self.analyze_table_structures(doc),
                'style_definitions': self.analyze_style_definitions(doc),
                'template_elements': self.identify_template_elements(doc)
            }
            
            # Generate insights
            self.generate_template_insights(results)
            
            return results
            
        except Exception as e:
            print(f"❌ Analysis error: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def analyze_basic_structure(self, doc):
        """Analisis struktur dasar dokumen"""
        
        print("\n📋 BASIC DOCUMENT STRUCTURE")
        print("-" * 50)
        
        basic_info = {
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'total_sections': len(doc.sections),
            'content_paragraphs': 0,
            'empty_paragraphs': 0
        }
        
        # Count content vs empty paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                basic_info['content_paragraphs'] += 1
            else:
                basic_info['empty_paragraphs'] += 1
        
        print(f"📊 Total Paragraphs: {basic_info['total_paragraphs']:,}")
        print(f"📝 Content Paragraphs: {basic_info['content_paragraphs']:,}")
        print(f"⭕ Empty Paragraphs: {basic_info['empty_paragraphs']:,}")
        print(f"📋 Total Tables: {basic_info['total_tables']:,}")
        print(f"📄 Total Sections: {basic_info['total_sections']:,}")
        
        return basic_info
    
    def analyze_content_in_detail(self, doc):
        """Analisis konten secara detail"""
        
        print("\n📖 DETAILED CONTENT ANALYSIS")
        print("-" * 50)
        
        content_analysis = {
            'total_words': 0,
            'total_characters': 0,
            'content_categories': {
                'titles': [],
                'headers': [],
                'body_paragraphs': [],
                'table_references': [],
                'instructions': [],
                'placeholders': []
            },
            'language_patterns': {},
            'content_samples': {}
        }
        
        # Analyze each paragraph
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
                
            # Count words and characters
            words = len(text.split())
            chars = len(text)
            content_analysis['total_words'] += words
            content_analysis['total_characters'] += chars
            
            # Categorize content
            self.categorize_paragraph_content(i, text, content_analysis['content_categories'])
        
        print(f"📊 Total Words: {content_analysis['total_words']:,}")
        print(f"📝 Total Characters: {content_analysis['total_characters']:,}")
        print()
        
        # Show category breakdown
        print("📂 CONTENT CATEGORIES:")
        for category, items in content_analysis['content_categories'].items():
            print(f"   {category.title()}: {len(items)} items")
            if items and len(items) > 0:
                # Show first few samples
                for j, item in enumerate(items[:2]):
                    sample_text = item['text'][:60] + '...' if len(item['text']) > 60 else item['text']
                    print(f"      Sample {j+1}: \"{sample_text}\"")
        
        return content_analysis
    
    def categorize_paragraph_content(self, index, text, categories):
        """Kategorikan konten paragraf"""
        
        text_lower = text.lower()
        
        # Check for placeholders or template markers
        if any(marker in text for marker in ['{', '}', '[', ']', '{{', '}}', 'XXXX', '___']):
            categories['placeholders'].append({
                'index': index,
                'text': text,
                'type': 'placeholder'
            })
        
        # Check for table references
        elif any(word in text_lower for word in ['table', 'tabel', 'figure', 'gambar', 'chart']):
            categories['table_references'].append({
                'index': index,
                'text': text,
                'type': 'table_reference'
            })
        
        # Check for instructions or commands
        elif any(word in text_lower for word in ['isi', 'fill', 'enter', 'input', 'masukkan', 'tuliskan']):
            categories['instructions'].append({
                'index': index,
                'text': text,
                'type': 'instruction'
            })
        
        # Check for titles (short, uppercase pattern)
        elif len(text) < 60 and (text.isupper() or text.istitle()) and len(text.split()) <= 8:
            categories['titles'].append({
                'index': index,
                'text': text,
                'type': 'title'
            })
        
        # Check for headers (medium length, formal)
        elif len(text) < 150 and any(char.isupper() for char in text) and text.endswith((':',)):
            categories['headers'].append({
                'index': index,
                'text': text,
                'type': 'header'
            })
        
        # Everything else is body text
        else:
            categories['body_paragraphs'].append({
                'index': index,
                'text': text,
                'type': 'body'
            })
    
    def analyze_formatting_patterns(self, doc):
        """Analisis pola formatting secara detail"""
        
        print("\n🎨 DETAILED FORMATTING ANALYSIS")
        print("-" * 50)
        
        formatting_data = {
            'font_families': {},
            'font_sizes': {},
            'text_styles': {},
            'paragraph_styles': {},
            'color_usage': {},
            'alignment_patterns': {}
        }
        
        # Analyze each paragraph's formatting
        for para_idx, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
                
            # Paragraph level formatting
            para_style = para.style.name if para.style else 'Normal'
            if para_style not in formatting_data['paragraph_styles']:
                formatting_data['paragraph_styles'][para_style] = 0
            formatting_data['paragraph_styles'][para_style] += 1
            
            # Alignment
            alignment = str(para.alignment) if para.alignment else 'Left'
            if alignment not in formatting_data['alignment_patterns']:
                formatting_data['alignment_patterns'][alignment] = 0
            formatting_data['alignment_patterns'][alignment] += 1
            
            # Run level formatting
            for run in para.runs:
                if not run.text.strip():
                    continue
                
                # Font family
                font_name = run.font.name or 'Default'
                if font_name not in formatting_data['font_families']:
                    formatting_data['font_families'][font_name] = 0
                formatting_data['font_families'][font_name] += 1
                
                # Font size
                font_size = str(run.font.size.pt) + 'pt' if run.font.size else 'Default'
                if font_size not in formatting_data['font_sizes']:
                    formatting_data['font_sizes'][font_size] = 0
                formatting_data['font_sizes'][font_size] += 1
                
                # Text styles
                styles = []
                if run.font.bold:
                    styles.append('Bold')
                if run.font.italic:
                    styles.append('Italic')
                if run.font.underline:
                    styles.append('Underline')
                
                style_combo = '+'.join(styles) if styles else 'Normal'
                if style_combo not in formatting_data['text_styles']:
                    formatting_data['text_styles'][style_combo] = 0
                formatting_data['text_styles'][style_combo] += 1
        
        # Display results
        print("🔤 FONT FAMILIES USED:")
        for font, count in sorted(formatting_data['font_families'].items(), key=lambda x: x[1], reverse=True):
            percentage = (count / sum(formatting_data['font_families'].values())) * 100
            print(f"   {font}: {count} uses ({percentage:.1f}%)")
        
        print("\n📏 FONT SIZES USED:")
        for size, count in sorted(formatting_data['font_sizes'].items(), key=lambda x: x[1], reverse=True):
            percentage = (count / sum(formatting_data['font_sizes'].values())) * 100
            print(f"   {size}: {count} uses ({percentage:.1f}%)")
        
        print("\n🎨 TEXT STYLES USED:")
        for style, count in sorted(formatting_data['text_styles'].items(), key=lambda x: x[1], reverse=True):
            percentage = (count / sum(formatting_data['text_styles'].values())) * 100
            print(f"   {style}: {count} uses ({percentage:.1f}%)")
        
        print("\n📄 PARAGRAPH STYLES USED:")
        for style, count in sorted(formatting_data['paragraph_styles'].items(), key=lambda x: x[1], reverse=True):
            percentage = (count / sum(formatting_data['paragraph_styles'].values())) * 100
            print(f"   {style}: {count} uses ({percentage:.1f}%)")
        
        return formatting_data
    
    def analyze_table_structures(self, doc):
        """Analisis struktur table secara detail"""
        
        print(f"\n📊 DETAILED TABLE ANALYSIS")
        print("-" * 50)
        
        if not doc.tables:
            print("❌ No tables found in template")
            return {'total_tables': 0, 'tables': []}
        
        table_analysis = {
            'total_tables': len(doc.tables),
            'tables': []
        }
        
        print(f"📋 Found {len(doc.tables)} table(s)")
        print()
        
        for table_idx, table in enumerate(doc.tables):
            table_info = {
                'index': table_idx,
                'dimensions': {
                    'rows': len(table.rows),
                    'columns': len(table.columns) if table.rows else 0
                },
                'content_analysis': {},
                'formatting_analysis': {},
                'structure_analysis': {}
            }
            
            print(f"📋 TABLE {table_idx + 1} ANALYSIS:")
            print(f"   📊 Dimensions: {table_info['dimensions']['rows']} rows × {table_info['dimensions']['columns']} columns")
            
            # Content analysis
            total_cells = 0
            filled_cells = 0
            cell_contents = []
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    total_cells += 1
                    cell_text = cell.text.strip()
                    if cell_text:
                        filled_cells += 1
                        cell_contents.append({
                            'row': row_idx,
                            'col': cell_idx,
                            'text': cell_text
                        })
            
            table_info['content_analysis'] = {
                'total_cells': total_cells,
                'filled_cells': filled_cells,
                'fill_percentage': (filled_cells / total_cells * 100) if total_cells > 0 else 0,
                'cell_contents': cell_contents
            }
            
            print(f"   📝 Content: {filled_cells}/{total_cells} cells filled ({table_info['content_analysis']['fill_percentage']:.1f}%)")
            
            # Show sample content
            if cell_contents:
                print(f"   📖 Sample Content:")
                for i, content in enumerate(cell_contents[:6]):  # First 6 cells
                    sample = content['text'][:40] + '...' if len(content['text']) > 40 else content['text']
                    print(f"      Row {content['row']+1}, Col {content['col']+1}: \"{sample}\"")
                    if i >= 5:
                        break
            
            # Formatting analysis for header row
            if table.rows:
                header_row = table.rows[0]
                if header_row.cells and header_row.cells[0].paragraphs:
                    first_para = header_row.cells[0].paragraphs[0]
                    if first_para.runs:
                        first_run = first_para.runs[0]
                        
                        header_format = {
                            'font_name': first_run.font.name or 'Unknown',
                            'font_size': first_run.font.size.pt if first_run.font.size else 'Unknown',
                            'bold': first_run.font.bold,
                            'italic': first_run.font.italic,
                            'underline': bool(first_run.font.underline)
                        }
                        
                        table_info['formatting_analysis']['header'] = header_format
                        print(f"   🎨 Header Format: {header_format['font_name']} {header_format['font_size']}pt" + 
                              (f" Bold" if header_format['bold'] else "") +
                              (f" Italic" if header_format['italic'] else "") +
                              (f" Underline" if header_format['underline'] else ""))
            
            # Detect if it's a data table template
            if filled_cells > 0:
                header_indicators = ['nama', 'name', 'id', 'column', 'kolom', 'type', 'tipe', 'description', 'deskripsi']
                has_headers = any(indicator in cell_text.lower() 
                                for cell_content in cell_contents 
                                for indicator in header_indicators
                                for cell_text in [cell_content['text']])
                
                table_info['structure_analysis']['likely_data_table'] = has_headers
                table_info['structure_analysis']['template_type'] = 'data_template' if has_headers else 'content_template'
                
                print(f"   🔍 Template Type: {table_info['structure_analysis']['template_type']}")
                print(f"   📊 Data Table: {'Yes' if has_headers else 'No'}")
            
            table_analysis['tables'].append(table_info)
            print()
        
        return table_analysis
    
    def analyze_style_definitions(self, doc):
        """Analisis definisi style dokumen"""
        
        print("\n🎭 STYLE DEFINITIONS ANALYSIS")
        print("-" * 50)
        
        styles_info = {
            'total_styles': 0,
            'style_details': {},
            'custom_styles': [],
            'builtin_styles': []
        }
        
        try:
            if hasattr(doc, 'styles'):
                styles_info['total_styles'] = len(doc.styles)
                print(f"📋 Total Styles Available: {styles_info['total_styles']}")
                print()
                
                print("📝 STYLE INVENTORY:")
                for style in doc.styles:
                    style_name = style.name
                    style_type = style.type.name if hasattr(style, 'type') else 'Unknown'
                    is_builtin = getattr(style, 'builtin', False)
                    
                    style_detail = {
                        'name': style_name,
                        'type': style_type,
                        'builtin': is_builtin
                    }
                    
                    if is_builtin:
                        styles_info['builtin_styles'].append(style_detail)
                    else:
                        styles_info['custom_styles'].append(style_detail)
                    
                    styles_info['style_details'][style_name] = style_detail
                    
                    status = "Built-in" if is_builtin else "Custom"
                    print(f"   • {style_name} ({style_type}) - {status}")
                
                print(f"\n📊 Style Summary:")
                print(f"   Built-in Styles: {len(styles_info['builtin_styles'])}")
                print(f"   Custom Styles: {len(styles_info['custom_styles'])}")
                
        except Exception as e:
            print(f"⚠️ Style analysis error: {e}")
        
        return styles_info
    
    def identify_template_elements(self, doc):
        """Identifikasi elemen-elemen template khusus"""
        
        print(f"\n🔍 TEMPLATE ELEMENTS IDENTIFICATION")
        print("-" * 50)
        
        template_elements = {
            'placeholders': [],
            'template_markers': [],
            'instruction_text': [],
            'sample_content': [],
            'formatting_examples': []
        }
        
        # Scan for template-specific patterns
        placeholder_patterns = [
            r'\{.*?\}',  # {placeholder}
            r'\[.*?\]',  # [placeholder]
            r'___+',     # ___blank___
            r'XXXX+',    # XXXX
            r'\$\{.*?\}', # ${variable}
        ]
        
        instruction_keywords = [
            'isi', 'fill', 'enter', 'input', 'masukkan', 'tuliskan',
            'ganti', 'replace', 'ubah', 'change', 'sesuaikan'
        ]
        
        sample_indicators = [
            'contoh', 'example', 'sample', 'misal', 'seperti'
        ]
        
        print("🔍 Scanning for template elements...")
        
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            text_lower = text.lower()
            
            # Check for placeholders
            import re
            for pattern in placeholder_patterns:
                matches = re.findall(pattern, text)
                if matches:
                    template_elements['placeholders'].extend([{
                        'paragraph': para_idx,
                        'text': text,
                        'matches': matches,
                        'pattern': pattern
                    }])
            
            # Check for instruction text
            if any(keyword in text_lower for keyword in instruction_keywords):
                template_elements['instruction_text'].append({
                    'paragraph': para_idx,
                    'text': text,
                    'type': 'instruction'
                })
            
            # Check for sample content indicators
            if any(indicator in text_lower for indicator in sample_indicators):
                template_elements['sample_content'].append({
                    'paragraph': para_idx,
                    'text': text,
                    'type': 'sample'
                })
        
        # Display findings
        print(f"📋 TEMPLATE ELEMENTS FOUND:")
        print(f"   🔹 Placeholders: {len(template_elements['placeholders'])}")
        print(f"   📝 Instructions: {len(template_elements['instruction_text'])}")
        print(f"   📖 Sample Content: {len(template_elements['sample_content'])}")
        
        # Show examples
        if template_elements['placeholders']:
            print(f"\n🔹 PLACEHOLDER EXAMPLES:")
            for i, placeholder in enumerate(template_elements['placeholders'][:3]):
                print(f"   {i+1}. \"{placeholder['text'][:60]}...\"")
                print(f"      Matches: {placeholder['matches']}")
        
        if template_elements['instruction_text']:
            print(f"\n📝 INSTRUCTION EXAMPLES:")
            for i, instruction in enumerate(template_elements['instruction_text'][:3]):
                sample = instruction['text'][:60] + '...' if len(instruction['text']) > 60 else instruction['text']
                print(f"   {i+1}. \"{sample}\"")
        
        return template_elements
    
    def generate_template_insights(self, analysis_results):
        """Generate comprehensive insights dari analysis"""
        
        print(f"\n" + "="*80)
        print("💡 COMPREHENSIVE TEMPLATE INSIGHTS")
        print("="*80)
        
        basic_info = analysis_results.get('basic_info', {})
        content_details = analysis_results.get('content_details', {})
        formatting_patterns = analysis_results.get('formatting_patterns', {})
        table_structures = analysis_results.get('table_structures', {})
        template_elements = analysis_results.get('template_elements', {})
        
        print(f"🎯 TEMPLATE CHARACTERISTICS:")
        print(f"   📊 Document Scale: {content_details.get('total_words', 0):,} words in {basic_info.get('content_paragraphs', 0)} paragraphs")
        print(f"   📋 Table Structure: {table_structures.get('total_tables', 0)} tables detected")
        print(f"   📝 Content Density: {basic_info.get('content_paragraphs', 0) / basic_info.get('total_paragraphs', 1):.1%} filled paragraphs")
        
        # Primary formatting
        font_families = formatting_patterns.get('font_families', {})
        font_sizes = formatting_patterns.get('font_sizes', {})
        
        if font_families:
            primary_font = max(font_families.keys(), key=lambda k: font_families[k])
            print(f"   🔤 Primary Font: {primary_font}")
        
        if font_sizes:
            primary_size = max(font_sizes.keys(), key=lambda k: font_sizes[k])
            print(f"   📏 Primary Size: {primary_size}")
        
        # Template nature
        placeholders = len(template_elements.get('placeholders', []))
        instructions = len(template_elements.get('instruction_text', []))
        
        print(f"\n📋 TEMPLATE NATURE:")
        print(f"   🔹 Placeholder Elements: {placeholders}")
        print(f"   📝 Instruction Elements: {instructions}")
        
        if placeholders > instructions:
            template_type = "Data Template (Form-based)"
        elif instructions > placeholders:
            template_type = "Instruction Template (Guide-based)"
        else:
            template_type = "Mixed Template"
        
        print(f"   🎭 Template Type: {template_type}")
        
        # Table insights
        if table_structures.get('total_tables', 0) > 0:
            print(f"\n📊 TABLE INSIGHTS:")
            for i, table_info in enumerate(table_structures.get('tables', [])):
                structure = table_info.get('structure_analysis', {})
                content = table_info.get('content_analysis', {})
                
                print(f"   Table {i+1}: {structure.get('template_type', 'Unknown')} " +
                      f"({content.get('fill_percentage', 0):.1f}% filled)")
        
        print(f"\n🎨 FORMATTING STRATEGY:")
        print(f"   ✅ Preserve detected font families")
        print(f"   ✅ Maintain size hierarchy")
        print(f"   ✅ Keep style patterns")
        print(f"   ✅ Replicate table structures")
        
        print(f"\n🚀 GENERATION RECOMMENDATIONS:")
        print(f"   1. Extract formatting → Apply to clean document")
        print(f"   2. Remove template content → Keep structure")
        print(f"   3. Fill with database data → Preserve format")
        print(f"   4. Validate output → Ensure consistency")
        
        print(f"\n✨ Template analysis complete! Use these insights for perfect generation.")

def main():
    """Main function untuk template reading"""
    
    print("🤖 Template Reader AI v2.0")
    print("📚 Comprehensive template analysis and logging")
    print("🔍 Extracting ALL template characteristics...")
    print()
    
    reader = TemplateReaderAI()
    result = reader.read_and_analyze_template()
    
    if result:
        print(f"\n🎉 TEMPLATE READING COMPLETE!")
        print(f"📊 All template details have been analyzed and logged")
        print(f"💡 Use these insights for perfect document generation!")
    else:
        print(f"\n❌ Template reading failed!")

if __name__ == '__main__':
    main()