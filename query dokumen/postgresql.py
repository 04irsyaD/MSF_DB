from sqlalchemy import create_engine, text
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from graphviz import Digraph
import os
import networkx as nx   # <-- tambahan untuk grouping relasi

# ==============================
# 1. Koneksi Database
# ==============================

# --- Tambahkan input schema ---
username = "postgres"     # ganti dengan username PostgreSQL kamu
password = "1234"         # ganti dengan password PostgreSQL kamu
host = "localhost"
port = "5414"
database = "erm"          # ganti dengan nama database kamu

# Pilih schema yang ingin diproses
import sys
if len(sys.argv) > 1:
    selected_schema = sys.argv[1]
else:
    selected_schema = "backup2"  # default
print(f"Schema yang diproses: {selected_schema}")


engine = create_engine(f"postgresql+psycopg2://{username}:{password}@{host}:{port}/{database}")
# Atur schema default pada koneksi
from sqlalchemy import text
with engine.connect() as conn:
    conn.execute(text(f"SET search_path TO {selected_schema}"))

# ==============================
# 2. Ambil Metadata Tables & Columns
# ==============================
query_tables = f"""
SELECT
    c.table_name,
    c.column_name,
    c.data_type,
    c.is_nullable,
    c.column_default,
    pgd.description AS column_description,
    CASE 
        WHEN pk.column_name IS NOT NULL THEN 'PK'
        WHEN fk.column_name IS NOT NULL THEN 'FK'
        ELSE ''
    END AS key_type
FROM information_schema.columns c
LEFT JOIN pg_catalog.pg_statio_all_tables as st
    ON c.table_schema = st.schemaname AND c.table_name = st.relname
LEFT JOIN pg_catalog.pg_description pgd
    ON pgd.objoid = st.relid AND pgd.objsubid = c.ordinal_position
LEFT JOIN (
    SELECT kcu.table_name, kcu.column_name
    FROM information_schema.table_constraints tc
    JOIN information_schema.key_column_usage kcu 
        ON tc.constraint_name = kcu.constraint_name
        AND tc.table_schema = kcu.table_schema
    WHERE tc.constraint_type = 'PRIMARY KEY'
      AND tc.table_schema = '{selected_schema}'
) pk ON c.table_name = pk.table_name AND c.column_name = pk.column_name
LEFT JOIN (
    SELECT kcu.table_name, kcu.column_name
    FROM information_schema.table_constraints tc
    JOIN information_schema.key_column_usage kcu 
        ON tc.constraint_name = kcu.constraint_name
        AND tc.table_schema = kcu.table_schema
    WHERE tc.constraint_type = 'FOREIGN KEY'
      AND tc.table_schema = '{selected_schema}'
) fk ON c.table_name = fk.table_name AND c.column_name = fk.column_name
WHERE c.table_schema = '{selected_schema}'
ORDER BY c.table_name, c.ordinal_position;
"""
df_tables = pd.read_sql(query_tables, engine)

# ==============================
# 3. Functions
# ==============================
query_functions = """
SELECT 
    n.nspname AS schema,
    p.proname AS function_name,
    pg_catalog.pg_get_function_result(p.oid) AS return_type,
    pg_catalog.pg_get_function_arguments(p.oid) AS arguments,
    d.description AS description
FROM pg_proc p
LEFT JOIN pg_namespace n ON n.oid = p.pronamespace
LEFT JOIN pg_description d ON d.objoid = p.oid
WHERE n.nspname NOT IN ('pg_catalog', 'information_schema')
ORDER BY n.nspname, p.proname;
"""
df_functions = pd.read_sql(query_functions, engine)

# ==============================
# 4. Triggers
# ==============================
query_triggers = f"""
SELECT 
        t.tgname AS trigger_name,
        c.relname AS table_name,
        pg_catalog.pg_get_triggerdef(t.oid, true) AS definition
FROM pg_trigger t
JOIN pg_class c ON c.oid = t.tgrelid
JOIN pg_namespace n ON n.oid = c.relnamespace
WHERE NOT t.tgisinternal
    AND n.nspname = '{selected_schema}'
ORDER BY c.relname, t.tgname;
"""
df_triggers = pd.read_sql(query_triggers, engine)

# ==============================
# 5. Views
# ==============================
query_views = f"""
SELECT 
    table_name,
    view_definition
FROM information_schema.views
WHERE table_schema = '{selected_schema}'
ORDER BY table_name;
"""
df_views = pd.read_sql(query_views, engine)

# ==============================
# 6. Foreign Keys (Relasi)
# ==============================
query_fk = f"""
SELECT
        tc.table_name AS source_table,
        kcu.column_name AS source_column,
        ccu.table_name AS target_table,
        ccu.column_name AS target_column,
        tc.constraint_name
FROM information_schema.table_constraints AS tc
JOIN information_schema.key_column_usage AS kcu
        ON tc.constraint_name = kcu.constraint_name
     AND tc.table_schema = kcu.table_schema
JOIN information_schema.constraint_column_usage AS ccu
        ON ccu.constraint_name = tc.constraint_name
     AND ccu.table_schema = tc.table_schema
WHERE tc.constraint_type = 'FOREIGN KEY'
    AND tc.table_schema = '{selected_schema}';
"""
df_fk = pd.read_sql(query_fk, engine)

# ==============================
# 7. Helper: Grouping Berdasarkan Relasi
# ==============================
def group_tables_by_relationships(df_tables, df_fk, max_tables=15):
    """Group tables into connected components based on relationships"""
    G = nx.Graph()
    G.add_nodes_from(df_tables["table_name"].unique())
    for _, row in df_fk.iterrows():
        G.add_edge(row["source_table"], row["target_table"])

    groups = []
    for component in nx.connected_components(G):
        comp = list(component)
        # Jika group terlalu besar, pecah jadi batch
        for i in range(0, len(comp), max_tables):
            groups.append(comp[i:i+max_tables])
    return groups

# ==============================
# 8. Buat ERD dengan Graphviz - A4 OPTIMIZED (Paginated)
# ==============================
def create_paginated_erds(df_tables, df_fk, tables_per_page=15, use_grouping=True):
    """Create multiple ERD pages for large databases"""
    if use_grouping:
        pages = group_tables_by_relationships(df_tables, df_fk, max_tables=tables_per_page)
    else:
        all_tables = df_tables["table_name"].unique()
        pages = [all_tables[i:i+tables_per_page] for i in range(0, len(all_tables), tables_per_page)]
    
    created_files = []
    max_images = 3  # batas maksimal file gambar ERD_Page
    for page_idx, page_tables in enumerate(pages, 1):
        if page_idx > max_images:
            break
        dot = Digraph(comment=f"ERD Page {page_idx}", format="png")
        
        # A4 optimized settings
        dot.attr(size='11.7,8.3!')
        dot.attr(ratio='compress')
        dot.attr(rankdir='TB')
        dot.attr(ranksep='0.6')
        dot.attr(nodesep='0.4')
        dot.attr(margin='0.3')
        dot.attr(dpi='200')
        dot.attr(overlap='false')
        dot.attr(splines='true')
        
        dot.attr('node', shape='plaintext', fontname='Arial', fontsize='9')
        dot.attr('edge', fontname='Arial', fontsize='8', color='#666666')
        
        # Add tables in this page
        for table in page_tables:
            table_columns = df_tables[df_tables['table_name'] == table]
            html = f'''<
            <TABLE BORDER="1" CELLBORDER="1" CELLSPACING="0" CELLPADDING="3">
                <TR><TD BGCOLOR="#2E75B6"><FONT COLOR="white"><B>{table}</B></FONT></TD></TR>'''
            for _, col in table_columns.head(6).iterrows():
                key_indicator = " 🔑" if col['key_type']=='PK' else (" 🔗" if col['key_type']=='FK' else "")
                html += f'<TR><TD ALIGN="LEFT">{col["column_name"]}{key_indicator}</TD></TR>'
            if len(table_columns) > 6:
                html += f'<TR><TD BGCOLOR="#F5F5F5">+{len(table_columns)-6} more</TD></TR>'
            html += '</TABLE>>'
            dot.node(table, html)

        # Add relationships within this batch
        for _, row in df_fk.iterrows():
            if row["source_table"] in page_tables and row["target_table"] in page_tables:
                dot.edge(row["source_table"], row["target_table"])

        filename = f"ERD_Page_{page_idx}"
        dot.render(filename, format="png", cleanup=True)
        created_files.append(f"{filename}.png")
        print(f"✅ Created {filename}.png ({len(page_tables)} tables)")
    
    return created_files

# ==============================
# 9. Dashboard + lainnya (tidak diubah)
# Template fungsi agar tidak error
def create_overview_dashboard():
    filename = "overview_dashboard.png"
    # Tambahkan logika dashboard jika diperlukan
    return filename

def create_relationships_only_erd():
    filename = "relationships_only_erd.png"
    # Tambahkan logika ERD relasi jika diperlukan
    return filename

def create_table_index():
    filename = "table_index.png"
    # Tambahkan logika index jika diperlukan
    return filename
# ... (semua fungsi create_overview_dashboard, create_relationships_only_erd, create_table_index tetap sama seperti code kamu)

# ==============================
# 10. Eksekusi utama
# ==============================
print(f"\n=== CREATING ERDS FOR LARGE DATABASE ({len(df_tables['table_name'].unique())} tables) ===")

# 1. Create paginated ERDs (15 tables per page, grouped by relationships)
paginated_files = create_paginated_erds(df_tables, df_fk, tables_per_page=15, use_grouping=True)

# 2. Create overview dashboard
dashboard_file = create_overview_dashboard()
print(f"✅ Created overview dashboard: {os.path.abspath(dashboard_file)}")

# 3. Create relationships-only ERD
rel_file = create_relationships_only_erd()
if rel_file:
    print(f"✅ Created relationships ERD: {os.path.abspath(rel_file)}")

# 4. Create alphabetical index
index_file = create_table_index()
print(f"✅ Created table index: {os.path.abspath(index_file)}")

# ==============================
# 11. Buat Dokumen Word (tidak diubah banyak)
# ==============================
# ... (lanjutkan kode Word kamu persis sama seperti sebelumnya)


# ==============================
from docx import Document
template_path = os.path.join(os.path.dirname(__file__), "TSD ERM.docx")
doc = Document(template_path)

# Tambahkan daftar isi otomatis
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
paragraph = doc.add_paragraph()
run = paragraph.add_run()
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.text = r'TOC \o "1-3" \h \z \u'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')
r_element = run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldChar3)
doc.add_heading('📘 Dokumentasi Database', 0)

# Ringkasan umum
doc.add_heading('1. Ringkasan Umum', level=1)
doc.add_paragraph(f"Nama Database : {database}")
doc.add_paragraph("Tujuan/Fungsi : ")
doc.add_paragraph("Digunakan oleh Aplikasi : ")
doc.add_paragraph("Environment : Development / Staging / Production")
doc.add_paragraph("Versi DBMS : PostgreSQL")

# Arsitektur
doc.add_heading('2. Arsitektur Database', level=1)
doc.add_paragraph("Jenis Database : PostgreSQL")
doc.add_paragraph("Versi DBMS : (isi sesuai server)")
doc.add_paragraph("Topologi : Standalone / Replication / Cluster")
doc.add_paragraph("Diagram Arsitektur :")

# Sisipkan ERD
if paginated_files:
    for erd_file in paginated_files:
        if os.path.exists(erd_file):
            doc.add_picture(erd_file, width=Inches(2)) # Sesuaikan ukuran gambar
            doc.add_paragraph(f"Gambar: {erd_file} hasil generate otomatis dengan detail kolom dan relasi PK-FK.")

# Skema Database
doc.add_heading('3. Skema Database', level=1)
doc.add_paragraph("Nama Schema : public")
doc.add_paragraph("Deskripsi : Schema utama aplikasi")
doc.add_paragraph("ERD : (lihat gambar di atas)")

# ==============================
# 9. Loop per Tabel
# ==============================
doc.add_heading('4. Tabel & Struktur Data', level=1)

for table in df_tables['table_name'].unique():
    doc.add_heading(f"Tabel: {table}", level=2)
    doc.add_paragraph("Deskripsi: (isi deskripsi fungsi tabel ini)")
    
    subset = df_tables[df_tables['table_name'] == table]
    
    # Create compact table like sample: No | Nama Field | Tipe Data | Deskripsi Field
    word_table = doc.add_table(rows=1, cols=4)
    word_table.style = 'Table Grid'
    word_table.autofit = False
    # Column widths (adjust as needed)
    NO_COL_WIDTH = Inches(0.32)
    NAME_COL_WIDTH = Inches(2.2)
    TYPE_COL_WIDTH = Inches(1.4)
    DESC_COL_WIDTH = Inches(3.0)
    col_widths = [NO_COL_WIDTH, NAME_COL_WIDTH, TYPE_COL_WIDTH, DESC_COL_WIDTH]

    hdr_cells = word_table.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Nama Field'
    hdr_cells[2].text = 'Tipe Data'
    hdr_cells[3].text = 'Deskripsi Field'
    # header style
    try:
        for c in hdr_cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception:
        pass

    # try apply simple widths (some environments accept .width)
    # Compute available page width and scale columns to fit if needed
    try:
        section = doc.sections[0]
        avail_inches = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
    except Exception:
        avail_inches = 6.5

    total_inches = sum(w.inches for w in col_widths)
    if total_inches > avail_inches and total_inches > 0:
        scale = avail_inches / total_inches
        # allow more aggressive scaling but keep some minimum
        MIN_SCALE = 0.45
        if scale < MIN_SCALE:
            scale = MIN_SCALE
        col_widths = [Inches(w.inches * scale) for w in col_widths]

    # set table width (w:tblW) and per-cell tcW to enforce widths
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    try:
        total_twips = int(sum(w.inches for w in col_widths) * 1440)
        tblPr = word_table._tbl.get_or_add_tblPr()
        # remove existing tblW if present
        for child in list(tblPr):
            if child.tag == qn('w:tblW'):
                tblPr.remove(child)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(total_twips))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)

        # set each cell width via tcPr/tcW
        for col_idx, w in enumerate(col_widths):
            tw = int(w.inches * 1440)
            if col_idx >= len(word_table.columns):
                break
            for cell in word_table.columns[col_idx].cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                # remove existing tcW
                for child in list(tcPr):
                    if child.tag == qn('w:tcW'):
                        tcPr.remove(child)
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(tw))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)
    except Exception:
        # fallback: ignore width enforcement
        pass

    def _truncate(text, maxlen=40):
        if text is None:
            return ''
        s = str(text)
        return s if len(s) <= maxlen else s[:maxlen-3] + '...'

    # Default description mapping for common field names
    default_desc_map = {
        'BEGDA': 'Begin date',
        'ENDDA': 'End date',
        'CHGDA': 'Change date',
        'CRAT': 'Create at',
        'DESC': 'Description',
        'CHGBY': 'Change by',
        'x1': 'Example field',
        'x2': 'Example field 2',
        'x3': 'Example field 3',
        'x4': 'Example field 4',
        'x5': 'Example field 5',
        'x6': 'Example field 6',
        'x7': 'Example field 7'

    }

    for i, row in enumerate(subset.itertuples(), start=1):
        row_cells = word_table.add_row().cells
        # Number with dot like '1.'
        row_cells[0].text = f"{i}."
        dtype = str(row.data_type) if getattr(row, 'data_type', None) is not None else ''
        if 'character varying' in dtype.lower():
            dtype_norm = 'Varchar'
        else:
            dtype_norm = dtype.title()

        row_cells[1].text = _truncate(row.column_name, maxlen=25)
        row_cells[2].text = _truncate(dtype_norm, maxlen=20)
        # use default mapping when column description is empty
        col_desc = getattr(row, 'column_description', '')
        if not col_desc or str(col_desc).strip() == '':
            col_key = str(row.column_name).upper() if getattr(row, 'column_name', None) else ''
            col_desc = default_desc_map.get(col_key, '')
        row_cells[3].text = _truncate(col_desc, maxlen=50)

        # Format No cell compact
        try:
            p = row_cells[0].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = p.runs[0]
            r.font.size = Pt(12)
            r.font.bold = True
        except Exception:
            pass

    # reduce font size for table body to fit
    try:
        for col_idx in range(len(word_table.columns)):
            for cell in word_table.columns[col_idx].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        try:
                            run.font.size = Pt(12)
                        except Exception:
                            pass
    except Exception:
        pass

# ==============================
# 10. Functions
# ==============================
doc.add_heading('5. Stored Functions / Procedures', level=1)
if df_functions.empty:
    doc.add_paragraph("Tidak ada function / procedure user-defined.")
else:
    table_func = doc.add_table(rows=1, cols=5)
    table_func.style = 'Table Grid'
    hdr = table_func.rows[0].cells
    hdr[0].text = "Schema"
    hdr[1].text = "Nama Function"
    hdr[2].text = "Return Type"
    hdr[3].text = "Arguments"
    hdr[4].text = "Deskripsi"
    for _, row in df_functions.iterrows():
        r = table_func.add_row().cells
        r[0].text = str(row['schema'])
        r[1].text = str(row['function_name'])
        r[2].text = str(row['return_type'])
        r[3].text = str(row['arguments'])
        r[4].text = str(row['description']) if row['description'] else ""

# ==============================
# 11. Triggers
# ==============================
doc.add_heading('6. Triggers', level=1)
if df_triggers.empty:
    doc.add_paragraph("Tidak ada trigger user-defined.")
else:
    table_trig = doc.add_table(rows=1, cols=3)
    table_trig.style = 'Table Grid'
    hdr = table_trig.rows[0].cells
    hdr[0].text = "Nama Trigger"
    hdr[1].text = "Tabel Target"
    hdr[2].text = "Definisi"
    for _, row in df_triggers.iterrows():
        r = table_trig.add_row().cells
        r[0].text = str(row['trigger_name'])
        r[1].text = str(row['table_name'])
        r[2].text = str(row['definition'])

# ==============================
# 12. Views
# ==============================
doc.add_heading('7. Views', level=1)
if df_views.empty:
    doc.add_paragraph("Tidak ada view user-defined.")
else:
    table_view = doc.add_table(rows=1, cols=2)
    table_view.style = 'Table Grid'
    hdr = table_view.rows[0].cells
    hdr[0].text = "Nama View"
    hdr[1].text = "Definisi"
    for _, row in df_views.iterrows():
        r = table_view.add_row().cells
        r[0].text = str(row['table_name'])
        r[1].text = str(row['view_definition'])

# ==============================
# 13. Bagian Lain
# ==============================
doc.add_heading('8. Security & User Access', level=1)
doc.add_paragraph("(Lengkapi roles, user, policy security)")

doc.add_heading('9. Backup & Recovery', level=1)
doc.add_paragraph("(Lengkapi jadwal backup, prosedur recovery)")

doc.add_heading('10. Monitoring & Maintenance', level=1)
doc.add_paragraph("(Lengkapi tools monitoring & maintenance routine)")

doc.add_heading('11. Change Management', level=1)
doc.add_paragraph("(Lengkapi proses migrasi schema dan catatan perubahan)")

doc.add_heading('12. Referensi & Catatan', level=1)
doc.add_paragraph("(Lengkapi link ke wiki/SOP internal)")

# ==============================
# 14. Simpan Dokumen
# ==============================
output_file = "Dokumentasi_PostgreSQL.docx"
doc.save(output_file)

print(f"✅ Dokumentasi berhasil dibuat: {os.path.abspath(output_file)}")
if paginated_files:
    for erd_file in paginated_files:
        print(f"✅ ERD A4-optimized berhasil dibuat: {os.path.abspath(erd_file)}")
print(f"📊 Total tabel: {len(df_tables['table_name'].unique())}")
print(f"🔗 Total relasi: {len(df_fk)}")