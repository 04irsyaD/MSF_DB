"""Template TSD wajib berupa kerangka, bukan dokumen terisi."""

from pathlib import Path

import pytest

TEMPLATE = Path(__file__).resolve().parent.parent / "templates" / "tsd.docx"


@pytest.fixture(scope="module")
def dokumen():
    if not TEMPLATE.exists():
        pytest.skip("templates/tsd.docx belum dibuat")
    from docx import Document

    return Document(str(TEMPLATE))


def _teks_penuh(dokumen) -> str:
    bagian = [p.text for p in dokumen.paragraphs]
    for tabel in dokumen.tables:
        for baris in tabel.rows:
            bagian.extend(sel.text for sel in baris.cells)
    return "\n".join(bagian)


def test_ukuran_template_wajar():
    """
    Sumbernya 7,7 MB. Template yang masih membawa gambar berat akan
    memperbesar setiap dokumen hasil generate, bukan hanya dirinya sendiri.
    """
    if not TEMPLATE.exists():
        pytest.skip("templates/tsd.docx belum dibuat")

    assert TEMPLATE.stat().st_size < 2_000_000


def test_seksi_per_tabel_hanya_tersisa_satu(dokumen):
    """109 seksi berulang harus menyusut menjadi satu badan perulangan."""
    header_field = [
        t
        for t in dokumen.tables
        if t.rows
        and [s.text.strip() for s in t.rows[0].cells][:2] == ["No", "Nama Field"]
    ]

    assert len(header_field) == 1


def test_tabel_induk_memakai_tiga_baris_perulangan(dokumen):
    """
    Perulangan baris docxtpl butuh tiga baris data: tag for, isi, tag endfor.
    Seluruh baris yang memuat tag akan digantikan oleh tag itu sendiri, jadi
    menaruh ketiganya pada satu baris membuat baris isi ikut lenyap.
    """
    induk = [
        t
        for t in dokumen.tables
        if t.rows
        and [s.text.strip() for s in t.rows[0].cells][:2] == ["No", "Nama Tabel"]
    ]

    assert len(induk) == 1
    assert len(induk[0].rows) == 4


def test_tag_perulangan_terpasang(dokumen):
    teks = _teks_penuh(dokumen)

    assert "{%tr for c in t.columns %}" in teks
    assert "{%tr for t in tables %}" in teks
    assert "{%tr endfor %}" in teks
    assert "{%p for t in tables %}" in teks
    assert "{%p endfor %}" in teks


def test_variabel_kolom_terpasang(dokumen):
    teks = _teks_penuh(dokumen)

    for variabel in ("c.no", "c.name", "c.data_type_label", "c.description"):
        assert "{{ " + variabel + " }}" in teks


def test_judul_seksi_menjadi_dinamis(dokumen):
    teks = _teks_penuh(dokumen)

    assert "3.4.{{ loop.index }} Tabel {{ t.name }}" in teks


def test_tag_tidak_pecah_antar_run(dokumen):
    """
    Tag yang terbelah beberapa run membuat render gagal. Penyisipan lewat
    python-docx mencegahnya; test ini menjaga agar penyuntingan manual di
    Word kelak tidak merusaknya tanpa ketahuan.
    """
    for paragraf in dokumen.paragraphs:
        if "{%" in paragraf.text or "{{" in paragraf.text:
            utuh = any("{%" in run.text or "{{" in run.text for run in paragraf.runs)

            assert utuh, f"tag pecah antar run: {paragraf.text[:60]}"


def test_data_nyata_sudah_dibuang(dokumen):
    """Nama tabel dari dokumen sumber tidak boleh tertinggal di template."""
    teks = _teks_penuh(dokumen)

    for nama in ("t_ChainAnalysis", "t_BusinessUnit", "m_grisklist"):
        assert nama not in teks


def test_seksi_view_dibuang(dokumen):
    """Generator tidak menghasilkan dokumentasi view."""
    teks = _teks_penuh(dokumen)

    assert "V_gris_irislist" not in teks


def test_template_benar_benar_dapat_dirender():
    """
    Test struktur saja tidak membuktikan template berfungsi. Tag yang salah
    susun tetap lolos pemeriksaan struktur, tetapi menjatuhkan render dengan
    'Encountered unknown tag endfor'.
    """
    if not TEMPLATE.exists():
        pytest.skip("templates/tsd.docx belum dibuat")

    import io
    from dataclasses import dataclass, field
    from typing import List

    from docx import Document
    from docxtpl import DocxTemplate
    from jinja2.sandbox import SandboxedEnvironment

    @dataclass
    class KolomUji:
        no: int
        name: str
        data_type_label: str
        description: str

    @dataclass
    class TabelUji:
        name: str
        summary: str
        columns: List[KolomUji] = field(default_factory=list)

    tabel = [
        TabelUji(
            "mst_nasabah",
            "Data induk nasabah.",
            [KolomUji(1, "id", "bigint", "Nomor identitas.")],
        ),
        TabelUji(
            "trx_pembayaran",
            "Transaksi pembayaran.",
            [
                KolomUji(1, "id", "bigint", "Identitas transaksi."),
                KolomUji(2, "jumlah", "numeric", "Nominal pembayaran."),
            ],
        ),
    ]

    dokumen = DocxTemplate(str(TEMPLATE))
    dokumen.render(
        {
            "project_name": "Proyek Uji",
            "generated_at": "11 August 2026",
            "author": "Penguji",
            "tables": tabel,
        },
        SandboxedEnvironment(autoescape=False),
    )
    penyangga = io.BytesIO()
    dokumen.save(penyangga)

    hasil = Document(io.BytesIO(penyangga.getvalue()))
    judul = [
        p.text for p in hasil.paragraphs if p.style.name == "Heading 3" and "3.4." in p.text
    ]
    tabel_field = [
        t
        for t in hasil.tables
        if t.rows
        and [s.text.strip() for s in t.rows[0].cells][:2] == ["No", "Nama Field"]
    ]

    assert judul == ["3.4.1 Tabel mst_nasabah", "3.4.2 Tabel trx_pembayaran"]
    assert len(tabel_field) == 2
    assert len(tabel_field[1].rows) == 3
    assert tabel_field[1].rows[2].cells[3].text == "Nominal pembayaran."
