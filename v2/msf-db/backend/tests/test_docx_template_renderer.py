"""Renderer mengisi template; keamanannya bertumpu pada sandbox dan enum."""

import io

import pytest
from docx import Document

from app.models.schemas import StructureTemplate
from app.services.doc_model import SUMBER_AI, ColumnDoc, DocumentModel, TableDoc
from app.services.renderers import docx_template_renderer as modul
from app.services.renderers.docx_template_renderer import render_docx


def model_contoh():
    return DocumentModel(
        project_name="Proyek Uji",
        generated_at="11 August 2026",
        tables=[
            TableDoc(
                name="t_access",
                summary="Menyimpan hak akses.",
                columns=[
                    ColumnDoc(
                        no=1,
                        name="id",
                        data_type_label="bigint",
                        description="Nomor identitas.",
                        source=SUMBER_AI,
                    )
                ],
            )
        ],
    )


def _template_kecil(tujuan, isi_tambahan=""):
    """
    Template fixture, bukan template TSD asli, agar test tetap cepat.

    Perulangan baris memakai tiga baris data karena docxtpl mengganti
    seluruh baris yang memuat tag dengan tag itu sendiri.
    """
    dokumen = Document()
    dokumen.add_paragraph("{{ project_name }}")
    dokumen.add_paragraph("{{ generated_at }}")
    if isi_tambahan:
        dokumen.add_paragraph(isi_tambahan)

    tabel = dokumen.add_table(rows=4, cols=2)
    tabel.rows[0].cells[0].text = "No"
    tabel.rows[0].cells[1].text = "Nama Tabel"
    tabel.rows[1].cells[0].text = "{%tr for t in tables %}"
    tabel.rows[2].cells[0].text = "{{ loop.index }}"
    tabel.rows[2].cells[1].text = "{{ t.name }}"
    tabel.rows[3].cells[0].text = "{%tr endfor %}"
    dokumen.save(str(tujuan))


@pytest.fixture
def templates_dir(tmp_path, monkeypatch):
    monkeypatch.setattr(modul, "TEMPLATES_DIR", str(tmp_path))
    return tmp_path


def _teks(data: bytes) -> str:
    dokumen = Document(io.BytesIO(data))
    bagian = [p.text for p in dokumen.paragraphs]
    for tabel in dokumen.tables:
        for baris in tabel.rows:
            bagian.extend(sel.text for sel in baris.cells)
    return "\n".join(bagian)


def test_variabel_dan_perulangan_terisi(templates_dir):
    _template_kecil(templates_dir / "tsd.docx")

    hasil = render_docx(model_contoh(), StructureTemplate.MSF_TSD)

    teks = _teks(hasil)
    assert "Proyek Uji" in teks
    assert "t_access" in teks


def test_template_tidak_ada_memberi_pesan_jelas(templates_dir):
    with pytest.raises(FileNotFoundError) as info:
        render_docx(model_contoh(), StructureTemplate.MSF_TSD)

    assert "tsd.docx" in str(info.value)


def test_template_tanpa_berkas_terpetakan_ditolak(templates_dir):
    """Nilai enum yang belum punya berkas tidak boleh jatuh ke path tebakan."""
    with pytest.raises(ValueError):
        render_docx(model_contoh(), StructureTemplate.STANDARD)


def test_ekspresi_berbahaya_diblokir_sandbox(templates_dir):
    """
    docxtpl menjalankan Jinja2. Tanpa sandbox, ekspresi di dalam template
    dapat menjangkau atribut internal Python.
    """
    from jinja2.exceptions import SecurityError

    _template_kecil(
        templates_dir / "tsd.docx",
        isi_tambahan="{{ project_name.__class__.__mro__ }}",
    )

    with pytest.raises(SecurityError):
        render_docx(model_contoh(), StructureTemplate.MSF_TSD)


def test_nilai_request_tidak_pernah_menyentuh_path():
    """Pemetaan enum ke nama berkas wajib konstan di kode."""
    import inspect

    sumber = inspect.getsource(modul)

    assert "os.path.join(TEMPLATES_DIR" not in sumber
    assert "_BERKAS_TEMPLATE" in sumber


def test_penulis_masuk_konteks(templates_dir):
    _template_kecil(templates_dir / "tsd.docx", isi_tambahan="{{ author }}")

    hasil = render_docx(model_contoh(), StructureTemplate.MSF_TSD, author="Irsyad")

    assert "Irsyad" in _teks(hasil)


def test_template_asli_dapat_dirender(templates_dir):
    """
    Fixture kecil tidak membuktikan template TSD sungguhan berfungsi.
    Renderer diuji sekali terhadap berkas yang benar-benar dipakai produksi.
    """
    from pathlib import Path
    import shutil

    asli = Path(__file__).resolve().parent.parent / "templates" / "tsd.docx"
    if not asli.exists():
        pytest.skip("templates/tsd.docx belum dibuat")

    shutil.copy(asli, templates_dir / "tsd.docx")

    hasil = render_docx(model_contoh(), StructureTemplate.MSF_TSD, author="Penguji")

    teks = _teks(hasil)
    assert "3.4.1 Tabel t_access" in teks
    assert "Nomor identitas." in teks
