"""
DOCX Exporter — konversi Markdown ke Word (.docx).
"""

import re
import os
from io import BytesIO
from typing import Optional
from datetime import datetime
import structlog

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = structlog.get_logger()

TEMPLATES_DIR = os.getenv("TEMPLATES_DIR", "/app/templates")


class DocxExporter:
    """
    Konversi Markdown dokumentasi ke file Word (.docx).
    Menggunakan python-docx untuk generate dokumen yang rapi.
    """

    def __init__(self, template_path: Optional[str] = None):
        self.template_path = template_path

    def export(
        self,
        markdown_content: str,
        project_name: str = "Database Documentation",
        author: Optional[str] = None,
    ) -> bytes:
        """
        Konversi markdown ke DOCX, return bytes.
        """
        doc = Document()
        self._setup_document(doc, project_name, author)
        self._parse_and_write(doc, markdown_content)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _setup_document(self, doc: Document, project_name: str, author: Optional[str]):
        """Setup style dan properti dokumen"""
        # Page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)

        # Document properties
        core_props = doc.core_properties
        core_props.title = project_name
        if author:
            core_props.author = author
        core_props.created = datetime.now()
        core_props.subject = "Database Documentation"
        core_props.keywords = "database, documentation, MSF-APP"

        # Setup styles
        self._setup_styles(doc)

    def _setup_styles(self, doc: Document):
        """Setup custom styles"""
        styles = doc.styles

        # Normal text style
        try:
            normal = styles["Normal"]
            normal.font.name = "Calibri"
            normal.font.size = Pt(11)
        except Exception:
            pass

        # Heading 1
        try:
            h1 = styles["Heading 1"]
            h1.font.name = "Calibri"
            h1.font.size = Pt(20)
            h1.font.bold = True
            h1.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # Biru tua
        except Exception:
            pass

        # Heading 2
        try:
            h2 = styles["Heading 2"]
            h2.font.name = "Calibri"
            h2.font.size = Pt(16)
            h2.font.bold = True
            h2.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        except Exception:
            pass

        # Heading 3
        try:
            h3 = styles["Heading 3"]
            h3.font.name = "Calibri"
            h3.font.size = Pt(13)
            h3.font.bold = True
            h3.font.color.rgb = RGBColor(0x37, 0x63, 0xD0)
        except Exception:
            pass

    def _parse_and_write(self, doc: Document, markdown: str):
        """Parse markdown line by line dan tulis ke dokumen Word"""
        lines = markdown.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            # Heading 1
            if line.startswith("# ") and not line.startswith("## "):
                doc.add_heading(line[2:].strip(), level=1)
                i += 1
                continue

            # Heading 2
            if line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
                i += 1
                continue

            # Heading 3
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
                i += 1
                continue

            # Heading 4
            if line.startswith("#### "):
                p = doc.add_paragraph(line[5:].strip())
                if p.runs:
                    p.runs[0].bold = True
                i += 1
                continue

            # Horizontal rule
            if line.strip() in ("---", "***", "___"):
                self._add_horizontal_rule(doc)
                i += 1
                continue

            # Table (baris dimulai dengan |)
            if line.strip().startswith("|") and i + 1 < len(lines):
                # Kumpulkan baris tabel
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(doc, table_lines)
                continue

            # Bullet list
            if line.strip().startswith("- ") or line.strip().startswith("* "):
                text = line.strip()[2:]
                p = doc.add_paragraph(style="List Bullet")
                self._add_formatted_run(p, text)
                i += 1
                continue

            # Blockquote (⚠️ warning)
            if line.strip().startswith(">"):
                text = line.strip()[1:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(text)
                run.font.color.rgb = RGBColor(0x92, 0x40, 0x00)
                run.font.italic = True
                i += 1
                continue

            # Baris kosong
            if not line.strip():
                i += 1
                continue

            # Paragraf biasa
            if line.strip():
                p = doc.add_paragraph()
                self._add_formatted_run(p, line.strip())

            i += 1

    def _add_formatted_run(self, paragraph, text: str):
        """Tambah teks dengan format bold/italic/code"""
        # Split berdasarkan format markers
        parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_)", text)

        for part in parts:
            if not part:
                continue

            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif (part.startswith("*") and part.endswith("*")) or \
                 (part.startswith("_") and part.endswith("_")):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            else:
                paragraph.add_run(part)

    def _add_table(self, doc: Document, table_lines: list):
        """Parse dan tambah tabel Markdown ke Word"""
        if len(table_lines) < 2:
            return

        # Parse header
        header_line = table_lines[0]
        headers = [cell.strip() for cell in header_line.split("|") if cell.strip()]

        # Skip separator line (|---|---|)
        data_lines = [l for l in table_lines[2:] if l.strip() and not all(
            c in "-|: " for c in l
        )]

        if not headers:
            return

        # Buat tabel Word
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header.strip("*_"))
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # Putih
            # Background header
            self._set_cell_background(cell, "1E40AF")  # Biru

        # Data rows
        for row_line in data_lines:
            cells_data = [cell.strip() for cell in row_line.split("|") if cell.strip() or row_line.count("|") > 1]
            # Filter empty cells at start/end dari split
            cells_data = [c for c in [cell.strip() for cell in row_line.split("|")] if c != "" or True]
            # Ambil hanya yang relevan
            raw_cells = row_line.split("|")
            # Hapus elemen pertama dan terakhir jika kosong (dari | di awal/akhir)
            if raw_cells and raw_cells[0].strip() == "":
                raw_cells = raw_cells[1:]
            if raw_cells and raw_cells[-1].strip() == "":
                raw_cells = raw_cells[:-1]

            if not raw_cells:
                continue

            row_cells = table.add_row().cells
            for j, cell_text in enumerate(raw_cells):
                if j < len(row_cells):
                    row_cells[j].text = ""
                    p = row_cells[j].paragraphs[0]
                    self._add_formatted_run(p, cell_text.strip())

        # Tambah spasi setelah tabel
        doc.add_paragraph()

    def _set_cell_background(self, cell, hex_color: str):
        """Set background color pada cell"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _add_horizontal_rule(self, doc: Document):
        """Tambah garis horizontal sebagai paragraph border"""
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)


# Singleton
docx_exporter = DocxExporter()
