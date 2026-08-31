"""
Ubah dokumen TSD terisi menjadi template docxtpl.

Dijalankan manual, sekali, atau diulang bila dokumen sumber diperbarui.
Dokumen sumber TIDAK pernah diubah; keluaran ditulis ke berkas terpisah.

Pendekatannya mengisi template aslinya, bukan membangun ulang tampilannya.
Percobaan v1 di AI OLLMA/summary/scripts/ menghasilkan 43 skrip karena
menempuh jalur replikasi: mencocokkan font, border, dan lebar kolom satu per
satu lewat XML mentah, pekerjaan yang tidak pernah konvergen.

Seluruh tag Jinja disisipkan lewat python-docx sehingga selalu mendarat dalam
satu run. Tag yang diketik manusia di Word dapat terbelah beberapa run akibat
autocorrect, dan tag yang terbelah membuat render gagal.

Pakai:
    python scripts/siapkan_template_tsd.py <sumber.docx> <keluaran.docx> [--buang-gambar-besar N]
"""

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

HEADER_FIELD = ["No", "Nama Field", "Tipe Data", "Deskripsi Field"]
HEADER_INDUK = ["No", "Nama Tabel", "Deskripsi Tabel"]

TAG_MULAI_TABEL = "{%p for t in tables %}"
TAG_SELESAI_TABEL = "{%p endfor %}"

POLA_TABEL = re.compile(r"^3\.4\.\d+")
POLA_VIEW = re.compile(r"^3\.5\.\d+")


def _header(tabel: Table) -> list:
    if not tabel.rows:
        return []
    return [sel.text.strip() for sel in tabel.rows[0].cells]


def _buang(elemen) -> None:
    elemen.getparent().remove(elemen)


def _tulis_sel(sel, teks: str) -> None:
    """Menulis satu run utuh sehingga tag Jinja tidak pernah pecah."""
    sel.text = ""
    sel.paragraphs[0].add_run(teks)


def _pasang_perulangan_baris(tabel: Table, tag_for: str, isi_sel: list) -> None:
    """
    Susun perulangan baris docxtpl memakai tiga baris data.

    Ini bukan pilihan gaya. Pada docxtpl, SELURUH baris yang memuat tag
    {%tr ...%} digantikan oleh tag tersebut, jadi barisnya tidak bertahan.
    Menaruh for, isi, dan endfor pada satu baris yang sama membuat baris isi
    ikut lenyap, dan Jinja menemukan endfor tanpa pasangan.

    Susunannya karena itu: baris tag for, baris isi, baris tag endfor.
    """
    jumlah_dibutuhkan = 4  # header + tiga baris data
    while len(tabel.rows) > jumlah_dibutuhkan:
        _buang(tabel.rows[-1]._tr)
    while len(tabel.rows) < jumlah_dibutuhkan:
        tabel.add_row()

    baris_for, baris_isi, baris_endfor = tabel.rows[1], tabel.rows[2], tabel.rows[3]

    for sel in baris_for.cells:
        _tulis_sel(sel, "")
    _tulis_sel(baris_for.cells[0], tag_for)

    for indeks, sel in enumerate(baris_isi.cells):
        _tulis_sel(sel, isi_sel[indeks] if indeks < len(isi_sel) else "")

    for sel in baris_endfor.cells:
        _tulis_sel(sel, "")
    _tulis_sel(baris_endfor.cells[0], "{%tr endfor %}")


def _paragraf_baru(teks: str):
    """Paragraf polos berisi satu run, untuk tag tingkat paragraf docxtpl."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = teks
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


def _gaya_paragraf(dokumen, elemen) -> str:
    if elemen.tag != qn("w:p"):
        return ""
    return Paragraph(elemen, dokumen).style.name


def _teks_paragraf(dokumen, elemen) -> str:
    if elemen.tag != qn("w:p"):
        return ""
    return Paragraph(elemen, dokumen).text.strip()


def _rentang_seksi(dokumen, pola_penomoran) -> tuple:
    """
    Kembalikan (indeks_mulai, indeks_akhir) untuk satu seksi bernomor.

    Deteksi berbasis rentang, bukan kedekatan antar elemen. Urutan elemen di
    dokumen sumber TIDAK seragam: dari 109 tabel field, hanya 61 yang berada
    tepat sesudah heading; 46 didahului paragraf kosong dan 2 didahului
    caption. Asumsi kedekatan karena itu menyisakan puluhan tabel yatim.
    """
    anak = list(dokumen.element.body)
    mulai = None
    for indeks, elemen in enumerate(anak):
        if _gaya_paragraf(dokumen, elemen) == "Heading 3" and pola_penomoran.match(
            _teks_paragraf(dokumen, elemen)
        ):
            mulai = indeks
            break
    if mulai is None:
        return None, None

    akhir = len(anak)
    for indeks in range(mulai + 1, len(anak)):
        if _gaya_paragraf(dokumen, anak[indeks]) in ("Heading 1", "Heading 2"):
            akhir = indeks
            break
    return mulai, akhir


def _buang_toc_lama(dokumen) -> int:
    """
    Buang entri daftar isi tingkat tiga.

    Isinya adalah hasil field yang ter-cache dan memuat nama tabel dari
    dokumen sumber. Membiarkannya berarti setiap dokumen hasil generate
    membawa daftar isi milik sistem lain sampai pengguna menekan F9.
    """
    dibuang = 0
    for paragraf in list(dokumen.paragraphs):
        if paragraf.style.name in ("toc 3", "toc 4"):
            _buang(paragraf._p)
            dibuang += 1
    return dibuang


def _inventaris_gambar(sumber: Path) -> list:
    with zipfile.ZipFile(sumber) as arsip:
        return sorted(
            (
                (info.filename, info.file_size)
                for info in arsip.infolist()
                if info.filename.startswith("word/media/")
            ),
            key=lambda pasangan: pasangan[1],
            reverse=True,
        )


def _buang_gambar_besar(dokumen, batas_byte: int) -> int:
    """
    Buang elemen gambar yang berkas medianya melebihi batas.

    Gambar besar pada dokumen sumber adalah tangkapan layar diagram ERD milik
    sistem lain. Membiarkannya berarti setiap dokumen hasil generate membawa
    diagram yang tidak ada hubungannya, sekaligus beban ukuran berkas.

    Relasinya WAJIB ikut dilepas. Menghapus elemen gambar saja tidak cukup:
    berkas medianya tetap ikut tersimpan di dalam paket karena masih
    dirujuk relasi, sehingga ukuran berkas nyaris tidak berkurang.
    """
    bagian = dokumen.part
    dibuang = 0
    rid_dilepas = []

    # Ditelusuri lewat elemen blip, bukan lewat inline_shapes. Gambar terbesar
    # pada dokumen sumber berupa drawing anchored alias mengambang, dan
    # inline_shapes tidak pernah melihat jenis itu.
    for blip in list(bagian.element.body.iter(qn("a:blip"))):
        rid = blip.get(qn("r:embed"))
        if not rid:
            continue
        try:
            if len(bagian.related_parts[rid].blob) <= batas_byte:
                continue
        except KeyError:
            continue

        induk = blip
        while induk is not None and induk.tag != qn("w:drawing"):
            induk = induk.getparent()
        if induk is None:
            continue

        # Buang run pembungkusnya bila ada, agar tidak menyisakan run kosong.
        sasaran = induk.getparent() if induk.getparent().tag == qn("w:r") else induk
        _buang(sasaran)
        rid_dilepas.append(rid)
        dibuang += 1

    for rid in rid_dilepas:
        try:
            bagian.drop_rel(rid)
        except Exception:
            continue
    return dibuang


def siapkan(sumber: Path, keluaran: Path, batas_gambar_mb: float = 0) -> None:
    if not sumber.exists():
        raise SystemExit(f"Dokumen sumber tidak ditemukan: {sumber}")

    print(f"Membaca {sumber} ({sumber.stat().st_size / 1_048_576:.1f} MB)")
    for nama, ukuran in _inventaris_gambar(sumber)[:5]:
        print(f"  gambar terbesar: {nama} {ukuran / 1_048_576:.2f} MB")

    dokumen = Document(str(sumber))

    tabel_field = [t for t in dokumen.tables if _header(t) == HEADER_FIELD]
    tabel_induk = [t for t in dokumen.tables if _header(t) == HEADER_INDUK]
    if not tabel_field or not tabel_induk:
        raise SystemExit(
            "Struktur tabel tidak dikenali. Periksa apakah dokumen sumber "
            "masih memakai header 'No | Nama Field | Tipe Data | Deskripsi Field'."
        )

    print(f"Ditemukan {len(tabel_field)} tabel field, {len(tabel_induk)} tabel induk.")

    # Seksi 3.5 Deskripsi Field View dibuang seluruhnya. Generator tidak
    # menghasilkan dokumentasi view, sehingga membiarkannya berarti setiap
    # dokumen hasil membawa daftar view milik sistem lain.
    mulai_view, akhir_view = _rentang_seksi(dokumen, POLA_VIEW)
    if mulai_view is not None:
        anak = list(dokumen.element.body)
        # Heading 2 pembuka seksi view berada tepat sebelum sub-seksi pertama.
        awal_buang = mulai_view - 1 if mulai_view > 0 else mulai_view
        for elemen in anak[awal_buang:akhir_view]:
            _buang(elemen)
        print(f"  seksi 3.5 Deskripsi Field View dibuang ({akhir_view - awal_buang} elemen).")

    mulai, akhir = _rentang_seksi(dokumen, POLA_TABEL)
    if mulai is None:
        raise SystemExit("Tidak ada seksi 3.4.x yang dikenali.")

    anak = list(dokumen.element.body)
    contoh_heading = Paragraph(anak[mulai], dokumen)
    contoh_tabel = None
    dibuang = 0
    for elemen in anak[mulai + 1 : akhir]:
        if contoh_tabel is None and elemen.tag == qn("w:tbl"):
            calon = Table(elemen, dokumen)
            if _header(calon) == HEADER_FIELD:
                contoh_tabel = calon
                continue
        _buang(elemen)
        dibuang += 1

    if contoh_tabel is None:
        raise SystemExit("Tabel field contoh tidak ditemukan di dalam seksi 3.4.")

    print(f"  {dibuang} elemen seksi berulang dibuang, satu seksi contoh disisakan.")

    # Heading contoh menjadi judul dinamis.
    contoh_heading.text = ""
    contoh_heading.add_run("3.4.{{ loop.index }} Tabel {{ t.name }}")

    # Tabel field contoh menjadi badan perulangan kolom.
    _pasang_perulangan_baris(
        contoh_tabel,
        "{%tr for c in t.columns %}",
        ["{{ c.no }}", "{{ c.name }}", "{{ c.data_type_label }}", "{{ c.description }}"],
    )

    # Bungkus seksi contoh dengan perulangan tingkat paragraf.
    contoh_heading._p.addprevious(_paragraf_baru(TAG_MULAI_TABEL))
    contoh_tabel._tbl.addnext(_paragraf_baru(TAG_SELESAI_TABEL))

    # Tabel induk: satu baris data sebagai badan perulangan.
    _pasang_perulangan_baris(
        tabel_induk[0],
        "{%tr for t in tables %}",
        ["{{ loop.index }}", "{{ t.name }}", "{{ t.summary }}"],
    )

    dibuang_toc = _buang_toc_lama(dokumen)
    print(f"  {dibuang_toc} entri daftar isi lama dibuang.")

    if batas_gambar_mb > 0:
        jumlah = _buang_gambar_besar(dokumen, int(batas_gambar_mb * 1_048_576))
        print(f"  {jumlah} gambar melebihi {batas_gambar_mb} MB dibuang.")

    keluaran.parent.mkdir(parents=True, exist_ok=True)
    dokumen.save(str(keluaran))
    print(
        f"Template ditulis ke {keluaran} "
        f"({keluaran.stat().st_size / 1_048_576:.2f} MB)"
    )


def _argumen(argv: list) -> tuple:
    if len(argv) < 3:
        raise SystemExit(__doc__)
    batas = 0.0
    if "--buang-gambar-besar" in argv:
        batas = float(argv[argv.index("--buang-gambar-besar") + 1])
    return Path(argv[1]), Path(argv[2]), batas


if __name__ == "__main__":
    sumber_arg, keluaran_arg, batas_arg = _argumen(sys.argv)
    siapkan(sumber_arg, keluaran_arg, batas_arg)
